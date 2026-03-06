"""
Test script to verify equation extraction from documents
"""
import requests
import json
from docx import Document

# Configuration
BASE_URL = "http://127.0.0.1:8000/api/v1"
TEST_EMAIL = "test@example.com"
TEST_PASSWORD = "testpass123"

def create_sample_docx_with_equations():
    """Create a sample DOCX with equations for testing"""
    doc = Document()
    doc.add_heading('Test Document with Equations', 0)

    doc.add_paragraph('This document contains mathematical equations.')

    # Add some equation-like text
    doc.add_paragraph('Equation 1: E = mc^2 (6-1)')
    doc.add_paragraph('Equation 2: F = ma (6-2)')
    doc.add_paragraph('Equation 3: ax+v + b^2 = c^2 (6-3)')
    doc.add_paragraph('Equation 4: Integral f(x)dx = F(x) + C (6-4)')
    doc.add_paragraph('Equation 5: Sum(x_i) / n = mu (6-5)')

    doc.add_paragraph('Some regular text here.')

    filename = 'file_sample.docx'
    doc.save(filename)
    print(f"[OK] Created test document: {filename}")
    return filename

def register_and_login():
    """Register and login to get access token"""
    print("\n1. Registering user...")

    # Register
    response = requests.post(f"{BASE_URL}/auth/register", json={
        "email": TEST_EMAIL,
        "password": TEST_PASSWORD
    })

    if response.status_code == 200:
        print("[why is this passowrd user ] User registered successfully")
    elif response.status_code == 400:
        print("[INFO] User already exists, continuing...")
    else:
        print(f"[ERROR] Registration failed: {response.text}")
        return None

    # Login
    print("\n2. Logging in...")
    response = requests.post(f"{BASE_URL}/auth/login", json={
        "email": TEST_EMAIL,
        "password": TEST_PASSWORD
    })

    if response.status_code == 200:
        token = response.json()["access_token"]
        print("[OK] Login successful")
        return token
    else:
        print(f"[ERROR] Login failed: {response.text}")
        return None

def upload_document(token, filepath):
    """Upload document"""
    print(f"\n3. Uploading document: {filepath}")

    headers = {"Authorization": f"Bearer {token}"}

    with open(filepath, 'rb') as f:
        files = {'file': (filepath, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        response = requests.post(f"{BASE_URL}/upload/", files=files, headers=headers)

    if response.status_code == 200:
        data = response.json()
        document_id = data['document_id']
        print(f"[OK] Document uploaded successfully")
        print(f"  Document ID: {document_id}")
        print(f"  Filename: {data['filename']}")
        print(f"  Status: {data['status']}")
        return document_id
    else:
        print(f"[ERROR] Upload failed: {response.text}")
        return None

def process_document(token, document_id):
    """Process the uploaded document"""
    print(f"\n4. Processing document (ID: {document_id})...")
    print("   Note: First run will download Nougat model (~1.5GB), may take 1-2 minutes")

    headers = {"Authorization": f"Bearer {token}"}
    response = requests.post(f"{BASE_URL}/documents/{document_id}/process", headers=headers)

    if response.status_code == 200:
        data = response.json()
        print(f"[OK] Processing completed")
        print(f"  Status: {data['status']}")
        print(f"  Message: {data['message']}")
        return True
    else:
        print(f"[ERROR] Processing failed: {response.text}")
        return False

def get_document_details(token, document_id):
    """Get document details including extracted equations"""
    print(f"\n5. Fetching document details...")

    headers = {"Authorization": f"Bearer {token}"}
    response = requests.get(f"{BASE_URL}/documents/{document_id}", headers=headers)

    if response.status_code == 200:
        data = response.json()
        print(f"\n{'='*60}")
        print(f"DOCUMENT EXTRACTION RESULTS")
        print(f"{'='*60}")

        print(f"\nDocument Info:")
        print(f"  - Filename: {data.get('original_filename')}")
        print(f"  - Status: {data.get('status')}")
        print(f"  - File Type: {data.get('file_type')}")

        print(f"\nMetadata:")
        metadata = data.get('metadata', {})
        print(f"  - Total Pages: {metadata.get('total_pages', 0)}")
        print(f"  - Total Paragraphs: {metadata.get('total_paragraphs', 0)}")
        print(f"  - Total Equations: {metadata.get('total_equations', 0)}")
        print(f"  - Total Figures: {metadata.get('total_figures', 0)}")
        print(f"  - Total Tables: {metadata.get('total_tables', 0)}")

        print(f"\nExtracted Equations ({len(data.get('equations', []))}):")
        equations = data.get('equations', [])
        if equations:
            for idx, eq in enumerate(equations, 1):
                print(f"\n  Equation {idx}:")
                print(f"    ID: {eq.get('equation_id')}")
                print(f"    LaTeX: {eq.get('latex')}")
                print(f"    Number: {eq.get('number', 'N/A')}")
        else:
            print("  No equations extracted")

        print(f"\nText Content Preview (first 500 chars):")
        text = data.get('text_content', '')
        print(f"  {text[:500]}...")

        print(f"\nFigures: {len(data.get('figures', []))}")
        print(f"Tables: {len(data.get('tables', []))}")

        print(f"\n{'='*60}")

        return data
    else:
        print(f"[ERROR] Failed to get document: {response.text}")
        return None

def main():
    print("="*60)
    print("EQUATION EXTRACTION TEST")
    print("="*60)

    # Step 1: Create test document
    docx_file = create_sample_docx_with_equations()

    # Step 2: Register & Login
    token = register_and_login()
    if not token:
        print("\n[ERROR] Authentication failed, exiting...")
        return

    # Step 3: Upload document
    document_id = upload_document(token, docx_file)
    if not document_id:
        print("\n[ERROR] Upload failed, exiting...")
        return

    # Step 4: Process document
    success = process_document(token, document_id)
    if not success:
        print("\n[ERROR] Processing failed, exiting...")
        return

    # Step 5: Get results
    document_data = get_document_details(token, document_id)

    print("\n[OK] Test completed!")
    print("\nTo test with your own document:")
    print("  1. Replace 'test_equations.docx' with your file")
    print("  2. Run: python test_equation_extraction.py")

if __name__ == "__main__":
    main()
