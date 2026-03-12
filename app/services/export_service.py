import os
import re
import copy
import unicodedata
from typing import Optional, List
from docx import Document as DocxDocument
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from app.core.logger import get_logger
from app.core.config import settings
from app.services.renumbering_service import RenumberingService

logger = get_logger(__name__)


# ------------------------------------------------------------------ #
# Low-level helpers                                                    #
# ------------------------------------------------------------------ #

def _add_highlight(run, color="yellow"):
    """Add highlight color to a run. Common colors: yellow, green, cyan."""
    rPr = run._element.get_or_add_rPr()
    highlight = OxmlElement("w:highlight")
    highlight.set(qn("w:val"), color)
    rPr.append(highlight)


def _get_p_text(p_element):
    """Extract plain text from a w:p XML element."""
    return "".join(node.text or "" for node in p_element.iter(qn("w:t")))


def _is_heading_element(p_element):
    """Check if a w:p element is a heading or special/large-font paragraph."""
    pPr = p_element.find(qn("w:pPr"))
    if pPr is None:
        return False
    # Check style name
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is not None:
        style_val = (pStyle.get(qn("w:val")) or "").lower()
        # Catch all heading-like styles: Heading1, Title, TOC, BoxTitle, ChapterNum, etc.
        heading_keywords = (
            "heading", "title", "toc", "caption", "subtitle",
            "box", "callout", "chapter", "figure", "table",
            "header", "footer", "quote", "note", "sidebar",
        )
        if any(kw in style_val for kw in heading_keywords):
            return True
    # Check outline level
    if pPr.find(qn("w:outlineLvl")) is not None:
        return True
    return False


def _get_font_size_half_points(p_element):
    """Get the font size (in half-points) from a paragraph's first run.

    Returns None if no explicit font size is set.
    Word default body text is typically 20-24 half-points (10-12pt).
    """
    for r in p_element.iter(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is not None:
            sz = rPr.find(qn("w:sz"))
            if sz is not None:
                try:
                    return int(sz.get(qn("w:val")))
                except (ValueError, TypeError):
                    pass
    return None


def _is_body_text_paragraph(para):
    """Check if a paragraph is regular body text (not heading, not special, reasonable font size)."""
    if not para.text.strip() or not para.runs:
        return False
    if _is_heading_element(para._element):
        return False
    # Check style name — prefer "Normal", "BodyText", or similar
    style_name = (para.style.name or "").lower() if para.style else ""
    if "normal" in style_name or "body" in style_name:
        return True
    # Check font size — body text is typically 10-14pt (20-28 half-points)
    font_size = _get_font_size_half_points(para._element)
    if font_size is not None:
        if font_size > 28:  # > 14pt — likely a heading or callout
            return False
        if font_size >= 16:  # >= 8pt — reasonable body text
            return True
    # If text is long enough, it's likely body text (headings are usually short)
    if len(para.text.strip()) > 80:
        return True
    return False


def _find_body_style_ref(doc, near_element=None):
    """Find a body text paragraph in the document to use as style reference.

    If near_element is provided, searches for the best body text paragraph
    near that element first (same section), then falls back to a global search.
    This ensures AI content inherits formatting from the correct section
    (e.g., correct column layout, margins, indentation).
    """
    def _is_good_ref(para):
        """Check if paragraph is a good style reference (long body text with runs)."""
        if not para.text.strip() or not para.runs:
            return False
        if _is_heading_element(para._element):
            return False
        if _has_frame_props(para._element):
            return False
        if len(para.text.strip()) < 40:
            return False
        font_size = _get_font_size_half_points(para._element)
        if font_size is not None and font_size > 28:
            return False
        return True

    # If we have a target element, search nearby paragraphs first
    if near_element is not None:
        body = doc.element.body
        children = list(body)
        try:
            target_idx = children.index(near_element)
        except ValueError:
            target_idx = -1

        if target_idx >= 0:
            # Search backwards then forwards from target for a body text paragraph
            for offset in range(1, 30):
                for idx in [target_idx - offset, target_idx + offset]:
                    if 0 <= idx < len(children):
                        child = children[idx]
                        if child.tag == qn("w:p"):
                            from docx.text.paragraph import Paragraph
                            para = Paragraph(child, body)
                            if _is_good_ref(para):
                                return para

    # Global search: Priority 1 — long body text paragraphs
    for para in doc.paragraphs:
        if _is_good_ref(para):
            return para

    # Priority 2: Any paragraph with text and runs
    for para in doc.paragraphs:
        if para.text.strip() and para.runs and not _is_heading_element(para._element):
            return para

    # Last resort
    for para in doc.paragraphs:
        if para.text.strip() and para.runs:
            return para
    return None


# ------------------------------------------------------------------ #
# Create styled paragraph element (pure XML, no Paragraph class)      #
# ------------------------------------------------------------------ #

def _resolve_spacing_from_style(style_ref):
    """Resolve spacing values from the style reference paragraph.

    Word documents often define spacing in the style definition (styles.xml),
    not on individual paragraphs. python-docx's paragraph_format resolves
    through the full style inheritance chain and returns the effective values.

    Returns a dict of resolved spacing properties (in EMU/twips as needed).
    """
    pf = style_ref.paragraph_format
    result = {}

    # Line spacing — in Word XML: w:spacing/@w:line (twips) + w:lineRule
    if pf.line_spacing is not None:
        from docx.shared import Pt, Twips
        rule = pf.line_spacing_rule
        if rule is not None:
            from docx.enum.text import WD_LINE_SPACING
            if rule == WD_LINE_SPACING.MULTIPLE:
                # Multiple: value is a float (e.g., 1.15), stored as val * 240
                result["line"] = str(int(pf.line_spacing * 240))
                result["lineRule"] = "auto"
            elif rule == WD_LINE_SPACING.EXACTLY:
                # Exact: value is in EMU, convert to twips
                result["line"] = str(int(pf.line_spacing / 12700 * 20))
                result["lineRule"] = "exact"
            elif rule == WD_LINE_SPACING.AT_LEAST:
                result["line"] = str(int(pf.line_spacing / 12700 * 20))
                result["lineRule"] = "atLeast"
            else:
                # Single, 1.5, Double — stored as multiple
                result["line"] = str(int(pf.line_spacing * 240))
                result["lineRule"] = "auto"
        elif isinstance(pf.line_spacing, (int, float)):
            if pf.line_spacing < 5:
                # Likely a multiple (e.g., 1.0, 1.15, 1.5)
                result["line"] = str(int(pf.line_spacing * 240))
                result["lineRule"] = "auto"

    # Space before / after — in EMU, Word XML uses twips (1 twip = 12700 EMU)
    if pf.space_before is not None:
        result["before"] = str(int(pf.space_before / 12700 * 20))
    if pf.space_after is not None:
        result["after"] = str(int(pf.space_after / 12700 * 20))

    return result


def _resolve_indent_from_style(style_ref):
    """Resolve indentation values from the style reference paragraph."""
    pf = style_ref.paragraph_format
    result = {}

    if pf.first_line_indent is not None:
        val = int(pf.first_line_indent / 12700 * 20)
        if val >= 0:
            result["firstLine"] = str(val)
        else:
            result["hanging"] = str(abs(val))
    if pf.left_indent is not None:
        result["left"] = str(int(pf.left_indent / 12700 * 20))
    if pf.right_indent is not None:
        result["right"] = str(int(pf.right_indent / 12700 * 20))

    return result


def _resolve_font_from_style(style_ref):
    """Resolve font name and size from the style reference's run formatting."""
    result = {}
    if not style_ref.runs:
        return result

    font = style_ref.runs[0].font
    if font.name:
        result["name"] = font.name
    if font.size is not None:
        # font.size is in EMU, Word XML uses half-points (1pt = 2 half-points)
        result["sz"] = str(int(font.size / 12700 * 2))

    return result


def _create_styled_p(text, highlight_color=None, style_ref=None):
    """Create a w:p XML element that exactly matches the book's body text style.

    Two-phase approach:
    1. Deep-copy explicit pPr/rPr from the reference paragraph XML
    2. Fill in any missing properties by resolving through the style chain
       (python-docx's paragraph_format/font resolve inherited styles)

    This handles documents where spacing/font are defined in styles.xml
    rather than on individual paragraphs.
    """
    new_p = OxmlElement("w:p")

    # ── Phase 1: Deep-copy explicit paragraph properties ──────────────
    if style_ref is not None:
        source_pPr = style_ref._element.find(qn("w:pPr"))
        if source_pPr is not None:
            pPr = copy.deepcopy(source_pPr)
            # Remove dangerous positioning props
            for tag in [qn("w:framePr"), qn("w:sectPr")]:
                el = pPr.find(tag)
                if el is not None:
                    pPr.remove(el)
            # Remove numbering (don't join AI text to a list)
            numPr = pPr.find(qn("w:numPr"))
            if numPr is not None:
                pPr.remove(numPr)
        else:
            pPr = OxmlElement("w:pPr")
            pStyle_el = OxmlElement("w:pStyle")
            pStyle_el.set(qn("w:val"), "Normal")
            pPr.append(pStyle_el)

        # ── Phase 2: Fill in missing spacing from resolved style chain ──
        # If w:spacing is not explicitly on this paragraph, resolve it
        # from the style definition and add it explicitly.
        existing_spacing = pPr.find(qn("w:spacing"))
        resolved_spacing = _resolve_spacing_from_style(style_ref)

        if existing_spacing is None and resolved_spacing:
            spacing_el = OxmlElement("w:spacing")
            for attr, val in resolved_spacing.items():
                spacing_el.set(qn(f"w:{attr}"), val)
            pPr.append(spacing_el)
        elif existing_spacing is not None and resolved_spacing:
            # Fill in any attributes missing from explicit spacing
            for attr, val in resolved_spacing.items():
                if existing_spacing.get(qn(f"w:{attr}")) is None:
                    existing_spacing.set(qn(f"w:{attr}"), val)

        # Fill in missing indentation from resolved style
        existing_ind = pPr.find(qn("w:ind"))
        resolved_ind = _resolve_indent_from_style(style_ref)

        if existing_ind is None and resolved_ind:
            ind_el = OxmlElement("w:ind")
            for attr, val in resolved_ind.items():
                ind_el.set(qn(f"w:{attr}"), val)
            pPr.append(ind_el)
        elif existing_ind is not None and resolved_ind:
            for attr, val in resolved_ind.items():
                if existing_ind.get(qn(f"w:{attr}")) is None:
                    existing_ind.set(qn(f"w:{attr}"), val)

        # Fill in missing alignment from resolved style
        existing_jc = pPr.find(qn("w:jc"))
        if existing_jc is None and style_ref.paragraph_format.alignment is not None:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            align_map = {
                WD_ALIGN_PARAGRAPH.LEFT: "left",
                WD_ALIGN_PARAGRAPH.CENTER: "center",
                WD_ALIGN_PARAGRAPH.RIGHT: "right",
                WD_ALIGN_PARAGRAPH.JUSTIFY: "both",
                WD_ALIGN_PARAGRAPH.DISTRIBUTE: "distribute",
            }
            align_val = align_map.get(style_ref.paragraph_format.alignment)
            if align_val:
                jc_el = OxmlElement("w:jc")
                jc_el.set(qn("w:val"), align_val)
                pPr.append(jc_el)

    else:
        pPr = OxmlElement("w:pPr")
        pStyle_el = OxmlElement("w:pStyle")
        pStyle_el.set(qn("w:val"), "Normal")
        pPr.append(pStyle_el)

    new_p.append(pPr)

    # ── Run with text ─────────────────────────────────────────────────
    new_r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    if style_ref is not None and style_ref.runs:
        source_rPr = style_ref.runs[0]._element.find(qn("w:rPr"))
        if source_rPr is not None:
            rPr = copy.deepcopy(source_rPr)
            old_hl = rPr.find(qn("w:highlight"))
            if old_hl is not None:
                rPr.remove(old_hl)

        # Fill in missing font properties from resolved style
        resolved_font = _resolve_font_from_style(style_ref)
        if resolved_font:
            # Font family
            if rPr.find(qn("w:rFonts")) is None and "name" in resolved_font:
                rFonts = OxmlElement("w:rFonts")
                rFonts.set(qn("w:ascii"), resolved_font["name"])
                rFonts.set(qn("w:hAnsi"), resolved_font["name"])
                rFonts.set(qn("w:cs"), resolved_font["name"])
                rPr.insert(0, rFonts)
            # Font size
            if rPr.find(qn("w:sz")) is None and "sz" in resolved_font:
                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), resolved_font["sz"])
                rPr.append(sz)
                szCs = OxmlElement("w:szCs")
                szCs.set(qn("w:val"), resolved_font["sz"])
                rPr.append(szCs)

    # Force NOT bold (override whatever was copied)
    for tag_name in ["w:b", "w:bCs"]:
        existing = rPr.find(qn(tag_name))
        if existing is not None:
            existing.set(qn("w:val"), "0")
        else:
            el = OxmlElement(tag_name)
            el.set(qn("w:val"), "0")
            rPr.append(el)

    # Force NOT italic
    for tag_name in ["w:i", "w:iCs"]:
        existing = rPr.find(qn(tag_name))
        if existing is not None:
            existing.set(qn("w:val"), "0")
        else:
            el = OxmlElement(tag_name)
            el.set(qn("w:val"), "0")
            rPr.append(el)

    # Force color to auto (black)
    existing_color = rPr.find(qn("w:color"))
    if existing_color is not None:
        existing_color.set(qn("w:val"), "auto")
    else:
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "auto")
        rPr.append(color)

    # Add highlight
    if highlight_color:
        hl = OxmlElement("w:highlight")
        hl.set(qn("w:val"), highlight_color)
        rPr.append(hl)

    new_r.append(rPr)

    # Add text element
    new_t = OxmlElement("w:t")
    new_t.text = text
    new_t.set(qn("xml:space"), "preserve")
    new_r.append(new_t)

    new_p.append(new_r)
    return new_p


def _create_page_break_p():
    """Create a standalone paragraph containing only a hard page break.

    This is more reliable than w:pageBreakBefore because it works
    even when the insertion point is near floating images/text boxes.
    """
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    return p


def _remove_frame_props(p_element):
    """Remove w:framePr from a paragraph so text flows normally.

    Text frames have fixed dimensions — when replacement text is longer
    than the original, the extra text overflows and gets hidden.
    Removing framePr lets the paragraph flow as regular body text.
    """
    pPr = p_element.find(qn("w:pPr"))
    if pPr is not None:
        framePr = pPr.find(qn("w:framePr"))
        if framePr is not None:
            pPr.remove(framePr)
            return True
    return False


def _has_frame_props(p_element):
    """Check if a paragraph has w:framePr (is inside a text frame)."""
    pPr = p_element.find(qn("w:pPr"))
    if pPr is not None:
        return pPr.find(qn("w:framePr")) is not None
    return False


def _split_and_create_paragraphs(text, highlight_color=None, style_ref=None, page_break_before=False):
    """Split multi-line text into a list of styled w:p elements.

    Args:
        page_break_before: If True, prepend a hard page break paragraph
            to prevent overlap with existing content (images, frames, headers).
    """
    elements = []

    # Add a standalone page break paragraph first
    if page_break_before:
        elements.append(_create_page_break_p())

    lines = text.split("\n") if "\n" in text else [text]
    for line in lines:
        line = line.strip()
        if not line:
            continue
        elements.append(_create_styled_p(line, highlight_color, style_ref))

    return elements


# ------------------------------------------------------------------ #
# Body-level insertion helpers                                         #
# ------------------------------------------------------------------ #

def _append_to_body_end(doc, text, highlight_color=None, style_ref=None):
    """Append paragraph(s) at the very end of the document body.

    Inserts before the final w:sectPr (section properties) if present.
    Adds a page break before the content to prevent overlap.
    """
    body = doc.element.body

    # Find the LAST w:sectPr direct child (the final section properties)
    sect_prs = [child for child in body if child.tag == qn("w:sectPr")]
    last_sect_pr = sect_prs[-1] if sect_prs else None

    paragraphs = _split_and_create_paragraphs(
        text, highlight_color, style_ref, page_break_before=True
    )

    for p in paragraphs:
        if last_sect_pr is not None:
            last_sect_pr.addprevious(p)
        else:
            body.append(p)

    return len(paragraphs)


def _insert_after_body_element(body, target_element, text, highlight_color=None, style_ref=None):
    """Insert paragraph(s) after a specific body-level XML element.

    Adds a page break before the content to prevent overlap with
    existing positioned elements (images, headers, columns).
    """
    paragraphs = _split_and_create_paragraphs(
        text, highlight_color, style_ref, page_break_before=True
    )

    # Insert in reverse order so they end up in correct sequence
    for p in reversed(paragraphs):
        target_element.addnext(p)

    return len(paragraphs)


def _normalize_text(text):
    """Normalize text for fuzzy matching: lowercase, collapse whitespace, normalize quotes."""
    if not text:
        return ""
    # Normalize unicode (smart quotes → regular quotes, etc.)
    text = unicodedata.normalize("NFKD", text)
    # Replace all types of quotes/apostrophes with a standard one
    text = re.sub(r"[\u2018\u2019\u201A\u201B\u0060\u00B4]", "'", text)
    text = re.sub(r'[\u201C\u201D\u201E\u201F]', '"', text)
    # Replace en-dash, em-dash with hyphen
    text = re.sub(r"[\u2013\u2014]", "-", text)
    # Collapse whitespace
    text = re.sub(r"\s+", " ", text).strip().lower()
    return text


def _find_text_in_body(doc, search_text):
    """Find the body-level element containing the given text.

    Searches both body paragraphs AND paragraphs inside tables.
    Uses normalized fuzzy matching to handle smart quotes, special chars, etc.
    Returns the body-level element (w:p or w:tbl) that contains the match.
    """
    if not search_text:
        return None

    normalized_search = _normalize_text(search_text)
    # Also try with just the first 40 chars for partial matching
    short_search = normalized_search[:40] if len(normalized_search) > 40 else normalized_search
    body = doc.element.body

    # Pass 1: Try exact normalized match
    for child in body:
        tag = child.tag
        if tag == qn("w:p"):
            text = _normalize_text(_get_p_text(child))
            if normalized_search in text:
                return child
        elif tag == qn("w:tbl"):
            for p in child.iter(qn("w:p")):
                text = _normalize_text(_get_p_text(p))
                if normalized_search in text:
                    return child

    # Pass 2: Try partial match (first 40 chars) — handles truncated titles
    if len(normalized_search) > 20:
        for child in body:
            tag = child.tag
            if tag == qn("w:p"):
                text = _normalize_text(_get_p_text(child))
                if short_search in text:
                    return child
            elif tag == qn("w:tbl"):
                for p in child.iter(qn("w:p")):
                    text = _normalize_text(_get_p_text(p))
                    if short_search in text:
                        return child

    logger.debug("_find_text_in_body: no match for '%s'", search_text[:60])
    return None


def _is_section_boundary(p_element):
    """Check if a paragraph is a section/chapter boundary (heading or large-font title)."""
    if _is_heading_element(p_element):
        return True
    # Also detect by font size — anything >= 18pt (36 half-points) is likely a heading
    font_size = _get_font_size_half_points(p_element)
    if font_size is not None and font_size >= 36:
        text = _get_p_text(p_element).strip()
        # Large font + short text = likely a heading
        if text and len(text) < 100:
            return True
    return False


def _find_section_end_in_body(doc, start_element):
    """Find the last body-level content element in the current section.

    Starting from start_element, walks forward through body children
    until hitting the next heading/section boundary or end of document.
    Returns the element after which AI content should be inserted.
    Prefers non-framed paragraphs to avoid inserting next to text boxes.
    """
    body = doc.element.body
    children = list(body)

    try:
        start_idx = children.index(start_element)
    except ValueError:
        return start_element

    last_content = start_element
    last_non_framed = start_element  # Prefer non-framed insertion point

    for i in range(start_idx + 1, len(children)):
        child = children[i]
        tag = child.tag

        if tag == qn("w:sectPr"):
            break

        if tag == qn("w:p"):
            if _is_section_boundary(child):
                break
            # Only update last_content if this paragraph has text
            if _get_p_text(child).strip():
                last_content = child
                if not _has_frame_props(child):
                    last_non_framed = child
        elif tag == qn("w:tbl"):
            last_content = child
            last_non_framed = child

    # Prefer inserting after a non-framed element to avoid layout conflicts
    return last_non_framed if last_non_framed != start_element else last_content


# ------------------------------------------------------------------ #
# ExportService                                                        #
# ------------------------------------------------------------------ #

class ExportService:
    """Service for generating updated DOCX files with approved changes applied."""

    @staticmethod
    async def generate_updated_docx(
        document_data: dict,
        approved_changes: list,
        db=None,
        highlighted: bool = False,
    ) -> Optional[str]:
        """
        Generate an updated DOCX file by applying approved changes.

        Args:
            document_data: The document record from MongoDB.
            approved_changes: List of approved ChangeProposal dicts.
            db: Database connection (unused for now, reserved for future use).
            highlighted: If True, highlight changed text in yellow/green.

        Returns:
            Path to the generated DOCX file, or None on failure.
        """
        file_path = document_data.get("file_path")
        if not file_path or not os.path.exists(file_path):
            logger.error("Original file not found: %s", file_path)
            return None

        try:
            doc = DocxDocument(file_path)

            # Separate AI prompt changes (insertions) from regular changes (replacements)
            ai_prompt_changes = [
                c for c in approved_changes
                if c.get("change_type") == "ai_prompt"
                   and not c.get("old_content", "").startswith("[AI Prompt:")
                   or c.get("change_type") == "ai_prompt"
            ]
            regular_changes = [
                c for c in approved_changes
                if c.get("change_type") != "ai_prompt"
            ]

            # Sort regular changes by paragraph index (descending) so that
            # replacements don't shift indices for subsequent changes.
            sorted_changes = sorted(
                regular_changes,
                key=lambda c: c.get("paragraph_idx", 0),
                reverse=True,
            )

            replacements_made = 0
            insertions_made = 0

            # 1. Apply regular text replacements
            for change in sorted_changes:
                old_content = change.get("old_content", "")
                new_content = change.get("user_edited_content") or change.get("new_content", "")

                if not old_content or not new_content:
                    continue

                for paragraph in doc.paragraphs:
                    if old_content in paragraph.text:
                        if highlighted:
                            # Split paragraph text into: before | replaced | after
                            full_text = paragraph.text
                            idx = full_text.find(old_content)
                            before_text = full_text[:idx]
                            after_text = full_text[idx + len(old_content):]

                            # If paragraph is in a text frame and new text is longer,
                            # remove the frame so text doesn't get cut off
                            p_elem = paragraph._element
                            if len(new_content) > len(old_content):
                                _remove_frame_props(p_elem)

                            # Get formatting from first run to preserve font/size
                            ref_rPr = None
                            if paragraph.runs:
                                ref_rPr = paragraph.runs[0]._element.find(qn("w:rPr"))

                            # Remove ALL existing runs from the paragraph XML
                            for r_elem in list(p_elem.findall(qn("w:r"))):
                                p_elem.remove(r_elem)

                            def _make_run(text, hl_color=None):
                                """Create a run element with ref formatting and optional highlight."""
                                r = OxmlElement("w:r")
                                if ref_rPr is not None:
                                    new_rPr = copy.deepcopy(ref_rPr)
                                    # Remove any existing highlight from copied props
                                    old_hl = new_rPr.find(qn("w:highlight"))
                                    if old_hl is not None:
                                        new_rPr.remove(old_hl)
                                    if hl_color:
                                        hl = OxmlElement("w:highlight")
                                        hl.set(qn("w:val"), hl_color)
                                        new_rPr.append(hl)
                                    r.append(new_rPr)
                                elif hl_color:
                                    rPr = OxmlElement("w:rPr")
                                    hl = OxmlElement("w:highlight")
                                    hl.set(qn("w:val"), hl_color)
                                    rPr.append(hl)
                                    r.append(rPr)
                                t = OxmlElement("w:t")
                                t.text = text
                                t.set(qn("xml:space"), "preserve")
                                r.append(t)
                                return r

                            # Add runs: before (no highlight) + replaced (yellow) + after (no highlight)
                            if before_text:
                                p_elem.append(_make_run(before_text))
                            p_elem.append(_make_run(new_content, "yellow"))
                            if after_text:
                                p_elem.append(_make_run(after_text))

                            replacements_made += 1
                        else:
                            # Plain replacement (no highlight)
                            for run in paragraph.runs:
                                if old_content in run.text:
                                    run.text = run.text.replace(old_content, new_content)
                                    replacements_made += 1
                                    break
                            else:
                                full_text = paragraph.text
                                if old_content in full_text:
                                    new_full_text = full_text.replace(old_content, new_content)
                                    if paragraph.runs:
                                        paragraph.runs[0].text = new_full_text
                                        for run in paragraph.runs[1:]:
                                            run.text = ""
                                        replacements_made += 1
                        break

            # 2. Apply AI prompt insertions (sorted descending by paragraph_idx)
            sorted_ai_changes = sorted(
                ai_prompt_changes,
                key=lambda c: c.get("paragraph_idx", 0),
                reverse=True,
            )

            # Global fallback style reference (used when no nearby ref found)
            global_style_ref = _find_body_style_ref(doc)

            for change in sorted_ai_changes:
                new_content = change.get("user_edited_content") or change.get("new_content", "")
                if not new_content:
                    continue

                metadata = change.get("ai_prompt_metadata", {}) or {}
                placement = metadata.get("placement", "at_end")
                section_title = metadata.get("section_title", "")
                old_content = change.get("old_content", "")
                is_replace = (
                    placement == "replace_section"
                    and old_content
                    and not old_content.startswith("[AI Prompt:")
                )

                hl_color = "green" if highlighted else None

                if placement == "at_end" and not is_replace:
                    # ---- AT END: append to the very end of document body ----
                    style_ref = global_style_ref
                    count = _append_to_body_end(doc, new_content, hl_color, style_ref)
                    insertions_made += count
                    logger.info("Inserted %d AI paragraphs at document end", count)

                elif is_replace:
                    # ---- REPLACE SECTION: find old content and replace it ----
                    first_line = old_content.split('\n')[0].strip()
                    target_el = _find_text_in_body(doc, first_line) if first_line else None

                    if target_el is not None and target_el.tag == qn("w:p"):
                        # Find style ref near the target for correct section formatting
                        style_ref = _find_body_style_ref(doc, near_element=target_el) or global_style_ref
                        # Replace the paragraph's text
                        content_lines = [l.strip() for l in new_content.split("\n") if l.strip()]
                        if content_lines:
                            from docx.text.paragraph import Paragraph
                            target_para = Paragraph(target_el, doc.element.body)
                            for run in target_para.runs:
                                run.text = ""
                            if target_para.runs:
                                target_para.runs[0].text = content_lines[0]
                                if highlighted:
                                    _add_highlight(target_para.runs[0], "green")
                            # Insert remaining lines after
                            if len(content_lines) > 1:
                                remaining_text = "\n".join(content_lines[1:])
                                _insert_after_body_element(
                                    doc.element.body, target_el,
                                    remaining_text, hl_color, style_ref,
                                )
                        insertions_made += 1
                    else:
                        # Fallback: append at end
                        style_ref = global_style_ref
                        count = _append_to_body_end(doc, new_content, hl_color, style_ref)
                        insertions_made += count

                elif placement == "after_section":
                    # ---- AFTER SECTION: find section heading, insert at end of that section ----
                    insert_target = None
                    para_idx = change.get("paragraph_idx", 0)

                    # Strategy 1: Find by section title text (most reliable)
                    if section_title:
                        heading_el = _find_text_in_body(doc, section_title)
                        if heading_el is not None:
                            insert_target = _find_section_end_in_body(doc, heading_el)
                            logger.debug("Found section '%s' by title match", section_title[:40])

                    # Strategy 2: Fallback — try to find by old_content hint
                    if insert_target is None and old_content and not old_content.startswith("[AI Prompt:"):
                        first_line = old_content.split('\n')[0].strip()
                        if first_line:
                            found = _find_text_in_body(doc, first_line)
                            if found:
                                insert_target = _find_section_end_in_body(doc, found)
                                logger.debug("Found section by old_content match")

                    # Strategy 3: Fallback — use paragraph_idx with doc.paragraphs
                    # Even though idx may not be perfect, it gets content in the
                    # right general area instead of losing it at the end
                    if insert_target is None and para_idx < len(doc.paragraphs):
                        target_para = doc.paragraphs[min(para_idx, len(doc.paragraphs) - 1)]
                        insert_target = target_para._element
                        logger.debug(
                            "Using paragraph_idx %d fallback for section '%s'",
                            para_idx, section_title or "(unknown)",
                        )

                    if insert_target is not None:
                        # Find style ref near insertion point for correct formatting
                        style_ref = _find_body_style_ref(doc, near_element=insert_target) or global_style_ref
                        count = _insert_after_body_element(
                            doc.element.body, insert_target,
                            new_content, hl_color, style_ref,
                        )
                        insertions_made += count
                        logger.info(
                            "Inserted %d paragraphs after section '%s'",
                            count, section_title or "(fallback)",
                        )
                    else:
                        # Final fallback: append at end
                        logger.warning(
                            "Could not find section '%s' in DOCX — appending at end",
                            section_title,
                        )
                        style_ref = global_style_ref
                        count = _append_to_body_end(doc, new_content, hl_color, style_ref)
                        insertions_made += count

                else:
                    # Unknown placement — append at end as safe default
                    style_ref = global_style_ref
                    count = _append_to_body_end(doc, new_content, hl_color, style_ref)
                    insertions_made += count

            # Save to output directory
            os.makedirs(settings.OUTPUT_DIR, exist_ok=True)
            original_name = document_data.get("original_filename", "document.docx")
            base_name = os.path.splitext(original_name)[0]

            if highlighted:
                output_filename = f"{base_name}_highlighted.docx"
            else:
                output_filename = f"{base_name}_updated.docx"

            output_path = os.path.join(settings.OUTPUT_DIR, output_filename)

            # Validate references after applying changes
            text_content = document_data.get("text_content", "")
            figures = document_data.get("figures", [])
            tables = document_data.get("tables", [])
            equations = document_data.get("equations", [])

            if text_content:
                warnings = RenumberingService.validate_references(
                    text_content, figures, tables, equations
                )
                for warning in warnings:
                    logger.warning("Export reference check: %s", warning)

            doc.save(output_path)
            logger.info(
                "Generated %s DOCX: %s (%d replacements, %d insertions)",
                "highlighted" if highlighted else "clean",
                output_path,
                replacements_made,
                insertions_made,
            )
            return output_path

        except Exception as e:
            logger.error("Failed to generate updated DOCX: %s", str(e))
            return None

    @staticmethod
    async def generate_preview_pdf(
        document_data: dict,
        approved_changes: list,
    ) -> Optional[str]:
        """
        Generate a PDF preview that is pixel-perfect identical to the Highlighted DOCX.

        Steps:
            1. Generate the highlighted DOCX (reuses generate_updated_docx with highlighted=True)
            2. Convert that DOCX → PDF using docx2pdf (uses Microsoft Word on Windows)

        Returns path to the generated PDF file, or None on failure.
        """
        try:
            # 1. Generate the highlighted DOCX (same file the user would download)
            highlighted_docx_path = await ExportService.generate_updated_docx(
                document_data, approved_changes, highlighted=True
            )
            if not highlighted_docx_path or not os.path.exists(highlighted_docx_path):
                logger.error("Failed to generate highlighted DOCX for preview")
                return None

            # 2. Convert DOCX → PDF
            pdf_path = os.path.splitext(highlighted_docx_path)[0] + "_preview.pdf"

            try:
                from docx2pdf import convert
                convert(highlighted_docx_path, pdf_path)
            except Exception as e:
                logger.warning("docx2pdf failed (%s), trying LibreOffice fallback", e)
                # Fallback: try LibreOffice headless
                import subprocess
                result = subprocess.run(
                    [
                        "soffice", "--headless", "--convert-to", "pdf",
                        "--outdir", os.path.dirname(pdf_path),
                        highlighted_docx_path,
                    ],
                    capture_output=True, text=True, timeout=60,
                )
                # LibreOffice outputs to same dir with .pdf extension
                lo_pdf = os.path.splitext(highlighted_docx_path)[0] + ".pdf"
                if os.path.exists(lo_pdf) and lo_pdf != pdf_path:
                    os.rename(lo_pdf, pdf_path)
                if not os.path.exists(pdf_path):
                    logger.error("LibreOffice conversion also failed: %s", result.stderr)
                    return None

            if not os.path.exists(pdf_path):
                logger.error("PDF conversion produced no output")
                return None

            logger.info("Generated preview PDF: %s", pdf_path)
            return pdf_path

        except Exception as e:
            logger.error("Failed to generate preview PDF: %s", str(e))
            return None
