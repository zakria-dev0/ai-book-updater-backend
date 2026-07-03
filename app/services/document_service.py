# document_service.py
from app.utils.omml_to_latex import omml_to_latex
from docx import Document as DocxDocument
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph
from typing import List, Optional, Tuple
from app.models.document import Equation, Figure, Table, Position, DocumentMetadata
from app.services.image_service import ImageService
import base64
from io import BytesIO
from PIL import Image
import re
import math
import zipfile
import fitz  # PyMuPDF
import tempfile
import os
from docx2pdf import convert

class DOCXParser:
    """Parser for DOCX documents with optional Nougat support via PDF conversion"""

    def __init__(self, file_path: str, use_nougat: bool = True):
        self.file_path = file_path
        self.use_nougat = use_nougat
        self.doc = DocxDocument(file_path)
        self.equations: List[Equation] = []
        self.figures: List[Figure] = []
        self.tables: List[Table] = []
        self.text_content = ""
        self._real_page_count = self._read_page_count_from_properties()
        self._para_to_page = self._build_page_map()

    def _read_page_count_from_properties(self) -> Optional[int]:
        """
        Read real page count from DOCX extended properties (docProps/app.xml).
        Word stores the rendered page count here when the document is saved.
        """
        try:
            with zipfile.ZipFile(self.file_path, 'r') as z:
                if 'docProps/app.xml' in z.namelist():
                    from lxml import etree
                    app_xml = z.read('docProps/app.xml')
                    root = etree.fromstring(app_xml)
                    ns = {'ep': 'http://schemas.openxmlformats.org/officeDocument/2006/extended-properties'}
                    pages_el = root.find('ep:Pages', ns)
                    if pages_el is not None and pages_el.text:
                        count = int(pages_el.text)
                        if count > 0:
                            return count
        except Exception:
            pass
        return None

    def _build_page_map(self) -> dict:
        """
        Build a mapping of paragraph index → page number by scanning for
        page break elements (<w:br w:type="page"/> and <w:lastRenderedPageBreak/>)
        in each paragraph's XML.
        """
        W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        para_to_page: dict[int, int] = {}
        current_page = 1

        for para_idx, para in enumerate(self.doc.paragraphs):
            para_to_page[para_idx] = current_page
            # Check for explicit page breaks: <w:br w:type="page"/>
            for br in para._element.findall(f'.//{{{W_NS}}}br'):
                if br.get(f'{{{W_NS}}}type') == 'page':
                    current_page += 1
                    para_to_page[para_idx] = current_page
            # Check for rendered page breaks: <w:lastRenderedPageBreak/>
            if para._element.findall(f'.//{{{W_NS}}}lastRenderedPageBreak'):
                current_page += 1
                para_to_page[para_idx] = current_page

        # If the XML-based map found far fewer pages than the real count
        # (common when DOCX lacks <w:lastRenderedPageBreak/> markers),
        # fall back to proportional mapping using the real page count.
        total_paras = len(self.doc.paragraphs)
        if total_paras > 0 and self._real_page_count:
            xml_pages = len(set(para_to_page.values()))
            if self._real_page_count > xml_pages * 2:
                for para_idx in range(total_paras):
                    page = math.ceil((para_idx + 1) / total_paras * self._real_page_count)
                    para_to_page[para_idx] = max(1, page)

        return para_to_page

    def _page_for_para(self, para_idx: Optional[int]) -> Optional[int]:
        """Get the page number for a given paragraph index."""
        if para_idx is None:
            return None
        return self._para_to_page.get(para_idx)

    def parse(self) -> Tuple[str, List[Equation], List[Figure], List[Table], DocumentMetadata]:
        """Parse DOCX document and extract all content"""

        if self.use_nougat:
            try:
                # Convert DOCX to PDF and use Nougat for better equation extraction
                print("Using Nougat mode: Converting DOCX to PDF...")
                return self._parse_with_nougat()
            except Exception as e:
                print(f"Nougat conversion failed, using standard DOCX parsing: {e}")
                return self._parse_standard()
        else:
            return self._parse_standard()

    def _parse_with_nougat(self) -> Tuple[str, List[Equation], List[Figure], List[Table], DocumentMetadata]:
        """Convert DOCX to PDF and use Nougat for extraction"""
        # Create temporary PDF file
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_pdf:
            pdf_path = tmp_pdf.name

        try:
            # Convert DOCX to PDF
            print(f"Converting {self.file_path} to PDF...")
            convert(self.file_path, pdf_path)
            print(f"Conversion complete: {pdf_path}")

            # Use PDFParser with Nougat
            pdf_parser = PDFParser(pdf_path, use_nougat=True)
            text, equations, figures, tables, metadata = pdf_parser.parse()

            # Store results
            self.text_content = text
            self.equations = equations
            self.figures = figures
            self.tables = tables

            return text, equations, figures, tables, metadata

        finally:
            # Clean up temporary PDF
            if os.path.exists(pdf_path):
                try:
                    os.unlink(pdf_path)
                except:
                    pass

    def _parse_standard(self) -> Tuple[str, List[Equation], List[Figure], List[Table], DocumentMetadata]:
        """Standard DOCX parsing without Nougat"""
        # Extract text content
        self.text_content = self._extract_text()

        # Extract equations
        self.equations = self._extract_equations()

        # Extract figures (images)
        self.figures = self._extract_figures()

        # Extract tables
        self.tables = self._extract_tables()

        # Generate metadata
        metadata = self._generate_metadata()

        return self.text_content, self.equations, self.figures, self.tables, metadata
    
    def _extract_text(self) -> str:
        """Extract all text from document"""
        full_text = []
        for para in self.doc.paragraphs:
            full_text.append(para.text)
        return "\n".join(full_text)

    def _extract_paragraphs(self) -> list:
        """Extract every paragraph with rich metadata for downstream use.

        Returns a list of dicts, one per paragraph, preserving the original
        paragraph index so that outline / section-paragraph endpoints can
        work entirely from the DB without re-reading the DOCX.
        """
        import re as _re
        W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
        A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
        R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

        paragraphs = []
        for idx, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            style_name = para.style.name if para.style else ""

            # Determine formatting from runs
            runs = [r for r in para.runs if r.text.strip()]
            is_bold = all(r.bold for r in runs) if runs else False
            font_sizes = set()
            for r in runs:
                if r.font and r.font.size:
                    font_sizes.add(r.font.size.pt)

            # Check if paragraph has a drawing/image
            has_image = bool(para._element.findall(f'.//{{{W_NS}}}drawing'))

            # Collect r:embed IDs for images in this paragraph
            r_embeds = []
            for blip in para._element.findall(f'.//{{{A_NS}}}blip'):
                r_embed = blip.get(f'{{{R_NS}}}embed')
                if r_embed:
                    r_embeds.append(r_embed)

            # Get image dimensions from drawing extents
            image_cx = 0
            image_cy = 0
            extent = para._element.find(f'.//{{{WP_NS}}}extent')
            if extent is not None:
                image_cx = int(extent.get('cx', '0'))
                image_cy = int(extent.get('cy', '0'))

            # Check for OMML equations
            MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
            has_equation = bool(para._element.findall(f'.//{{{MATH_NS}}}oMath'))

            paragraphs.append({
                "idx": idx,
                "text": text,
                "style": style_name,
                "bold": is_bold,
                "font_sizes": sorted(font_sizes) if font_sizes else [],
                "length": len(text),
                "has_image": has_image,
                "r_embeds": r_embeds,
                "image_cx": image_cx,
                "image_cy": image_cy,
                "has_equation": has_equation,
                "page": self._para_to_page.get(idx, 1),
            })

        return paragraphs
    
    def _extract_equations(self) -> List[Equation]:
        """
        Extract equations from DOCX by parsing OMML (Office Math Markup Language) XML.
        Finds all <m:oMath> elements in both body paragraphs AND table cells.
        raw_omml is stored so it can later be converted to LaTeX via Mathpix or omml2latex.
        """
        from lxml import etree

        equations = []
        seen_omml = set()  # deduplicate by OMML XML content
        seen_latex = set()  # also deduplicate by final LaTeX text
        MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
        equation_number_pattern = r'\((\d+(?:[.\-]\d+)*|[A-Z](?:[.\-]\d+)+|[ivxlc]+)\)'

        def _process_paragraph(para, para_idx: int, source: str = "body"):
            """Extract equations from a single paragraph element."""
            omath_elements = para._element.findall(f'.//{{{MATH_NS}}}oMath')

            for eq_idx, omath in enumerate(omath_elements):
                omml_xml = etree.tostring(omath, encoding='unicode')

                # Deduplicate: skip if we've already seen this exact OMML
                omml_hash = hash(omml_xml)
                if omml_hash in seen_omml:
                    continue
                seen_omml.add(omml_hash)

                # Plain-text from <m:t> leaf nodes
                t_elements = omath.findall(f'.//{{{MATH_NS}}}t')
                eq_text = ''.join(t.text or '' for t in t_elements).strip()

                # Filter junk (symbol palettes, UI tables, alphabet listings)
                if self._is_junk_equation(eq_text, omml_xml):
                    continue

                # Convert OMML → LaTeX
                latex = omml_to_latex(omml_xml) or eq_text

                if self._is_junk_latex(latex):
                    continue

                # Deduplicate by normalized LaTeX text (catches near-identical table cell copies)
                latex_key = latex.strip().replace(' ', '')
                if latex_key in seen_latex:
                    continue
                seen_latex.add(latex_key)

                # Look for equation number in surrounding text
                para_text = para.text or ""
                eq_number_match = re.search(equation_number_pattern, para_text)
                eq_number = eq_number_match.group(0) if eq_number_match else None

                eq_id = f"eq_{source}_{para_idx}_{eq_idx}"
                equations.append(Equation(
                    equation_id=eq_id,
                    latex=latex,
                    raw_omml=omml_xml,
                    position=Position(page=self._page_for_para(para_idx), paragraph=para_idx),
                    number=eq_number
                ))

        # ── 1. Body paragraphs ──
        for para_idx, para in enumerate(self.doc.paragraphs):
            _process_paragraph(para, para_idx, "body")

        # ── 2. Table cells — many documents embed equations inside tables ──
        for t_idx, table in enumerate(self.doc.tables):
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _process_paragraph(para, t_idx + 10000, f"tbl{t_idx}")

        return equations

    @staticmethod
    def _is_junk_equation(eq_text: str, omml_xml: str) -> bool:
        """Return True if the OMML content is not a real math equation."""
        text_lower = eq_text.lower().strip()

        # 1. Empty
        if len(eq_text.strip()) == 0:
            return True

        # 2. Single character (just a lone symbol/letter, not an equation)
        if len(eq_text.strip()) == 1:
            return True

        # 3. Placeholder text from Word equation templates
        if text_lower in ('type equation here', 'type equation here.'):
            return True

        # 4. Alphabet listing / symbol palette (A,B,C,...Z or a,b,c,...z)
        stripped = eq_text.replace(',', '').replace(' ', '')
        if len(stripped) >= 20:
            unique_chars = set(stripped)
            alpha_chars = {c for c in unique_chars if c.isalpha()}
            if len(alpha_chars) >= 20 and all(c.isalpha() or c in ',. ' for c in eq_text):
                return True

        # 5. Contains Word UI / toolbar keywords
        ui_keywords = [
            'file', 'home', 'insert', 'design', 'layout', 'mailings',
            'review', 'view', 'add-ins', 'cover page', 'page break',
            'blank page', 'gallery', 'category', 'description', 'save in',
            'options', 'autotext', 'building block', 'search document',
            'normal', 'insert content only', 'online pictures',
        ]
        if any(kw in text_lower for kw in ui_keywords):
            return True

        # 6. Too many \\hline or tabular markers (formatting table, not equation)
        if text_lower.count('\\hline') >= 2 or '\\begin{tabular}' in text_lower:
            return True

        # 7. Mostly comma-separated single characters (symbol reference)
        parts = [p.strip() for p in eq_text.split(',')]
        if len(parts) >= 15 and all(len(p) <= 2 for p in parts):
            return True

        return False

    @staticmethod
    def _is_junk_latex(latex: str) -> bool:
        """Return True if the converted LaTeX is not a real equation."""
        stripped = latex.strip()

        # Single LaTeX command with no arguments (e.g., just \\partial)
        if re.match(r'^\\[a-zA-Z]+$', stripped) and len(stripped) < 20:
            return True

        # Just a single letter or digit
        if len(stripped) <= 1:
            return True

        # LaTeX tabular environment (Word UI table, not equation)
        if '\\begin{tabular}' in stripped:
            return True

        # Alphabet runs
        clean = stripped.replace(',', '').replace(' ', '')
        if len(clean) >= 20:
            alpha_only = ''.join(c for c in clean if c.isalpha())
            if len(alpha_only) >= 20:
                unique = set(alpha_only.lower())
                if len(unique) >= 18:
                    return True

        return False

    @staticmethod
    def _is_trivial_equation(latex: str) -> bool:
        """Return True if the LaTeX is a trivial inline math fragment, not a real equation.

        Catches standalone variables (Δv, μ, V∞), single variables with subscripts
        (V_{Earth}), and other short fragments that aren't full formulas.
        A real equation should contain a relational operator (=, <, >, ≤, ≥, ≈)
        or be long enough to be a meaningful expression.
        """
        import re as _re

        stripped = latex.strip()
        if not stripped:
            return True

        # Remove \text{} wrappers for length/content checking
        clean = _re.sub(r'\\(?:text|mathrm|mathbf)\{([^}]*)\}', r'\1', stripped)
        clean = clean.replace(' ', '')

        # Very short expressions without relational operators are trivial
        # (e.g., "\\Delta v", "V_{\\infty}", "\\mu", "R_{SOI}")
        has_relation = bool(_re.search(r'[=<>≤≥≈≠]|\\(?:eq|leq|geq|neq|approx|sim|equiv|le|ge)', stripped))

        if not has_relation:
            # No equals/comparison sign — only keep if it's long enough to be
            # a meaningful standalone expression (like a summation or integral)
            if len(clean) < 30:
                return True

        return False

    def _extract_figures(self) -> List[Figure]:
        """Extract images/figures from DOCX with full metadata for downstream analysis."""
        figures = []

        A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
        R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
        W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        # Build r_embed → (para_idx, cx, cy) map from drawings
        seen_embeds = set()
        embed_info: dict = {}  # r_embed -> {para_idx, cx, cy}
        for para_idx, para in enumerate(self.doc.paragraphs):
            drawings = para._element.findall(f'.//{{{W_NS}}}drawing')
            for drawing in drawings:
                blip = drawing.find(f'.//{{{A_NS}}}blip')
                if blip is not None:
                    r_embed = blip.get(f'{{{R_NS}}}embed')
                    if r_embed and r_embed not in seen_embeds:
                        seen_embeds.add(r_embed)
                        cx, cy = 0, 0
                        extent = drawing.find(f'.//{{{WP_NS}}}extent')
                        if extent is not None:
                            cx = int(extent.get('cx', '0'))
                            cy = int(extent.get('cy', '0'))
                        embed_info[r_embed] = {"para_idx": para_idx, "cx": cx, "cy": cy}

        fig_counter = 0
        for rel in self.doc.part.rels.values():
            if "image" not in rel.target_ref:
                continue
            try:
                image_data = rel.target_part.blob

                # Convert non-standard formats (EMF, WMF, TIFF, BMP) to PNG
                # so GPT Vision and browsers can display them
                try:
                    img = Image.open(BytesIO(image_data))
                    fmt = (img.format or "").upper()
                    if fmt not in ("PNG", "JPEG", "GIF", "WEBP"):
                        # Convert to PNG
                        if img.mode not in ("RGB", "RGBA", "L"):
                            img = img.convert("RGB")
                        buf = BytesIO()
                        img.save(buf, format="PNG")
                        image_data = buf.getvalue()
                except Exception:
                    pass  # If Pillow can't open it, keep original bytes

                # Downscale/re-encode before storing so large documents with
                # many full-resolution scans don't blow past the DB storage
                # quota (base64 adds ~33% on top of the raw bytes already).
                image_data = ImageService.compress_for_storage(image_data)

                image_base64 = base64.b64encode(image_data).decode('utf-8')
                info = embed_info.get(rel.rId, {})
                para_pos = info.get("para_idx")
                cx = info.get("cx", 0)
                cy = info.get("cy", 0)

                # Find caption from next paragraph
                caption = None
                if para_pos is not None and para_pos + 1 < len(self.doc.paragraphs):
                    next_text = self.doc.paragraphs[para_pos + 1].text.strip()
                    if next_text.lower().startswith("figure") or next_text.lower().startswith("fig"):
                        caption = next_text
                if not caption:
                    caption = self._find_figure_caption(fig_counter)

                figure = Figure(
                    figure_id=f"fig_{fig_counter}",
                    caption=caption,
                    image_base64=image_base64,
                    position=Position(page=self._page_for_para(para_pos), paragraph=para_pos),
                    number=self._extract_figure_number(caption) if caption else None,
                    r_embed=rel.rId,
                    size_bytes=len(image_data),
                    cx=cx,
                    cy=cy,
                )
                figures.append(figure)
                fig_counter += 1
            except Exception as e:
                print(f"Error extracting figure {fig_counter}: {e}")
                fig_counter += 1
                continue

        return figures
    
    def _extract_tables(self) -> List[Table]:
        """Extract tables from DOCX with equation-aware cell text extraction."""
        from lxml import etree

        tables = []
        MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

        # Walk ALL XML elements iteratively (avoids Python recursion limit on large docs).
        # Collect tables in document order alongside their para count.
        tables_in_order = []   # list of (lxml_tbl_element, para_count) in document order
        para_count = 0

        stack = list(reversed(list(self.doc.element.body)))
        while stack:
            element = stack.pop()
            local_tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
            if local_tag == 'p':
                para_count += 1
            elif local_tag == 'tbl':
                tables_in_order.append((element, para_count))
            stack.extend(reversed(list(element)))

        def _cell_text_with_equations(cell) -> str:
            """Extract cell text, converting embedded OMML equations to LaTeX."""
            parts = []
            for para in cell.paragraphs:
                # Check for OMML equations in the paragraph
                omath_elements = para._element.findall(f'.//{{{MATH_NS}}}oMath')
                if omath_elements:
                    for omath in omath_elements:
                        omml_xml = etree.tostring(omath, encoding='unicode')
                        latex = omml_to_latex(omml_xml)
                        if latex and latex.strip():
                            parts.append(latex.strip())
                        else:
                            # Fallback: plain text from <m:t> nodes
                            t_els = omath.findall(f'.//{{{MATH_NS}}}t')
                            eq_text = ''.join(t.text or '' for t in t_els).strip()
                            if eq_text:
                                parts.append(eq_text)
                else:
                    # No equations — use plain text
                    text = para.text or ""
                    if text.strip():
                        parts.append(text.strip())
            return ' '.join(parts) if parts else ""

        # ── Collect ALL "Table X-Y" caption paragraphs in document order ──
        # Use the body-level children to find caption paragraphs by their position
        # relative to table elements in the same walk.
        table_caption_pat = re.compile(r'(?:TABLE|Table)\s+\d+[\-\.]\d+', re.IGNORECASE)
        WPC_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        # Walk body children in order — track caption paragraphs and table elements
        caption_queue = []     # captions seen so far, waiting to be assigned
        table_captions = {}    # table_element_id -> caption text
        last_caption = None    # most recent caption paragraph text

        for child in self.doc.element.body:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'p':
                # Extract text from paragraph
                texts = []
                for t_el in child.iter(f'{{{WPC_NS}}}t'):
                    if t_el.text:
                        texts.append(t_el.text)
                para_text = ''.join(texts).strip()
                if para_text and table_caption_pat.search(para_text):
                    last_caption = para_text
            elif tag == 'tbl':
                # Assign the most recent caption to this table
                if last_caption:
                    table_captions[id(child)] = last_caption
                    last_caption = None  # consume — don't reuse for next table

        for idx, table in enumerate(self.doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [_cell_text_with_equations(cell) for cell in row.cells]
                table_data.append(row_data)

            para_pos = tables_in_order[idx][1] if idx < len(tables_in_order) else None

            # Look up caption by the table's underlying XML element
            caption = table_captions.get(id(table._tbl))

            table_obj = Table(
                table_id=f"tbl_{idx}",
                caption=caption,
                content=table_data,
                position=Position(page=self._page_for_para(para_pos), paragraph=para_pos),
                number=self._extract_table_number(caption) if caption else None
            )
            tables.append(table_obj)

        return tables
    
    def _find_figure_caption(self, figure_idx: int) -> Optional[str]:
        """Find caption for a figure"""
        # Look through paragraphs for "Figure X-Y" pattern
        figure_pattern = r'Figure\s+\d+-\d+[:.]\s*(.+)'
        
        for para in self.doc.paragraphs:
            match = re.search(figure_pattern, para.text, re.IGNORECASE)
            if match:
                return para.text.strip()
        
        return None
    
    def _find_table_caption(self, table_idx: int) -> Optional[str]:
        """Find caption for a table"""
        table_pattern = r'Table\s+\d+-\d+[:.]\s*(.+)'
        
        for para in self.doc.paragraphs:
            match = re.search(table_pattern, para.text, re.IGNORECASE)
            if match:
                return para.text.strip()
        
        return None
    
    def _extract_figure_number(self, caption: Optional[str]) -> Optional[str]:
        """Extract figure number from caption"""
        if not caption:
            return None
        match = re.search(r'Figure\s+(\d+-\d+)', caption, re.IGNORECASE)
        return match.group(1) if match else None
    
    def _extract_table_number(self, caption: Optional[str]) -> Optional[str]:
        """Extract table number from caption"""
        if not caption:
            return None
        match = re.search(r'Table\s+(\d+-\d+)', caption, re.IGNORECASE)
        return match.group(1) if match else None
    
    def _get_page_count(self) -> int:
        """
        Get total page count. Prefers the real count from DOCX extended properties;
        falls back to counting explicit page-break XML elements + section breaks.
        If that still seems too low, estimates from word count (~250 words/page).
        """
        if self._real_page_count:
            return self._real_page_count

        W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        # Count explicit page breaks
        page_breaks = sum(
            1 for el in self.doc.element.body.iter(f'{{{W_NS}}}br')
            if el.get(f'{{{W_NS}}}type') == 'page'
        )
        # Count section breaks (which also cause page breaks)
        section_breaks = sum(
            1 for el in self.doc.element.body.iter(f'{{{W_NS}}}sectPr')
        )
        xml_pages = max(1, page_breaks + section_breaks)

        # Estimate from word count as a sanity check (~250 words per page)
        total_words = sum(len(p.text.split()) for p in self.doc.paragraphs if p.text.strip())
        word_estimate = max(1, math.ceil(total_words / 250))

        # Use the higher of the two estimates
        return max(xml_pages, word_estimate)

    def _generate_metadata(self) -> DocumentMetadata:
        """Generate document metadata including title and author from core properties"""
        props = self.doc.core_properties
        return DocumentMetadata(
            title=props.title or None,
            author=props.author or None,
            total_pages=self._get_page_count(),
            total_paragraphs=len(self.doc.paragraphs),
            total_equations=len(self.equations),
            total_figures=len(self.figures),
            total_tables=len(self.tables),
        )


class PDFParser:
    """Parser for PDF documents using Nougat OCR for scientific documents"""

    def __init__(self, file_path: str, use_nougat: bool = True):
        self.file_path = file_path
        self.use_nougat = use_nougat
        self.doc = fitz.open(file_path)
        self.equations: List[Equation] = []
        self.figures: List[Figure] = []
        self.tables: List[Table] = []
        self.text_content = ""
        self._nougat_model = None

    def parse(self) -> Tuple[str, List[Equation], List[Figure], List[Table], DocumentMetadata]:
        """Parse PDF document and extract all content"""
        if self.use_nougat:
            try:
                self._parse_with_nougat()
            except Exception as e:
                print(f"Nougat parsing failed, falling back to PyMuPDF: {e}")
                self._parse_with_pymupdf()
        else:
            self._parse_with_pymupdf()

        metadata = self._generate_metadata()
        return self.text_content, self.equations, self.figures, self.tables, metadata

    def _parse_with_nougat(self):
        """Use Nougat for advanced PDF parsing (equations, tables, etc.)"""
        try:
            from nougat import NougatModel
            from nougat.postprocessing import markdown_compatible
            from nougat.utils.checkpoint import get_checkpoint

            # Load model (lazy loading - only once)
            if self._nougat_model is None:
                print("Loading Nougat model...")
                checkpoint = get_checkpoint(None, model_tag="0.1.0-small")
                self._nougat_model = NougatModel.from_pretrained(checkpoint)
                print("Nougat model loaded")

            # Process PDF with Nougat
            print(f"Processing PDF with Nougat: {self.file_path}")
            output = self._nougat_model.predict(self.file_path)

            # Nougat returns markdown with LaTeX equations
            markdown_text = markdown_compatible(output)
            self.text_content = markdown_text

            # Extract equations from markdown (LaTeX is between $ or $$)
            self._extract_equations_from_markdown(markdown_text)

            # Extract figures using PyMuPDF (Nougat doesn't extract images)
            self._extract_figures_pymupdf()

            # Extract tables from markdown
            self._extract_tables_from_markdown(markdown_text)

        except ImportError:
            print("Nougat not installed, falling back to PyMuPDF")
            self._parse_with_pymupdf()
        except Exception as e:
            print(f"Nougat error: {e}")
            raise

    def _parse_with_pymupdf(self):
        """Fallback: Use PyMuPDF for basic PDF parsing"""
        self.text_content = self._extract_text_pymupdf()
        self.equations = self._extract_equations_pymupdf()
        self.figures = self._extract_figures_pymupdf()

    def _extract_text_pymupdf(self) -> str:
        """Extract all text from PDF using PyMuPDF"""
        full_text = []
        for page in self.doc:
            full_text.append(page.get_text())
        return "\n".join(full_text)

    def _extract_equations_pymupdf(self) -> List[Equation]:
        """Basic equation detection in PDF text"""
        equations = []
        equation_pattern = r'\(([\d+-]+)\)'
        for page_num, page in enumerate(self.doc):
            text = page.get_text()
            for line_idx, line in enumerate(text.split("\n")):
                if any(char in line for char in ['=', '∫', '∑', '√', 'μ', 'π']):
                    eq_number_match = re.search(equation_pattern, line)
                    eq_number = eq_number_match.group(0) if eq_number_match else None
                    equations.append(Equation(
                        equation_id=f"eq_p{page_num}_{line_idx}",
                        latex=line.strip(),
                        position=Position(page=page_num),
                        number=eq_number
                    ))
        return equations

    def _extract_figures_pymupdf(self) -> List[Figure]:
        """Extract images from PDF pages using PyMuPDF"""
        figures = []
        for page_num, page in enumerate(self.doc):
            for img_idx, img in enumerate(page.get_images(full=True)):
                try:
                    xref = img[0]
                    base_image = self.doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_bytes = ImageService.compress_for_storage(image_bytes)
                    image_base64 = base64.b64encode(image_bytes).decode("utf-8")
                    figures.append(Figure(
                        figure_id=f"fig_p{page_num}_{img_idx}",
                        image_base64=image_base64,
                        position=Position(page=page_num)
                    ))
                except Exception as e:
                    print(f"Error extracting PDF figure on page {page_num}: {e}")
                    continue
        return figures

    def _extract_equations_from_markdown(self, markdown: str):
        """Extract LaTeX equations from Nougat's markdown output"""
        # Match inline equations: $...$
        inline_pattern = r'\$([^$]+)\$'
        # Match display equations: $$...$$
        display_pattern = r'\$\$([^$]+)\$\$'

        # Extract display equations first
        for idx, match in enumerate(re.finditer(display_pattern, markdown)):
            latex = match.group(1).strip()
            # Calculate approximate position
            char_pos = match.start()
            line_num = markdown[:char_pos].count('\n') + 1
            page_num = self._estimate_page_from_line(line_num)

            self.equations.append(Equation(
                equation_id=f"eq_nougat_{idx}",
                latex=latex,
                position=Position(page=page_num, line=line_num),
                number=None
            ))

        # Extract inline equations
        offset = len(self.equations)
        for idx, match in enumerate(re.finditer(inline_pattern, markdown)):
            latex = match.group(1).strip()
            # Skip if already captured in display equations
            if not any(eq.latex == latex for eq in self.equations):
                # Calculate approximate position
                char_pos = match.start()
                line_num = markdown[:char_pos].count('\n') + 1
                page_num = self._estimate_page_from_line(line_num)

                self.equations.append(Equation(
                    equation_id=f"eq_nougat_inline_{idx}",
                    latex=latex,
                    position=Position(page=page_num, line=line_num),
                    number=None
                ))

    def _extract_tables_from_markdown(self, markdown: str):
        """Extract tables from Nougat's markdown output"""
        # Markdown tables have format:
        # | Header1 | Header2 |
        # |---------|---------|
        # | Cell1   | Cell2   |

        table_pattern = r'(\|.+\|[\r\n]+(?:\|[-:\s|]+\|[\r\n]+)?(?:\|.+\|[\r\n]*)+)'

        for idx, match in enumerate(re.finditer(table_pattern, markdown, re.MULTILINE)):
            table_text = match.group(1).strip()
            rows = [line.strip() for line in table_text.split('\n') if line.strip()]

            # Calculate approximate position
            char_pos = match.start()
            line_num = markdown[:char_pos].count('\n') + 1
            page_num = self._estimate_page_from_line(line_num)

            # Parse markdown table
            table_data = []
            for row in rows:
                if re.match(r'\|[-:\s|]+\|', row):  # Skip separator row
                    continue
                cells = [cell.strip() for cell in row.split('|')[1:-1]]
                table_data.append(cells)

            if table_data:
                self.tables.append(Table(
                    table_id=f"tbl_nougat_{idx}",
                    caption=None,
                    content=table_data,
                    position=Position(page=page_num, line=line_num),
                    number=None
                ))

    def _estimate_page_from_line(self, line_num: int) -> int:
        """Estimate page number from line number (rough approximation)"""
        # Assume ~40-50 lines per page on average
        LINES_PER_PAGE = 45
        page = max(0, (line_num - 1) // LINES_PER_PAGE)
        # Make sure page doesn't exceed total pages
        return min(page, len(self.doc) - 1) if len(self.doc) > 0 else 0

    def _generate_metadata(self) -> DocumentMetadata:
        """Generate PDF document metadata"""
        return DocumentMetadata(
            total_pages=len(self.doc),
            total_paragraphs=0,
            total_equations=len(self.equations),
            total_figures=len(self.figures),
            total_tables=len(self.tables)
        )