from fastapi import APIRouter, Depends, HTTPException, status, BackgroundTasks, Request, UploadFile, File, Form
from typing import List, Optional, Dict
from pydantic import BaseModel, Field
from datetime import datetime
import asyncio
import uuid
import copy
import shutil
import os
import re
import subprocess
import base64

from app.core.security import get_current_user_dep, verify_download_token, decode_token
from app.database.connection import get_database
from app.database.repositories.session_repo import SessionRepository
from app.database.repositories.document_repo import DocumentRepository
from app.models.session import (
    SessionStatus, EditorialRules, OutlineItem, IssueType, Severity,
    DiagnosticSummary, TriggerType, PatchStatus,
)
from app.core.logger import get_logger
from app.core.rate_limit import limiter
from app.core.config import settings
from app.services.image_service import ImageService

logger = get_logger(__name__)

router = APIRouter(prefix="/sessions", tags=["Editorial Pipeline"])


# ── Helpers ──────────────────────────────────────────────────────────────────

def _get_repos(db):
    return SessionRepository(db), DocumentRepository(db)


def _resolve_file_path(stored_path: str) -> str:
    """Resolve a file path from DB, falling back to local UPLOAD_DIR if the stored
    (production) path doesn't exist. This allows the same MongoDB to work on both
    the production server and a local dev machine."""
    if stored_path and os.path.exists(stored_path):
        return stored_path
    if stored_path:
        filename = os.path.basename(stored_path)
        local_path = os.path.join(settings.UPLOAD_DIR, filename)
        if os.path.exists(local_path):
            logger.info("Resolved production path to local: %s -> %s", stored_path, local_path)
            return local_path
        # Also try absolute path from UPLOAD_DIR
        abs_local = os.path.abspath(local_path)
        if os.path.exists(abs_local):
            logger.info("Resolved production path to local: %s -> %s", stored_path, abs_local)
            return abs_local
        logger.warning("File not found locally either: %s (tried %s)", stored_path, abs_local)
    return stored_path  # return original (will fail downstream with a clear error)


async def _get_openai_key(db) -> str:
    """Get OpenAI API key from DB first, then fall back to .env."""
    try:
        doc = await db.settings.find_one({"key": "openai_api_key"})
        if doc and doc.get("value"):
            return doc["value"]
    except Exception:
        pass
    return settings.OPENAI_API_KEY


# ── Request / Response Models ────────────────────────────────────────────────

class CreateSessionRequest(BaseModel):
    document_id: str


class CreateSessionResponse(BaseModel):
    session_id: str
    document_id: str
    status: str
    message: str


class RulesRequest(BaseModel):
    date_cutoff: Optional[str] = None
    preserve_sections: List[str] = []
    voice_preservation: bool = True
    citation_style: str = "inline"
    confidence_threshold: float = Field(default=0.5, ge=0.0, le=1.0)
    allowed_source_types: List[str] = ["government", "academic", "news", "technical", "commercial"]
    excluded_topics: List[str] = []
    max_sentence_change_pct: float = Field(default=80.0, ge=0.0, le=100.0)


class OutlineConfirmRequest(BaseModel):
    selections: Dict[str, bool]  # {outline_item_id: in_scope}


class OpportunitySelectRequest(BaseModel):
    selections: Dict[str, bool]  # {opportunity_id: selected}


class PlanApproveRequest(BaseModel):
    plan_id: str
    approved: bool = True
    rejected: bool = False


class EvidenceDecisionRequest(BaseModel):
    evidence_id: str
    accepted: bool


class PatchReviewRequest(BaseModel):
    action: str  # "approve", "reject", "edit"
    editor_revision: Optional[str] = None


class DatedStatementResolveRequest(BaseModel):
    statement_id: str
    resolution: str  # "still_current", "flag_for_patch", "acceptable"


class CreateCustomIssueRequest(BaseModel):
    title: str
    # issue_kind: "find_replace" (default, legacy behaviour) or "research"
    issue_kind: str = "find_replace"
    # Find & Replace fields (required when issue_kind == "find_replace")
    find_text: Optional[str] = ""
    replace_with: Optional[str] = ""
    # Research fields (required when issue_kind == "research")
    research_prompt: Optional[str] = ""
    scope: str = "whole_document"  # "whole_document" or section name
    severity: str = "medium"  # "high", "medium", "low"


class ApproveResearchRequest(BaseModel):
    """Finalize placement for a research-type custom issue.
    The user has reviewed + (optionally) edited the research draft and
    chosen where in the document it should be inserted.
    """
    content: str  # final edited content
    section_text: str  # heading of the target section
    paragraph_index: int = -1  # -1 = section-level (insert after heading)
    paragraph_text: str = ""  # for robustness: finds paragraph by text
    position: str = "after"  # "before" | "after" | "replace"


class ResetToStepRequest(BaseModel):
    target_status: str  # e.g. "created", "rules_confirmed", "outline_extracted", etc.


# ══════════════════════════════════════════════════════════════════════════════
# SESSION CRUD
# ══════════════════════════════════════════════════════════════════════════════

@router.post("", response_model=CreateSessionResponse)
async def create_session(
    req: CreateSessionRequest,
    user=Depends(get_current_user_dep),
):
    """Find existing active session or create a new one for a document."""
    db = get_database()
    session_repo, doc_repo = _get_repos(db)

    doc = await doc_repo.find_by_id(req.document_id, lightweight=True)
    if not doc:
        raise HTTPException(404, "Document not found")
    if doc.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")
    if doc.get("status") not in ("completed", "export_ready"):
        raise HTTPException(400, "Document must be processed before creating a session")

    # Req 5: Find existing session (not error) before creating new
    # Resume even "exported" sessions so user can navigate all completed stages
    existing_sessions = await session_repo.find_sessions_by_document(req.document_id)
    for s in existing_sessions:
        if s.get("status") != "error":
            logger.info("Resuming existing session %s for document %s (status=%s)", s["id"], req.document_id, s.get("status"))
            return CreateSessionResponse(
                session_id=s["id"],
                document_id=req.document_id,
                status=s["status"],
                message="Resuming existing session.",
            )

    session_data = {
        "document_id": req.document_id,
        "user_id": user["email"],
        "status": SessionStatus.CREATED.value,
        "rules": None,
        "outline": [],
        "diagnostic": None,
        "working_doc_path": None,
        "created_at": datetime.utcnow(),
        "updated_at": datetime.utcnow(),
    }
    session_id = await session_repo.create_session(session_data)

    return CreateSessionResponse(
        session_id=session_id,
        document_id=req.document_id,
        status=SessionStatus.CREATED.value,
        message="Editorial session created. Proceed to define rules.",
    )


@router.get("/{session_id}")
async def get_session(session_id: str, user=Depends(get_current_user_dep)):
    """Get session details including current status."""
    db = get_database()
    session_repo = SessionRepository(db)
    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")
    return session


@router.get("/document/{document_id}")
async def get_sessions_for_document(document_id: str, user=Depends(get_current_user_dep)):
    """List all sessions for a document."""
    db = get_database()
    session_repo = SessionRepository(db)
    sessions = await session_repo.find_sessions_by_document(document_id)
    # Filter to user's sessions unless admin
    if user.get("role") != "admin":
        sessions = [s for s in sessions if s.get("user_id") == user["email"]]
    return {"document_id": document_id, "sessions": sessions}


@router.get("/document/{document_id}/patches")
async def get_document_patches(document_id: str, user=Depends(get_current_user_dep)):
    """Get all patches across all sessions for a document, with session dates."""
    db = get_database()
    session_repo = SessionRepository(db)
    doc_repo = DocumentRepository(db)

    doc = await doc_repo.find_by_id(document_id, lightweight=True)
    if not doc:
        raise HTTPException(404, "Document not found")

    sessions = await session_repo.find_sessions_by_document(document_id)
    if user.get("role") != "admin":
        sessions = [s for s in sessions if s.get("user_id") == user["email"]]

    all_patches = []
    for sess in sessions:
        session_id = sess.get("id", "")
        patches = await session_repo.find_patches(session_id)
        for p in patches:
            p["session_created_at"] = sess.get("created_at")
            p["session_updated_at"] = sess.get("updated_at")
            p["session_status"] = sess.get("status")
        all_patches.extend(patches)

    return {
        "document_id": document_id,
        "document_title": doc.get("metadata", {}).get("title") or doc.get("original_filename", "Untitled"),
        "patches": all_patches,
        "total": len(all_patches),
        "approved": sum(1 for p in all_patches if p.get("status") in ("approved", "edited")),
        "rejected": sum(1 for p in all_patches if p.get("status") == "rejected"),
        "pending": sum(1 for p in all_patches if p.get("status") == "pending"),
    }


@router.delete("/{session_id}")
async def delete_session(session_id: str, user=Depends(get_current_user_dep)):
    """Delete a session and all related data."""
    db = get_database()
    session_repo = SessionRepository(db)
    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")
    await session_repo.delete_session(session_id)
    return {"message": "Session deleted", "session_id": session_id}


# ── Pipeline step order for reset logic ──────────────────────────────────────
_STATUS_ORDER = [
    SessionStatus.CREATED.value,
    SessionStatus.RULES_CONFIRMED.value,
    SessionStatus.OUTLINE_EXTRACTED.value,
    SessionStatus.DIAGNOSTIC_COMPLETE.value,
    SessionStatus.OPPORTUNITIES_SELECTED.value,
    SessionStatus.RESEARCH_PLANNED.value,
    SessionStatus.RESEARCHING.value,
    SessionStatus.RESEARCH_DONE.value,
    SessionStatus.EVIDENCE_REVIEWED.value,
    SessionStatus.PATCHES_GENERATED.value,
    SessionStatus.EDITS_APPLIED.value,
    SessionStatus.EXPORTED.value,
]


@router.post("/{session_id}/reset-to-step")
async def reset_to_step(
    session_id: str,
    req: ResetToStepRequest,
    user=Depends(get_current_user_dep),
):
    """Reset session to a previous step, clearing all downstream data."""
    db = get_database()
    session_repo = SessionRepository(db)
    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    target = req.target_status
    if target not in _STATUS_ORDER:
        raise HTTPException(400, f"Invalid target status: {target}")

    target_idx = _STATUS_ORDER.index(target)
    current_idx = _STATUS_ORDER.index(session["status"]) if session["status"] in _STATUS_ORDER else 0

    if target_idx > current_idx:
        raise HTTPException(400, "Cannot reset forward — only backward")

    # Clear downstream data based on target
    update_fields: dict = {"status": target}

    # If resetting to before diagnostic, clear diagnostic + everything downstream
    if target_idx < _STATUS_ORDER.index(SessionStatus.DIAGNOSTIC_COMPLETE.value):
        update_fields["diagnostic"] = None
        await session_repo.delete_opportunities(session_id)

    # If resetting to before research planning, clear plans + downstream
    if target_idx < _STATUS_ORDER.index(SessionStatus.RESEARCH_PLANNED.value):
        await session_repo.delete_research_plans(session_id)

    # If resetting to before evidence, clear evidence + downstream
    if target_idx < _STATUS_ORDER.index(SessionStatus.EVIDENCE_REVIEWED.value):
        await session_repo.delete_evidence_items(session_id)

    # If resetting to before patches, clear patches + downstream
    if target_idx < _STATUS_ORDER.index(SessionStatus.PATCHES_GENERATED.value):
        await session_repo.delete_patches(session_id)

    # If resetting to before apply, clear working doc
    if target_idx < _STATUS_ORDER.index(SessionStatus.EDITS_APPLIED.value):
        update_fields["working_doc_path"] = None

    await session_repo.update_session(session_id, update_fields)
    logger.info("Session %s reset to step: %s", session_id, target)
    return {"message": f"Session reset to {target}", "session_id": session_id, "status": target}


# ══════════════════════════════════════════════════════════════════════════════
# STAGE 1: RULES CONFIRMATION
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/{session_id}/rules")
async def set_rules(
    session_id: str,
    req: RulesRequest,
    user=Depends(get_current_user_dep),
):
    """Define editorial rules for the session."""
    db = get_database()
    session_repo = SessionRepository(db)
    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    rules = req.model_dump()
    await session_repo.update_session(session_id, {
        "rules": rules,
        "status": SessionStatus.RULES_CONFIRMED.value,
    })

    return {
        "session_id": session_id,
        "status": SessionStatus.RULES_CONFIRMED.value,
        "rules": rules,
        "message": "Rules confirmed. Proceed to extract outline.",
    }


# ══════════════════════════════════════════════════════════════════════════════
# STAGE 2: OUTLINE EXTRACTION
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/{session_id}/extract-outline")
async def extract_outline(
    session_id: str,
    user=Depends(get_current_user_dep),
):
    """Extract document outline (headings) for scope selection."""
    db = get_database()
    session_repo = SessionRepository(db)
    doc_repo = DocumentRepository(db)

    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")
    # Allow re-running if session is at or past rules_confirmed
    current_status = session.get("status")
    if current_status in _STATUS_ORDER:
        if _STATUS_ORDER.index(current_status) < _STATUS_ORDER.index(SessionStatus.RULES_CONFIRMED.value):
            raise HTTPException(400, "Rules must be confirmed first")
    else:
        raise HTTPException(400, "Rules must be confirmed first")

    # ── Load paragraph metadata from DB (extracted once during POST /process) ──
    doc = await doc_repo.find_with_paragraphs(session["document_id"])
    if not doc:
        raise HTTPException(404, "Document not found")

    db_paragraphs = doc.get("paragraphs", [])
    if not db_paragraphs:
        raise HTTPException(400, "Document has not been processed yet. Run POST /process first.")

    outline_items = []

    import re as _re
    import json as _json
    from collections import Counter, defaultdict

    def _strip_trailing_page_number(txt: str) -> str:
        cleaned = _re.sub(r'\s+\d{1,4}\s*$', '', txt)
        return cleaned if cleaned else txt

    def _normalize_heading(txt: str) -> str:
        cleaned = _strip_trailing_page_number(txt)
        return _re.sub(r'\s+', ' ', cleaned).strip().lower()

    # ── Detect which approach to use (all from DB paragraphs) ─────
    has_heading_styles = any(
        _re.match(r"[Hh]eading\s*#?\s*(\d+)", p.get("style", ""))
        for p in db_paragraphs if p.get("text")
    )

    # Collect all paragraphs matching X.X pattern (e.g., "1.1 Introduction")
    numbered_candidates = []
    for p in db_paragraphs:
        text = p.get("text", "")
        if not text:
            continue
        numbered_match = _re.match(r"^(\d+(?:\.\d+)+)\s+\S", text)
        if numbered_match and p.get("length", 0) < 150:
            numbered_candidates.append(p)

    has_numbered_headings = len(numbered_candidates) >= 3

    try:
        # ── APPROACH 1 (Fast path): Regex on numbered X.X patterns ────
        if has_numbered_headings:
            logger.info("Outline: Found %d numbered heading candidates — using regex (no GPT)", len(numbered_candidates))

            # ── Filter 1: Detect page headers (text with trailing tab/spaces + page number) ──
            # Pattern: "1.1\tIntroduction and Overview\t5" or "1.2 The Life Cycle 9"
            # These have a number at the very end separated by whitespace/tab
            _page_header_re = _re.compile(r'^.+[\t\s]+\d{1,4}\s*$')

            # ── Filter 2: Detect TOC region ──
            # TOC entries cluster at the start (first ~20 non-empty paragraphs), are non-bold
            # Find first bold numbered candidate to mark end of TOC region
            first_bold_numbered_idx = None
            for p in numbered_candidates:
                if p.get("bold", False):
                    first_bold_numbered_idx = p.get("idx", 0)
                    break

            # ── Filter 3: Build normalized text → best candidate map ──
            # When same heading appears multiple times (TOC + real + page header),
            # prefer: bold without trailing page number > bold with trailing > non-bold
            heading_groups = defaultdict(list)  # normalized_text -> [candidate, ...]
            for p in numbered_candidates:
                text = p.get("text", "")
                normalized = _normalize_heading(text)
                heading_groups[normalized].append(p)

            seen_heading_texts = set()
            for normalized, group in heading_groups.items():
                # Pick the BEST candidate from the group
                # Priority: bold + no page-header pattern + not in TOC region
                best = None
                for p in group:
                    text = p.get("text", "")
                    idx = p.get("idx", 0)
                    is_bold = p.get("bold", False)
                    is_page_header = bool(_page_header_re.match(text))
                    is_toc = (not is_bold and first_bold_numbered_idx is not None
                              and idx < first_bold_numbered_idx)

                    # Skip header/footer/caption/toc styles
                    style_lower = p.get("style", "").lower()
                    if any(skip in style_lower for skip in ("header", "footer", "caption", "toc")):
                        continue

                    # Score: higher is better
                    score = 0
                    if is_bold:
                        score += 10
                    if not is_page_header:
                        score += 5
                    if not is_toc:
                        score += 3

                    if best is None or score > best[0]:
                        best = (score, p)

                if best is None:
                    continue

                p = best[1]
                text = p.get("text", "")
                idx = p.get("idx", 0)

                # Parse numbering depth for level
                num_match = _re.match(r"^(\d+(?:\.\d+)+)", text)
                if not num_match:
                    continue
                num_str = num_match.group(1)
                depth = num_str.count(".")
                if depth <= 1:
                    level = 1
                elif depth == 2:
                    level = 2
                else:
                    level = 3

                seen_heading_texts.add(normalized)
                clean_text = _strip_trailing_page_number(text)
                # Also strip trailing tabs/whitespace that may remain
                clean_text = _re.sub(r'[\t]+', ' ', clean_text).strip()
                outline_items.append({
                    "id": str(uuid.uuid4())[:8],
                    "text": clean_text,
                    "level": level,
                    "in_scope": True,
                    "paragraph_index": idx,
                })

            # Sort by paragraph index to maintain document order
            outline_items.sort(key=lambda x: x["paragraph_index"])

        # ── APPROACH 2: Heading styles + GPT filter ───────────────────
        elif has_heading_styles:
            logger.info("Outline: Using Heading styles with GPT filter")
            raw_headings = []
            for p in db_paragraphs:
                text = p.get("text", "")
                if not text:
                    continue
                heading_match = _re.match(r"[Hh]eading\s*#?\s*(\d+)", p.get("style", ""))
                if heading_match:
                    raw_headings.append({
                        "idx": p["idx"],
                        "text": text[:200],
                        "style": p.get("style", ""),
                        "heading_level": int(heading_match.group(1)),
                        "bold": p.get("bold", False),
                        "font_sizes": p.get("font_sizes", []),
                        "length": p.get("length", 0),
                    })

            from openai import AsyncOpenAI
            client = AsyncOpenAI(api_key=(await _get_openai_key(db)))

            gpt_filter_prompt = (
                "You are analyzing heading-styled paragraphs from a DOCX document.\n"
                "Your job is to identify which are REAL content section headings and which are NOT.\n\n"
                "INCLUDE only real content section headings — these are the main topics/sections of the document.\n"
                "Examples of real headings: '2.1 Early Space Explorers', 'Astronomy Begins', 'The Age of Rockets'\n\n"
                "EXCLUDE these types (they are NOT real section headings):\n"
                "- Book titles, chapter titles that are just a title without section content below\n"
                "- Instructional/callout boxes: 'In This Section You will Learn To...', 'You Should Already Know...'\n"
                "- Review/exercise sections: 'Section Review', 'Mission Problems', 'For Discussion'\n"
                "- Fun fact/sidebar callouts: 'Astro Fun Fact...', any callout box headings\n"
                "- Supplementary sections: 'For Further Reading', 'References', 'Contributor'\n"
                "- Case study/profile meta-headings: 'Mission Overview', 'Mission Data', 'Mission Impact'\n"
                "- The word 'Outline' by itself\n"
                "- Any heading that is clearly metadata, not a content section\n\n"
                "For each heading you keep, assign a level:\n"
                "- level 1: Major numbered section (e.g. '2.1 Early Space Explorers')\n"
                "- level 2: Sub-section (e.g. 'Astronomy Begins', 'The Age of Rockets')\n"
                "- level 3: Sub-sub-section\n"
                "- If the heading has a section number like X.X, use that for level\n"
                "- If no number, determine level from context and heading_level in the data\n\n"
                "Here are the heading-styled paragraphs:\n\n"
                f"{_json.dumps(raw_headings, indent=None)}\n\n"
                "Return ONLY a JSON array of the real section headings, each with:\n"
                '{"idx": <paragraph_index>, "text": "<heading text>", "level": <1|2|3>}\n\n'
                "Return ONLY the JSON array, no explanation or markdown. "
                "If none qualify, return []."
            )

            try:
                response = await client.chat.completions.create(
                    model=settings.GPT_MODEL,
                    messages=[{"role": "user", "content": gpt_filter_prompt}],
                    temperature=0.0,
                    max_tokens=4000,
                )
                raw = response.choices[0].message.content.strip()
                if raw.startswith("```"):
                    raw = _re.sub(r'^```(?:json)?\s*', '', raw)
                    raw = _re.sub(r'\s*```$', '', raw)

                gpt_headings = _json.loads(raw)
                logger.info("GPT filtered to %d real section headings", len(gpt_headings))

                seen_heading_texts = set()
                for h in gpt_headings:
                    text = h.get("text", "").strip()
                    idx = h.get("idx", 0)
                    level = h.get("level", 1)
                    if not text:
                        continue
                    clean_text = _strip_trailing_page_number(text)
                    normalized = _normalize_heading(text)
                    if normalized in seen_heading_texts:
                        continue
                    seen_heading_texts.add(normalized)
                    outline_items.append({
                        "id": str(uuid.uuid4())[:8],
                        "text": clean_text,
                        "level": level,
                        "in_scope": True,
                        "paragraph_index": idx,
                    })

            except Exception as gpt_err:
                logger.error("GPT heading filter failed: %s — using all headings as fallback", gpt_err)
                seen_heading_texts = set()
                for rh in raw_headings:
                    normalized = _normalize_heading(rh["text"])
                    if normalized in seen_heading_texts:
                        continue
                    seen_heading_texts.add(normalized)
                    outline_items.append({
                        "id": str(uuid.uuid4())[:8],
                        "text": _strip_trailing_page_number(rh["text"]),
                        "level": rh["heading_level"],
                        "in_scope": True,
                        "paragraph_index": rh["idx"],
                    })

        # ── APPROACH 3 (Rare): No numbered headings AND no Heading styles ─
        else:
            logger.info("Outline: No numbered headings and no Heading styles — GPT fallback")

            para_metadata = []
            for p in db_paragraphs:
                text = p.get("text", "")
                if not text:
                    continue
                style_lower = p.get("style", "").lower()
                if any(skip in style_lower for skip in ("header", "footer", "caption", "toc")):
                    continue
                if _re.match(r'^\d+(\.\d+)*\s*$', text):
                    continue
                para_metadata.append({
                    "idx": p["idx"],
                    "text": text[:200],
                    "bold": p.get("bold", False),
                    "font_sizes": p.get("font_sizes", []),
                    "style": p.get("style", ""),
                    "length": p.get("length", 0),
                })

            para_metadata = para_metadata[:300]

            from openai import AsyncOpenAI
            client = AsyncOpenAI(api_key=(await _get_openai_key(db)))

            gpt_prompt = (
                "You are analyzing a DOCX document to identify section headings.\n"
                "The document has no Heading styles and no numbered headings.\n"
                "Your job is to identify which paragraphs are REAL section headings.\n\n"
                "STRICT RULES:\n"
                "1. Section headings are typically: short (< 100 chars), bold, larger font\n"
                "2. DO NOT include: body text, TOC entries, page headers, author names\n"
                "3. Assign levels: level 1 for main sections, level 2 for sub-sections, level 3 for sub-sub\n\n"
                "Here are the paragraphs:\n\n"
                f"{_json.dumps(para_metadata, indent=None)}\n\n"
                "Return ONLY a JSON array of heading objects:\n"
                '{"idx": <paragraph_index>, "text": "<heading text>", "level": <1|2|3>}\n\n'
                "Return ONLY the JSON array. If none found, return []."
            )

            try:
                response = await client.chat.completions.create(
                    model=settings.GPT_MODEL,
                    messages=[{"role": "user", "content": gpt_prompt}],
                    temperature=0.0,
                    max_tokens=4000,
                )
                raw = response.choices[0].message.content.strip()
                if raw.startswith("```"):
                    raw = _re.sub(r'^```(?:json)?\s*', '', raw)
                    raw = _re.sub(r'\s*```$', '', raw)

                gpt_headings = _json.loads(raw)
                logger.info("GPT identified %d headings (rare fallback)", len(gpt_headings))

                seen_heading_texts = set()
                for h in gpt_headings:
                    text = h.get("text", "").strip()
                    idx = h.get("idx", 0)
                    level = h.get("level", 1)
                    if not text:
                        continue
                    clean_text = _strip_trailing_page_number(text)
                    normalized = _normalize_heading(text)
                    if normalized in seen_heading_texts:
                        continue
                    seen_heading_texts.add(normalized)
                    outline_items.append({
                        "id": str(uuid.uuid4())[:8],
                        "text": clean_text,
                        "level": level,
                        "in_scope": True,
                        "paragraph_index": idx,
                    })

            except Exception as gpt_err:
                logger.error("GPT heading detection failed: %s — falling back to heuristics", gpt_err)
                seen_heading_texts = set()
                for pm in para_metadata:
                    text = pm["text"]
                    idx = pm["idx"]
                    numbered = _re.match(r"^(\d+(?:\.\d+)*)\s+[A-Z]", text)
                    if numbered and pm["bold"] and pm["length"] < 120:
                        depth = numbered.group(1).count(".")
                        if depth <= 1:
                            level = 1
                        elif depth == 2:
                            level = 2
                        else:
                            level = 3
                        clean_text = _strip_trailing_page_number(text)
                        normalized = _normalize_heading(text)
                        if normalized in seen_heading_texts:
                            continue
                        seen_heading_texts.add(normalized)
                        outline_items.append({
                            "id": str(uuid.uuid4())[:8],
                            "text": clean_text,
                            "level": level,
                            "in_scope": True,
                            "paragraph_index": idx,
                        })

    except Exception as e:
        logger.error("Outline extraction failed: %s", e)
        # Ultimate fallback: regex on text_content
        text_content = doc.get("text_content", "")
        lines = text_content.split("\n")
        for idx, line in enumerate(lines):
            line_stripped = line.strip()
            if _re.match(r"^\d+(\.\d+)*\s+[A-Z]", line_stripped) or (
                line_stripped.isupper() and 3 < len(line_stripped) < 100
            ):
                outline_items.append({
                    "id": str(uuid.uuid4())[:8],
                    "text": line_stripped,
                    "level": 1 if _re.match(r"^\d+\s", line_stripped) else 2,
                    "in_scope": True,
                    "paragraph_index": idx,
                })

    # ── Final sort: by section number (e.g. 7.1 < 7.2 < 7.3), then paragraph_index ──
    def _section_sort_key(item):
        m = _re.match(r'^(\d+(?:\.\d+)*)', item.get("text", ""))
        if m:
            # Convert "7.2.1" → (7, 2, 1) for proper numeric sorting
            return tuple(int(x) for x in m.group(1).split('.'))
        return (item.get("paragraph_index", 0),)

    outline_items.sort(key=_section_sort_key)

    await session_repo.update_session(session_id, {
        "outline": outline_items,
        "status": SessionStatus.OUTLINE_EXTRACTED.value,
    })

    return {
        "session_id": session_id,
        "status": SessionStatus.OUTLINE_EXTRACTED.value,
        "outline": outline_items,
        "total_sections": len(outline_items),
        "message": "Outline extracted. Select sections to analyze.",
    }


@router.put("/{session_id}/outline")
async def confirm_outline(
    session_id: str,
    req: OutlineConfirmRequest,
    user=Depends(get_current_user_dep),
):
    """Confirm which sections are in scope."""
    db = get_database()
    session_repo = SessionRepository(db)
    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    outline = session.get("outline", [])
    for item in outline:
        if item["id"] in req.selections:
            item["in_scope"] = req.selections[item["id"]]

    await session_repo.update_session(session_id, {"outline": outline})

    in_scope = sum(1 for i in outline if i["in_scope"])
    return {
        "session_id": session_id,
        "in_scope_count": in_scope,
        "total_sections": len(outline),
        "message": f"{in_scope} sections selected for analysis.",
    }


# ══════════════════════════════════════════════════════════════════════════════
# STAGE 3: DIAGNOSTIC REVIEW
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/{session_id}/run-diagnostic")
async def run_diagnostic(
    session_id: str,
    background_tasks: BackgroundTasks,
    user=Depends(get_current_user_dep),
):
    """Run AI diagnostic to identify issues in selected sections."""
    db = get_database()
    session_repo = SessionRepository(db)
    doc_repo = DocumentRepository(db)

    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")
    # Allow re-running if session is at or past outline_extracted
    current_status = session.get("status")
    if current_status in _STATUS_ORDER:
        if _STATUS_ORDER.index(current_status) < _STATUS_ORDER.index(SessionStatus.OUTLINE_EXTRACTED.value):
            raise HTTPException(400, "Outline must be extracted first")
    else:
        raise HTTPException(400, "Outline must be extracted first")

    doc = await doc_repo.find_by_id(session["document_id"], analysis_mode=True)
    if not doc:
        raise HTTPException(404, "Document not found")

    background_tasks.add_task(
        _run_diagnostic_task, session_id, session, doc
    )

    return {
        "session_id": session_id,
        "status": "running",
        "message": "Diagnostic analysis started. Poll session status for updates.",
    }


async def _run_diagnostic_task(session_id: str, session: dict, doc: dict):
    """Background task: scan each in-scope section individually for thorough coverage."""
    from openai import AsyncOpenAI
    import json
    import asyncio

    db = get_database()
    session_repo = SessionRepository(db)

    try:
        client = AsyncOpenAI(api_key=(await _get_openai_key(db)))
        rules = session.get("rules", {}) or {}
        outline = session.get("outline", [])
        text_content = doc.get("text_content", "")

        # Get in-scope section texts
        in_scope_headings = [o for o in outline if o.get("in_scope", True)]

        # Build individual section texts from paragraph indices
        paragraphs = text_content.split("\n")
        section_texts = []
        for i, heading in enumerate(in_scope_headings):
            start = heading.get("paragraph_index", 0)
            if i + 1 < len(in_scope_headings):
                end = in_scope_headings[i + 1].get("paragraph_index", len(paragraphs))
            else:
                end = len(paragraphs)
            section_text = "\n".join(paragraphs[start:end]).strip()
            if len(section_text) > 50:  # skip tiny/empty sections
                section_texts.append({
                    "section": heading.get("text", ""),
                    "text": section_text[:8000],
                })

        # Build rule instructions
        confidence_threshold = rules.get("confidence_threshold", 0.5)

        rule_lines = []
        if rules.get("date_cutoff"):
            rule_lines.append(f"- Flag facts that may be outdated relative to {rules['date_cutoff']}.")
        else:
            rule_lines.append("- Flag any facts, statistics, figures, or claims that may be outdated or no longer accurate.")
        if rules.get("excluded_topics"):
            rule_lines.append(f"- Do NOT flag issues related to these topics: {', '.join(rules['excluded_topics'])}.")
        if rules.get("preserve_sections"):
            rule_lines.append(f"- Do NOT flag issues in these sections: {', '.join(rules['preserve_sections'])}.")
        rules_block = "\n".join(rule_lines)

        system_msg = "You are a meticulous document auditor. Your job is to find EVERY sentence that contains potentially outdated information, stale references, temporal language, or changed real-world status. Be extremely thorough — missing an issue is worse than flagging one that turns out to be fine. IMPORTANT: Never spell out Greek letters — always preserve the original symbols (Δ, α, β, θ, Σ, etc.) and subscript/superscript notation exactly as written. Return only valid JSON."

        # Delete old opportunities for re-run
        await session_repo.delete_opportunities(session_id)

        # ── Scan each section with its own AI call ──
        api_errors = []  # Track API-level failures (quota, auth)

        async def scan_section(section_data: dict) -> list:
            section_name = section_data["section"]
            section_text = section_data["text"]

            prompt = f"""Audit this section of a document for content that may need updating.

Section: "{section_name}"

Rules:
{rules_block}

Read every sentence below carefully and flag ALL of the following:
1. **Outdated facts** — statistics, numbers, figures, counts, mission outcomes, technology specs, organizational names, or claims that may have changed since the text was written.
2. **Broken or stale citations** — references to URLs, reports, or sources that may no longer exist or have been superseded.
3. **Date references** — explicit dates ("as of 2018", "in 2015", "planned for 2020") and temporal language ("currently", "recently", "upcoming", "today", "now") that may now be inaccurate.
4. **Changed status** — anything described as "planned", "proposed", "under development", "upcoming", "new", "latest" that may have since launched, been completed, failed, been cancelled, merged, shut down, or otherwise changed.

For EACH issue found, return a JSON object:
- "sentence": the EXACT sentence from the text (copy it word-for-word, do not paraphrase)
- "section_ref": "{section_name}"
- "issue_type": one of "outdated_fact", "broken_citation", "date_reference", "changed_status"
- "severity": "high" (very likely wrong/outdated), "medium" (probably needs checking), or "low" (minor/cosmetic)
- "brief_reason": one-line explanation of why this may need updating
- "confidence": float 0.0-1.0

Only include issues with confidence >= {confidence_threshold}.
Be thorough. Check EVERY sentence for dates, numbers, statistics, status words, and temporal language.
Return ONLY a valid JSON array. No markdown, no explanation. If no issues found, return [].

Text:
{section_text}"""

            try:
                response = await client.chat.completions.create(
                    model=settings.GPT_MODEL,
                    messages=[
                        {"role": "system", "content": system_msg},
                        {"role": "user", "content": prompt},
                    ],
                    temperature=0.2,
                    max_tokens=4000,
                )
                raw = response.choices[0].message.content.strip()
                if raw.startswith("```"):
                    raw = raw.split("\n", 1)[1].rsplit("```", 1)[0]
                return json.loads(raw)
            except Exception as e:
                error_str = str(e)
                logger.warning("Diagnostic failed for section '%s': %s", section_name, e)
                # Track quota/auth errors to surface to the user
                if "insufficient_quota" in error_str or "429" in error_str:
                    api_errors.append("quota")
                elif "invalid_api_key" in error_str or "401" in error_str:
                    api_errors.append("auth")
                else:
                    api_errors.append("other")
                return []

        # Run sections in parallel batches of 3 to avoid rate limits
        all_issues = []
        batch_size = 3
        for batch_start in range(0, len(section_texts), batch_size):
            batch = section_texts[batch_start:batch_start + batch_size]
            results = await asyncio.gather(*[scan_section(s) for s in batch])
            for issues_list in results:
                if isinstance(issues_list, list):
                    all_issues.extend(issues_list)

        # If ALL sections failed due to API errors, surface the error
        if api_errors and len(api_errors) >= len(section_texts):
            if "quota" in api_errors:
                error_msg = (
                    "OpenAI API quota exceeded. Your API key has no remaining credits. "
                    "Please check your billing at platform.openai.com or update your API key in Admin → API Keys."
                )
            elif "auth" in api_errors:
                error_msg = (
                    "OpenAI API key is invalid or expired. "
                    "Please update your API key in Admin → API Keys."
                )
            else:
                error_msg = "OpenAI API requests failed for all sections. Please try again later."
            await session_repo.update_session(session_id, {
                "status": SessionStatus.ERROR.value,
                "error_message": error_msg,
            })
            logger.error("Diagnostic aborted for session %s: %s", session_id, error_msg)
            return

        # Deduplicate by sentence text
        seen_sentences = set()
        opportunities = []
        for issue in all_issues:
            sentence = issue.get("sentence", "").strip()
            if not sentence or sentence in seen_sentences:
                continue
            conf = issue.get("confidence", 0.5)
            if conf < confidence_threshold:
                continue
            seen_sentences.add(sentence)
            opportunities.append({
                "opportunity_id": str(uuid.uuid4())[:8],
                "session_id": session_id,
                "section_ref": issue.get("section_ref", ""),
                "original_sentence": sentence,
                "issue_type": issue.get("issue_type", "outdated_fact"),
                "severity": issue.get("severity", "medium"),
                "confidence": conf,
                "brief_reason": issue.get("brief_reason", ""),
                "selected": False,
            })

        await session_repo.create_opportunities(opportunities)

        # Build summary
        high = sum(1 for o in opportunities if o["severity"] == "high")
        medium = sum(1 for o in opportunities if o["severity"] == "medium")
        low = sum(1 for o in opportunities if o["severity"] == "low")
        by_type = {}
        for o in opportunities:
            t = o["issue_type"]
            by_type[t] = by_type.get(t, 0) + 1

        diagnostic = {
            "total_issues": len(opportunities),
            "high_count": high,
            "medium_count": medium,
            "low_count": low,
            "by_type": by_type,
        }

        await session_repo.update_session(session_id, {
            "diagnostic": diagnostic,
            "status": SessionStatus.DIAGNOSTIC_COMPLETE.value,
        })

        logger.info("Diagnostic complete for session %s: %d issues found across %d sections",
                     session_id, len(opportunities), len(section_texts))

    except Exception as e:
        logger.error("Diagnostic failed for session %s: %s", session_id, e)
        await session_repo.update_session(session_id, {
            "status": SessionStatus.ERROR.value,
            "error_message": str(e),
        })


@router.get("/{session_id}/diagnostic")
async def get_diagnostic(session_id: str, user=Depends(get_current_user_dep)):
    """Get diagnostic summary and opportunities."""
    db = get_database()
    session_repo = SessionRepository(db)
    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    opportunities = await session_repo.find_opportunities(session_id)
    result = {
        "session_id": session_id,
        "status": session.get("status"),
        "diagnostic": session.get("diagnostic"),
        "opportunities": opportunities,
    }
    if session.get("error_message"):
        result["error_message"] = session["error_message"]
    return result


# ══════════════════════════════════════════════════════════════════════════════
# STAGE 4: OPPORTUNITY SELECTION
# ══════════════════════════════════════════════════════════════════════════════

@router.put("/{session_id}/opportunities")
async def select_opportunities(
    session_id: str,
    req: OpportunitySelectRequest,
    user=Depends(get_current_user_dep),
):
    """Select which opportunities to pursue."""
    db = get_database()
    session_repo = SessionRepository(db)
    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    total = await session_repo.update_opportunity_selections(session_id, req.selections)

    selected_count = sum(1 for v in req.selections.values() if v)
    if selected_count == 0:
        raise HTTPException(400, "At least one opportunity must be selected")

    await session_repo.update_session(session_id, {
        "status": SessionStatus.OPPORTUNITIES_SELECTED.value,
    })

    return {
        "session_id": session_id,
        "status": SessionStatus.OPPORTUNITIES_SELECTED.value,
        "updated": total,
        "selected_count": selected_count,
        "message": f"{selected_count} opportunities selected. Proceed to research planning.",
    }


# ── Custom Issue Creation ─────────────────────────────────────────────────────

@router.post("/{session_id}/custom-issue")
async def create_custom_issue(
    session_id: str,
    req: CreateCustomIssueRequest,
    user=Depends(get_current_user_dep),
):
    """Create a custom issue added manually by the user.

    Two kinds:
    - "find_replace" (default): legacy find-and-replace behaviour.
    - "research": user provides a research prompt; the agent will later run
      web research + LLM synthesis and save drafted content onto this
      opportunity for the user to place in Step 7.
    """
    db = get_database()
    session_repo = SessionRepository(db)
    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    issue_kind = (req.issue_kind or "find_replace").lower()
    if issue_kind not in ("find_replace", "research"):
        raise HTTPException(400, "issue_kind must be 'find_replace' or 'research'")

    if issue_kind == "find_replace":
        if not (req.find_text or "").strip() or not (req.replace_with or "").strip():
            raise HTTPException(400, "find_text and replace_with are required for find_replace")
        custom_opp = {
            "opportunity_id": str(uuid.uuid4())[:8],
            "session_id": session_id,
            "section_ref": req.scope if req.scope != "whole_document" else "Whole Document",
            "original_sentence": req.find_text,
            "issue_type": "custom_replacement",
            "severity": req.severity,
            "confidence": 1.0,
            "brief_reason": req.title,
            "selected": True,
            "is_custom": True,
            "custom_kind": "find_replace",
            "replace_with": req.replace_with,
            "scope": req.scope,
        }
    else:  # research
        if not (req.research_prompt or "").strip():
            raise HTTPException(400, "research_prompt is required for research")
        custom_opp = {
            "opportunity_id": str(uuid.uuid4())[:8],
            "session_id": session_id,
            "section_ref": "Research: " + req.title[:60],
            "original_sentence": "",
            "issue_type": "custom_research",
            "severity": req.severity,
            "confidence": 1.0,
            "brief_reason": req.title,
            "selected": True,
            "is_custom": True,
            "custom_kind": "research",
            "research_prompt": req.research_prompt,
            # filled in later by _plan_research_task:
            "research_result": "",
            "research_citations": [],
            "research_status": "pending",
            # filled in later when user approves placement:
            "research_placement": None,
            "scope": "research_insert",
        }

    await session_repo.create_opportunities([custom_opp])
    if issue_kind == "find_replace":
        logger.info("Custom find/replace created for session %s: '%s' → '%s'",
                    session_id, (req.find_text or "")[:50], (req.replace_with or "")[:50])
    else:
        logger.info("Custom research issue created for session %s: '%s'",
                    session_id, req.title[:60])

    # Remove MongoDB _id (ObjectId) before returning — FastAPI can't serialize it
    custom_opp.pop("_id", None)

    return {
        "session_id": session_id,
        "opportunity": custom_opp,
        "message": "Custom issue created successfully.",
    }


@router.get("/{session_id}/custom-research-issues")
async def list_custom_research_issues(
    session_id: str,
    user=Depends(get_current_user_dep),
):
    """Return all research-kind custom issues for a session, including the
    current research_result / status so the Step 7 UI can render them.
    """
    db = get_database()
    session_repo = SessionRepository(db)
    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    opps = await session_repo.find_opportunities(session_id)
    out = []
    for opp in opps:
        if not opp.get("is_custom"):
            continue
        if opp.get("custom_kind") != "research":
            continue
        opp.pop("_id", None)
        out.append(opp)
    return {"session_id": session_id, "issues": out, "total": len(out)}


@router.post("/{session_id}/custom-research-issues/{opportunity_id}/approve")
async def approve_custom_research_issue(
    session_id: str,
    opportunity_id: str,
    req: ApproveResearchRequest,
    user=Depends(get_current_user_dep),
):
    """Finalize a research-type custom issue: the user has reviewed (and
    possibly edited) the drafted content and chosen where to insert it in
    the document. We create a pending patch with `research_insert_meta`
    that the clean + tracked exports apply at the chosen location.
    """
    db = get_database()
    session_repo = SessionRepository(db)
    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")
    if req.position not in ("before", "after", "replace"):
        raise HTTPException(400, "position must be 'before', 'after', or 'replace'")
    if not (req.content or "").strip():
        raise HTTPException(400, "content must not be empty")
    if not (req.section_text or "").strip():
        raise HTTPException(400, "section_text is required")

    opps = await session_repo.find_opportunities(session_id)
    opp = next(
        (o for o in opps if o.get("opportunity_id") == opportunity_id and o.get("custom_kind") == "research"),
        None,
    )
    if not opp:
        raise HTTPException(404, "Research custom issue not found")

    research_insert_meta = {
        "content": req.content.strip(),
        "section_text": req.section_text.strip(),
        "paragraph_index": req.paragraph_index,
        "paragraph_text": (req.paragraph_text or "").strip(),
        "position": req.position,
        "title": opp.get("brief_reason", ""),
        "citations": opp.get("research_citations", []),
    }

    pos_label = {"before": "Before", "after": "After", "replace": "Replace"}[req.position]
    label = opp.get("brief_reason", "Custom research")

    # Remove any prior research-insert patches for this opportunity so repeat
    # approvals replace rather than stack.
    existing_patches = await session_repo.find_patches(session_id)
    for p in existing_patches:
        if (
            p.get("opportunity_id") == opportunity_id
            and p.get("research_insert_meta") is not None
        ):
            try:
                await session_repo.delete_patch(session_id, p.get("patch_id"))
            except Exception:
                pass

    patch = {
        "patch_id": str(uuid.uuid4())[:8],
        "opportunity_id": opportunity_id,
        "session_id": session_id,
        "original_sentence": (
            f"[Research insert {pos_label.lower()} paragraph: {req.paragraph_text[:60]}]"
            if req.paragraph_index >= 0 and req.paragraph_text.strip()
            else f"[Research insert {pos_label.lower()} section: {req.section_text[:60]}]"
        ),
        "revised_sentence": req.content.strip()[:400],
        "citation": "; ".join(
            f"[{c.get('index')}] {c.get('url', '')}"
            for c in opp.get("research_citations", [])[:5]
        ) or "User-directed research",
        "rationale": f"User research: {label}",
        "confidence": 1.0,
        "change_pct": 100.0,
        # Already approved by user in the research review card — no second
        # review needed in the regular patch list.
        "status": PatchStatus.APPROVED.value,
        "editor_revision": None,
        "reviewed_at": datetime.utcnow().isoformat(),
        "section_ref": (
            f"{pos_label} paragraph in section: {req.section_text}"
            if req.paragraph_index >= 0 and req.paragraph_text.strip()
            else f"{pos_label} section: {req.section_text}"
        ),
        "is_custom": True,
        "custom_kind": "research",
        "research_insert_meta": research_insert_meta,
    }

    await session_repo.create_patches([patch])

    # Persist the placement + edited content on the opportunity
    await session_repo.update_opportunity(
        session_id,
        opportunity_id,
        {
            "research_placement": research_insert_meta,
            "research_status": "approved",
            "research_result": req.content.strip(),
        },
    )
    patch.pop("_id", None)

    logger.info(
        "Custom research issue approved for session %s: opp=%s section='%s' position=%s",
        session_id, opportunity_id, req.section_text[:40], req.position,
    )

    return {
        "session_id": session_id,
        "patch": patch,
        "message": "Research content queued for insertion.",
    }


# ══════════════════════════════════════════════════════════════════════════════
# STAGE 5: RESEARCH PLANNING
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/{session_id}/plan-research")
async def plan_research(
    session_id: str,
    background_tasks: BackgroundTasks,
    user=Depends(get_current_user_dep),
):
    """AI generates research plans for each selected opportunity."""
    db = get_database()
    session_repo = SessionRepository(db)
    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")
    # Allow re-running if session is at or past opportunities_selected
    current_status = session.get("status")
    if current_status in _STATUS_ORDER:
        if _STATUS_ORDER.index(current_status) < _STATUS_ORDER.index(SessionStatus.OPPORTUNITIES_SELECTED.value):
            raise HTTPException(400, "Opportunities must be selected first")
    else:
        raise HTTPException(400, "Opportunities must be selected first")

    background_tasks.add_task(_plan_research_task, session_id, session)

    return {
        "session_id": session_id,
        "status": "planning",
        "message": "Research planning started. Poll session for updates.",
    }


async def _run_custom_research(
    client, research_prompt: str, title: str, allowed_sources: list,
    chapter_text: str = "",
) -> tuple[str, list]:
    """Run web research for a user-defined research topic and synthesize it
    into drafted content suitable for insertion into a book chapter.

    Returns (draft_content, citations). Draft is a single multi-paragraph
    string. Citations is a list of {title, url, source_type} dicts.

    When *chapter_text* is provided, GPT receives an excerpt of the chapter
    so that generated content is relevant to the document (not generic).

    Uses Tavily for web search (same provider as the automated research
    pipeline) and GPT for synthesis. If Tavily is not configured, falls back
    to pure GPT synthesis (user will see a note).
    """
    from app.services.research_service import TavilyResearchService
    import json

    tavily = TavilyResearchService()

    # Sanitize chapter text: remove control characters and null bytes that
    # break JSON serialization when sent to the OpenAI API.
    if chapter_text:
        import unicodedata
        chapter_text = "".join(
            ch for ch in chapter_text
            if ch in ("\n", "\r", "\t") or (not unicodedata.category(ch).startswith("C"))
        )

    # Build a short chapter summary hint for search planning
    chapter_hint = ""
    if chapter_text and chapter_text.strip():
        # Use first ~500 chars to give GPT a sense of the chapter topic
        chapter_hint = f"\nChapter context (the book chapter this is for):\n\"{chapter_text.strip()[:500]}...\"\n"

    # 1) Run 2-3 parallel searches to gather sources
    search_plan_prompt = f"""The user wants the agent to research the following topic for a technical book update:
Title: {title}
Research request: {research_prompt}
{chapter_hint}
Generate 3 focused web search queries (English, specific, recent) that will
find authoritative, up-to-date information RELEVANT TO THE CHAPTER TOPIC.
Do NOT generate generic queries — tailor them to the chapter's specific domain.
Return ONLY a JSON object with a single key "queries" containing a list of 3 query strings. No other text."""

    try:
        plan_resp = await client.chat.completions.create(
            model=settings.GPT_MODEL,
            messages=[
                {"role": "system", "content": "You are a research planner. Output strict JSON."},
                {"role": "user", "content": search_plan_prompt},
            ],
            temperature=0.2,
            max_tokens=400,
        )
        raw = plan_resp.choices[0].message.content.strip()
        if raw.startswith("```"):
            raw = raw.split("\n", 1)[1].rsplit("```", 1)[0]
        queries = json.loads(raw).get("queries", [])[:3]
    except Exception as exc:
        logger.warning("Custom research: query planning failed: %s", exc)
        queries = [research_prompt]

    # 2) Execute searches in parallel
    sources: list = []
    if tavily.is_configured and queries:
        tasks = [tavily.search(q, max_results=4) for q in queries]
        batches = await asyncio.gather(*tasks, return_exceptions=True)
        for batch in batches:
            if isinstance(batch, Exception):
                logger.warning("Custom research: search error: %s", batch)
                continue
            for r in batch or []:
                # ResearchResult objects — dump to dict
                try:
                    sources.append({
                        "title": getattr(r, "title", "") or "",
                        "url": getattr(r, "url", "") or "",
                        "excerpt": (getattr(r, "content", "") or "")[:600],
                        "source_type": getattr(r, "source_type", "") or "",
                    })
                except Exception:
                    pass
        # Deduplicate by URL
        seen = set()
        deduped = []
        for s in sources:
            u = s.get("url", "")
            if u and u not in seen:
                seen.add(u)
                deduped.append(s)
        sources = deduped[:10]

    # 3) Synthesize drafted content for the book
    if sources:
        evidence_block = "\n\n".join(
            f"[{i + 1}] {s['title']}\n    {s['url']}\n    {s['excerpt']}"
            for i, s in enumerate(sources)
        )
    else:
        evidence_block = "(No external sources found — rely on general knowledge. Flag uncertainty.)"

    # Build chapter context excerpt — truncate to ~3000 chars to fit within
    # token budget while giving GPT enough context about the chapter topic.
    chapter_context_block = ""
    if chapter_text and chapter_text.strip():
        excerpt = chapter_text.strip()[:3000]
        if len(chapter_text.strip()) > 3000:
            excerpt += "\n... [truncated]"
        chapter_context_block = f"""
CHAPTER CONTEXT (the chapter this content will be inserted into):
\"\"\"
{excerpt}
\"\"\"

CRITICAL: The drafted content MUST be directly relevant to this chapter's
subject matter. Do NOT write generic content. Tailor every paragraph to
the specific topic, domain, and systems discussed in the chapter above.
"""

    synth_prompt = f"""You are drafting new content for a technical book update.
The user has requested research on:
Title: {title}
Request: {research_prompt}
Allowed source types: {', '.join(allowed_sources)}
{chapter_context_block}
Web research gathered the following sources:
{evidence_block}

Write exactly 3 to 4 SHORT paragraphs (4-6 sentences each) of book-ready
prose that the editor can insert into a chapter. Keep total length under
800 words. Requirements:
- The content MUST be relevant to the chapter's specific topic — not generic.
- Match the measured, technical tone of a reference book (not a blog post).
- Include concrete current facts, numbers, and dates from the sources where relevant.
- Use inline citations as bracketed numbers [1], [2] matching the source list above.
- Do NOT include headings, bullet lists, or markdown — plain paragraphs only.
- CRITICAL: Do NOT include section numbers (1.3, 1.4, 2.1, etc.) or subsection headings at the start of paragraphs. The editor places the content at the right location. Start each paragraph directly with prose content.
- Do NOT preface with 'Here is...' — output the content directly.
- Preserve Greek symbols (Δ, α, β) and unit subscripts (v₁) as-is if used.

Output only the drafted paragraphs. No numbering, no headings, no titles.
IMPORTANT: Maximum 4 paragraphs. Be concise and focused."""

    synth_resp = await client.chat.completions.create(
        model=settings.GPT_MODEL,
        messages=[
            {"role": "system", "content": "You are a precise technical editor writing reference-book content. Keep your output concise — maximum 4 short paragraphs, under 800 words total."},
            {"role": "user", "content": synth_prompt},
        ],
        temperature=0.4,
        max_tokens=1000,
    )
    draft = (synth_resp.choices[0].message.content or "").strip()

    # Post-process: strip section numbers that GPT may insert despite instructions.
    # Matches lines starting with "1.3", "1.4 Title", "2.1.", etc. and removes
    # the number prefix so the paragraph starts with prose content.
    _sec_num_re = re.compile(r"^\d+\.\d+\.?\s*", re.MULTILINE)
    draft = _sec_num_re.sub("", draft)
    # Also strip any markdown heading markers (## 1.3 Title → Title)
    _md_heading_re = re.compile(r"^#{1,4}\s*(?:\d+\.\d+\.?\s*)?", re.MULTILINE)
    draft = _md_heading_re.sub("", draft)
    # Remove any leading/trailing blank lines that result from stripping
    draft = "\n\n".join(p.strip() for p in draft.split("\n\n") if p.strip())

    citations = [
        {
            "index": i + 1,
            "title": s.get("title", ""),
            "url": s.get("url", ""),
            "source_type": s.get("source_type", ""),
        }
        for i, s in enumerate(sources)
    ]
    return draft, citations


async def _plan_research_task(session_id: str, session: dict):
    """Background task: generate research plans for selected opportunities."""
    from openai import AsyncOpenAI
    import json

    db = get_database()
    session_repo = SessionRepository(db)

    try:
        client = AsyncOpenAI(api_key=(await _get_openai_key(db)))
        rules = session.get("rules", {}) or {}
        allowed_sources = rules.get("allowed_source_types", ["government", "academic", "news", "technical", "commercial"])

        selected = await session_repo.find_selected_opportunities(session_id)
        if not selected:
            await session_repo.update_session(session_id, {
                "status": SessionStatus.ERROR.value,
                "error_message": "No opportunities selected",
            })
            return

        # Fetch chapter text from the document (already stored in DB from
        # initial parsing) so research content is tailored to the chapter.
        chapter_text = ""
        try:
            doc_id = session.get("document_id", "")
            if doc_id:
                doc_repo = DocumentRepository(db)
                doc_with_text = await doc_repo.find_with_paragraphs(doc_id)
                if doc_with_text:
                    chapter_text = doc_with_text.get("text_content", "") or ""
                    logger.info("Custom research: loaded chapter text (%d chars) from document %s",
                                len(chapter_text), doc_id)
        except Exception as ctx_err:
            logger.warning("Custom research: could not load chapter text: %s", ctx_err)

        # Delete old plans for re-run
        await session_repo.delete_research_plans(session_id)

        # Helper to strip control characters that break OpenAI JSON serialization
        def _sanitize_for_api(text: str) -> str:
            if not text:
                return ""
            import unicodedata
            return "".join(
                ch for ch in text
                if ch in ("\n", "\r", "\t") or not unicodedata.category(ch).startswith("C")
            )

        plans = []
        for opp in selected:
            opp_id = opp.get("opportunity_id", opp.get("id", ""))

            # Custom issues: two sub-kinds
            if opp.get("is_custom"):
                custom_kind = opp.get("custom_kind", "find_replace")

                if custom_kind == "research":
                    # Run real web research + LLM synthesis, save draft to the
                    # opportunity so the user can review it in Step 7. Failures
                    # are non-fatal — the user will see an error state in the UI.
                    try:
                        draft_content, citations = await _run_custom_research(
                            client,
                            opp.get("research_prompt", ""),
                            opp.get("brief_reason", ""),
                            allowed_sources,
                            chapter_text=chapter_text,
                        )
                        await session_repo.update_opportunity(
                            session_id,
                            opp_id,
                            {
                                "research_result": draft_content,
                                "research_citations": citations,
                                "research_status": "ready",
                            },
                        )
                    except Exception as exc:
                        logger.warning(
                            "Custom research failed for opp %s: %s", opp_id, exc,
                        )
                        await session_repo.update_opportunity(
                            session_id,
                            opp_id,
                            {
                                "research_status": "error",
                                "research_error": str(exc)[:500],
                            },
                        )

                    plans.append({
                        "plan_id": str(uuid.uuid4())[:8],
                        "opportunity_id": opp_id,
                        "session_id": session_id,
                        "search_queries": [],
                        "target_sources": [],
                        "facts_to_verify": [f"Custom research: {opp.get('brief_reason', '')}"],
                        "approved": True,
                        "is_custom": True,
                        "custom_kind": "research",
                    })
                    continue

                # Legacy find_replace fast-path
                plans.append({
                    "plan_id": str(uuid.uuid4())[:8],
                    "opportunity_id": opp_id,
                    "session_id": session_id,
                    "search_queries": [],
                    "target_sources": [],
                    "facts_to_verify": [f"Custom replacement: '{opp.get('original_sentence', '')}' → '{opp.get('replace_with', '')}'"],
                    "approved": True,
                    "is_custom": True,
                    "custom_kind": "find_replace",
                    "replace_with": opp.get("replace_with", ""),
                })
                continue

            # Sanitize text that may contain control characters from the DOCX
            _claim = _sanitize_for_api(opp.get('original_sentence', ''))
            _issue = _sanitize_for_api(opp.get('brief_reason', ''))

            prompt = f"""Given this outdated claim from a technical document:
Claim: "{_claim}"
Issue: {_issue}
Allowed source types: {', '.join(allowed_sources)}

Generate a research plan to verify/update this claim. Return JSON with:
- "search_queries": list of 2-4 specific search queries
- "target_sources": list of specific domain names or organizations to search
- "facts_to_verify": list of specific factual claims to check

Return ONLY valid JSON. No other text."""

            response = await client.chat.completions.create(
                model=settings.GPT_MODEL,
                messages=[
                    {"role": "system", "content": "You are a research planner for technical document updates. Generate focused, efficient research plans."},
                    {"role": "user", "content": prompt},
                ],
                temperature=0.3,
                max_tokens=1000,
            )

            raw = response.choices[0].message.content.strip()
            if raw.startswith("```"):
                raw = raw.split("\n", 1)[1].rsplit("```", 1)[0]
            plan_data = json.loads(raw)

            plans.append({
                "plan_id": str(uuid.uuid4())[:8],
                "opportunity_id": opp_id,
                "session_id": session_id,
                "search_queries": plan_data.get("search_queries", []),
                "target_sources": plan_data.get("target_sources", []),
                "facts_to_verify": plan_data.get("facts_to_verify", []),
                "approved": False,
            })

        await session_repo.create_research_plans(plans)
        await session_repo.update_session(session_id, {
            "status": SessionStatus.RESEARCH_PLANNED.value,
        })

        logger.info("Research planning complete for session %s: %d plans", session_id, len(plans))

    except Exception as e:
        logger.error("Research planning failed for session %s: %s", session_id, e)
        error_str = str(e)
        if "insufficient_quota" in error_str or ("429" in error_str and "quota" in error_str.lower()):
            error_msg = (
                "OpenAI API quota exceeded. Your API key has no remaining credits. "
                "Please check your billing at platform.openai.com or update your API key in Admin → API Keys."
            )
        elif "invalid_api_key" in error_str or "401" in error_str:
            error_msg = (
                "OpenAI API key is invalid or expired. "
                "Please update your API key in Admin → API Keys."
            )
        else:
            error_msg = error_str
        await session_repo.update_session(session_id, {
            "status": SessionStatus.ERROR.value,
            "error_message": error_msg,
        })


@router.get("/{session_id}/research-plans")
async def get_research_plans(session_id: str, user=Depends(get_current_user_dep)):
    """Get all research plans for a session."""
    db = get_database()
    session_repo = SessionRepository(db)
    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    plans = await session_repo.find_research_plans(session_id)
    opportunities = await session_repo.find_selected_opportunities(session_id)

    # Build lookup for opportunity info
    opp_map = {}
    for o in opportunities:
        oid = o.get("opportunity_id", o.get("id", ""))
        opp_map[oid] = o

    # Enrich plans with opportunity data
    enriched = []
    for p in plans:
        p["opportunity"] = opp_map.get(p.get("opportunity_id", ""), {})
        enriched.append(p)

    return {
        "session_id": session_id,
        "plans": enriched,
        "total": len(enriched),
    }


@router.put("/{session_id}/research-plans/{plan_id}/approve")
async def approve_plan(
    session_id: str,
    plan_id: str,
    req: PlanApproveRequest,
    user=Depends(get_current_user_dep),
):
    """Approve or reject a research plan."""
    db = get_database()
    session_repo = SessionRepository(db)
    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")

    updated = await session_repo.approve_research_plan(plan_id, req.approved, req.rejected)
    return {"plan_id": plan_id, "approved": req.approved, "rejected": req.rejected, "updated": updated}


# ══════════════════════════════════════════════════════════════════════════════
# STAGE 6: EVIDENCE REVIEW (run research + review)
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/{session_id}/run-research")
async def run_research(
    session_id: str,
    background_tasks: BackgroundTasks,
    user=Depends(get_current_user_dep),
):
    """Run research for all approved plans using Tavily API."""
    db = get_database()
    session_repo = SessionRepository(db)
    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    approved_plans = await session_repo.find_approved_plans(session_id)
    if not approved_plans:
        raise HTTPException(400, "No approved research plans. Approve at least one plan first.")
    logger.info("Starting research for session %s: %d approved plans", session_id, len(approved_plans))
    for p in approved_plans[:3]:
        logger.info("  Plan %s queries: %s", p.get("plan_id", "?"), p.get("search_queries", []))

    # Set status to RESEARCHING immediately
    await session_repo.update_session(session_id, {"status": SessionStatus.RESEARCHING.value})

    background_tasks.add_task(_run_research_task, session_id, session, approved_plans)

    return {
        "session_id": session_id,
        "status": "researching",
        "plans_to_run": len(approved_plans),
        "message": "Research started. Poll session for updates.",
    }


async def _run_research_task(session_id: str, session: dict, approved_plans: list):
    """Background task: fast, focused research — capped plans, high parallelism, incremental saves."""
    import asyncio
    import httpx

    db = get_database()
    session_repo = SessionRepository(db)

    try:
        # Resolve Tavily API key: check MongoDB first, fall back to .env
        tavily_key = settings.TAVILY_API_KEY
        db_key_doc = await db.settings.find_one({"key": "tavily_api_key"})
        if db_key_doc and db_key_doc.get("value"):
            tavily_key = db_key_doc["value"]

        # Delete old evidence for re-run
        await session_repo.delete_evidence_items(session_id)

        # Cap at 20 plans max to keep research fast
        plans_to_run = approved_plans[:20]
        if len(approved_plans) > 20:
            logger.info("Capping research from %d to 20 plans for session %s",
                        len(approved_plans), session_id)

        # Pre-fetch all opportunities in one go to avoid N+1 DB queries
        opp_ids = list(set(p.get("opportunity_id", "") for p in plans_to_run if p.get("opportunity_id")))
        all_opps = {}
        for opp_id in opp_ids:
            opp = await session_repo.find_opportunity(opp_id)
            if opp:
                all_opps[opp_id] = opp

        async def research_plan(plan: dict, http_client: httpx.AsyncClient) -> list:
            """Research a single plan: 1 focused Tavily search, keep top 2 results."""
            plan_id = plan.get("plan_id", plan.get("id", ""))
            queries = plan.get("search_queries", [])
            target_sources = plan.get("target_sources", [])

            opp = all_opps.get(plan.get("opportunity_id", ""))
            original_sentence = opp.get("original_sentence", "") if opp else ""

            evidence = []
            seen_urls = set()

            # Use only the FIRST query (most targeted)
            query = queries[0] if queries else ""
            if not query:
                logger.warning("Plan %s has no search queries, skipping", plan_id)
                return evidence

            try:
                logger.info("Tavily search for plan %s: %s", plan_id, query[:80])
                payload = {
                    "api_key": tavily_key,
                    "query": query,
                    "max_results": 2,
                    "search_depth": "basic",
                }
                # Only add include_domains if we have actual domains
                if target_sources:
                    payload["include_domains"] = target_sources[:3]

                tavily_resp = await http_client.post(
                    "https://api.tavily.com/search",
                    json=payload,
                    timeout=15,
                )
                if tavily_resp.status_code != 200:
                    err_body = tavily_resp.text[:300]
                    logger.warning("Tavily returned %d for plan %s: %s", tavily_resp.status_code, plan_id, err_body)
                    # Detect API key expiry, auth errors, or usage limit exceeded — raise to abort all research
                    if tavily_resp.status_code in (401, 403, 432):
                        detail_msg = ""
                        try:
                            detail_msg = tavily_resp.json().get("detail", {}).get("error", "")
                        except Exception:
                            pass
                        if tavily_resp.status_code == 432:
                            raise RuntimeError(
                                "Tavily API usage limit exceeded. "
                                "Please upgrade your Tavily plan or update your API key."
                            )
                        raise RuntimeError(
                            f"Tavily API key expired or invalid (HTTP {tavily_resp.status_code}). "
                            "Please update your TAVILY_API_KEY in the server configuration."
                        )
                    return evidence
                results = tavily_resp.json().get("results", [])
                logger.info("Tavily returned %d results for plan %s", len(results), plan_id)

                for r in results[:2]:
                    url = r.get("url", "")
                    if url in seen_urls:
                        continue
                    seen_urls.add(url)

                    tavily_score = r.get("score", 0.5)
                    excerpt = r.get("content", "")[:800]
                    claim_words = set(original_sentence.lower().split())
                    excerpt_words = set(excerpt.lower().split())
                    overlap = len(claim_words & excerpt_words)
                    keyword_boost = min(overlap / max(len(claim_words), 1), 0.3)
                    relevance = round(min(tavily_score + keyword_boost, 1.0), 2)

                    evidence.append({
                        "evidence_id": str(uuid.uuid4())[:8],
                        "research_plan_id": plan_id,
                        "session_id": session_id,
                        "source_url": url,
                        "source_title": r.get("title", ""),
                        "excerpt": excerpt,
                        "relevance_score": relevance,
                        "accepted": None,
                    })

            except RuntimeError:
                raise  # Let API key / usage limit errors propagate to abort all research
            except Exception as e:
                logger.warning("Research query failed for plan %s: %s", plan_id, e)

            return evidence

        # Run in parallel batches of 10 — save after each batch for incremental progress
        total_evidence = 0
        async with httpx.AsyncClient() as http_client:
            batch_size = 10
            for batch_start in range(0, len(plans_to_run), batch_size):
                batch = plans_to_run[batch_start:batch_start + batch_size]
                results = await asyncio.gather(
                    *[research_plan(p, http_client) for p in batch]
                )
                batch_evidence = []
                for evidence_list in results:
                    batch_evidence.extend(evidence_list)

                # Save this batch immediately so the frontend sees progress
                if batch_evidence:
                    await session_repo.create_evidence_items(batch_evidence)
                    total_evidence += len(batch_evidence)
                    logger.info("Research batch saved: %d items (total %d) for session %s",
                                len(batch_evidence), total_evidence, session_id)

        # Mark research as done so frontend stops polling
        await session_repo.update_session(session_id, {"status": SessionStatus.RESEARCH_DONE.value})
        logger.info("Research complete for session %s: %d evidence items", session_id, total_evidence)

    except Exception as e:
        logger.error("Research failed for session %s: %s", session_id, e)
        error_str = str(e)
        if "Tavily" in error_str:
            error_msg = error_str  # Already user-friendly from RuntimeError above
        elif "insufficient_quota" in error_str or ("429" in error_str and "quota" in error_str.lower()):
            error_msg = (
                "OpenAI API quota exceeded. Your API key has no remaining credits. "
                "Please check your billing at platform.openai.com or update your API key in Admin → API Keys."
            )
        elif "invalid_api_key" in error_str or "401" in error_str:
            error_msg = (
                "OpenAI API key is invalid or expired. "
                "Please update your API key in Admin → API Keys."
            )
        else:
            error_msg = error_str
        await session_repo.update_session(session_id, {
            "status": SessionStatus.ERROR.value,
            "error_message": error_msg,
        })


@router.get("/{session_id}/evidence")
async def get_evidence(session_id: str, user=Depends(get_current_user_dep)):
    """Get all evidence items grouped by research plan."""
    db = get_database()
    session_repo = SessionRepository(db)
    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    evidence = await session_repo.find_evidence_items(session_id)
    plans = await session_repo.find_research_plans(session_id)
    opportunities = await session_repo.find_selected_opportunities(session_id)

    # Build lookups
    plan_map = {p.get("plan_id", p.get("id", "")): p for p in plans}
    opp_map = {o.get("opportunity_id", o.get("id", "")): o for o in opportunities}

    # Group evidence by opportunity
    grouped = {}
    for e in evidence:
        plan = plan_map.get(e.get("research_plan_id", ""), {})
        opp_id = plan.get("opportunity_id", "unknown")
        if opp_id not in grouped:
            grouped[opp_id] = {
                "opportunity": opp_map.get(opp_id, {}),
                "evidence": [],
            }
        grouped[opp_id]["evidence"].append(e)

    return {
        "session_id": session_id,
        "evidence_groups": list(grouped.values()),
        "total_evidence": len(evidence),
        "decided": sum(1 for e in evidence if e.get("accepted") is not None),
        "undecided": sum(1 for e in evidence if e.get("accepted") is None),
    }


@router.put("/{session_id}/evidence/{evidence_id}")
async def decide_evidence(
    session_id: str,
    evidence_id: str,
    req: EvidenceDecisionRequest,
    user=Depends(get_current_user_dep),
):
    """Accept or reject an evidence item."""
    db = get_database()
    session_repo = SessionRepository(db)
    updated = await session_repo.decide_evidence(evidence_id, req.accepted)

    # Check if all evidence has been decided
    evidence = await session_repo.find_evidence_items(session_id)
    all_decided = all(e.get("accepted") is not None for e in evidence)

    if all_decided and evidence:
        await session_repo.update_session(session_id, {
            "status": SessionStatus.EVIDENCE_REVIEWED.value,
        })

    return {
        "evidence_id": evidence_id,
        "accepted": req.accepted,
        "updated": updated,
        "all_decided": all_decided,
    }


# ══════════════════════════════════════════════════════════════════════════════
# STAGE 7: PATCH GENERATION & APPROVAL
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/{session_id}/generate-patches")
async def generate_patches(
    session_id: str,
    background_tasks: BackgroundTasks,
    user=Depends(get_current_user_dep),
):
    """Generate replacement patches using accepted evidence."""
    db = get_database()
    session_repo = SessionRepository(db)
    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    background_tasks.add_task(_generate_patches_task, session_id, session)

    return {
        "session_id": session_id,
        "status": "generating_patches",
        "message": "Patch generation started. Poll session for updates.",
    }


async def _generate_patches_task(session_id: str, session: dict):
    """Background task: generate sentence-level patches using AI."""
    from openai import AsyncOpenAI
    import json
    import unicodedata as _ucd

    def _sanitize(text: str) -> str:
        if not text:
            return ""
        return "".join(
            ch for ch in text
            if ch in ("\n", "\r", "\t") or not _ucd.category(ch).startswith("C")
        )

    db = get_database()
    session_repo = SessionRepository(db)

    try:
        client = AsyncOpenAI(api_key=(await _get_openai_key(db)))
        rules = session.get("rules", {}) or {}
        max_change_pct = rules.get("max_sentence_change_pct", 80.0)
        citation_style = rules.get("citation_style", "inline")
        voice_preservation = rules.get("voice_preservation", True)

        selected = await session_repo.find_selected_opportunities(session_id)

        # Delete old patches for re-run
        await session_repo.delete_patches(session_id)

        logger.info("Generating patches for %d selected opportunities (parallel)", len(selected))

        patch_api_errors = []  # Track API-level failures (quota, auth)

        # Note: when diagnostic re-runs, opportunities get new IDs but research plans
        # still reference old IDs. The fallback pool above handles this gracefully.

        # Pre-fetch ALL accepted evidence for the session as fallback
        # (needed when diagnostic re-run creates new opportunity IDs that don't match plan IDs)
        all_session_evidence = await session_repo.find_evidence_items(session_id)
        all_accepted_evidence = [e for e in all_session_evidence if e.get("accepted") is True]
        logger.info("Session has %d accepted evidence items (fallback pool)", len(all_accepted_evidence))

        async def _generate_one_patch(opp):
            """Generate a single patch for one opportunity — runs in parallel."""
            opp_id = opp.get("opportunity_id", opp.get("id", ""))

            # Custom issues: create patch directly without GPT — user already provided replacement
            if opp.get("is_custom"):
                custom_kind = opp.get("custom_kind", "find_replace")
                # Research-kind custom issues have their own Step 7 review flow
                # (user edits + chooses placement → approve-research endpoint
                # creates the final patch). Skip auto-patch-gen here.
                if custom_kind == "research":
                    return None
                return {
                    "patch_id": str(uuid.uuid4())[:8],
                    "opportunity_id": opp_id,
                    "session_id": session_id,
                    "original_sentence": opp.get("original_sentence", ""),
                    "revised_sentence": opp.get("replace_with", ""),
                    "citation": "Manual replacement by editor",
                    "rationale": opp.get("brief_reason", "Custom find-and-replace"),
                    "confidence": 1.0,
                    "change_pct": 100,
                    "status": PatchStatus.PENDING.value,
                    "editor_revision": None,
                    "reviewed_at": None,
                    "section_ref": opp.get("section_ref", ""),
                    "is_custom": True,
                    "scope": opp.get("scope", "whole_document"),
                }

            accepted_evidence = await session_repo.find_accepted_evidence_for_opportunity(opp_id)

            if not accepted_evidence and all_accepted_evidence:
                # Fallback: use all accepted session evidence when opp→plan chain is broken
                accepted_evidence = all_accepted_evidence
                logger.info("  Patch gen opp %s: using %d session-level accepted evidence (fallback)", opp_id, len(accepted_evidence))

            if not accepted_evidence:
                logger.warning("  Patch gen opp %s: SKIPPED — no accepted evidence", opp_id)
                return None  # skip opportunities with no accepted evidence

            evidence_text = "\n".join(
                f"- {e.get('source_title', '')}: {e.get('excerpt', '')[:300]}"
                for e in accepted_evidence
            )

            style_instructions = ""
            if voice_preservation:
                style_instructions = "IMPORTANT: Preserve the original author's writing style, tone, and voice. "
            style_instructions += f"Citation style: {citation_style}. "
            style_instructions += f"Maximum change: {max_change_pct}% of the original sentence. "

            prompt = f"""Given this original sentence from a technical document and the research evidence, write a replacement sentence.

Original: "{_sanitize(opp.get('original_sentence', ''))}"
Section: {_sanitize(opp.get('section_ref', ''))}
Issue: {_sanitize(opp.get('brief_reason', ''))}

Evidence:
{_sanitize(evidence_text)}

{style_instructions}

Return ONLY valid JSON with:
- "revised_sentence": the replacement text (do NOT add citation numbers like [N])
- "citation": short inline citation for the source used
- "reference_entry": ONE properly formatted bibliography entry for the PRIMARY source. Format MUST match this academic style: 'Last, First. Year. "Article Title." Organization/Publisher. URL' — ONLY use the ACTUAL source_url from the evidence above (do NOT invent or guess URLs). If no real URL exists in the evidence, omit the URL. Do NOT include raw excerpts, garbled text, or placeholder URLs. If no clear source, return empty string "".
- "rationale": brief explanation of why this change is needed
- "confidence": float 0.0-1.0
- "change_pct": estimated percentage of the sentence that changed

Return ONLY valid JSON. No other text."""

            try:
                response = await client.chat.completions.create(
                    model=settings.GPT_MODEL,
                    messages=[
                        {"role": "system", "content": "You are a precise technical editor. Write concise, accurate replacement sentences that update outdated information while preserving the document's style. CRITICAL: Never spell out Greek letters — always use the original symbols (Δ, α, β, θ, Σ, π, etc.). Never convert subscripts/superscripts to plain text (keep v₁ not v1, Δv not Delta v)."},
                        {"role": "user", "content": prompt},
                    ],
                    temperature=0.3,
                    max_tokens=1000,
                )

                raw = response.choices[0].message.content.strip()
                if raw.startswith("```"):
                    raw = raw.split("\n", 1)[1].rsplit("```", 1)[0]
                patch_data = json.loads(raw)

                # Enforce max change percentage
                change_pct = patch_data.get("change_pct", 50)
                if change_pct > max_change_pct:
                    return None  # skip patches that change too much

                return {
                    "patch_id": str(uuid.uuid4())[:8],
                    "opportunity_id": opp_id,
                    "session_id": session_id,
                    "original_sentence": opp.get("original_sentence", ""),
                    "revised_sentence": patch_data.get("revised_sentence", ""),
                    "citation": patch_data.get("citation", ""),
                    "reference_entry": patch_data.get("reference_entry", ""),
                    "rationale": patch_data.get("rationale", ""),
                    "confidence": patch_data.get("confidence", 0.5),
                    "change_pct": change_pct,
                    "status": PatchStatus.PENDING.value,
                    "editor_revision": None,
                    "reviewed_at": None,
                    "section_ref": opp.get("section_ref", ""),
                }
            except Exception as e:
                error_str = str(e)
                logger.warning("Patch generation failed for opp %s: %s", opp_id, e)
                if "insufficient_quota" in error_str or ("429" in error_str and "quota" in error_str.lower()):
                    patch_api_errors.append("quota")
                elif "invalid_api_key" in error_str or "401" in error_str:
                    patch_api_errors.append("auth")
                return None

        # Run all OpenAI calls in parallel
        results = await asyncio.gather(*[_generate_one_patch(opp) for opp in selected])
        patches = [p for p in results if p is not None]

        # Check if all non-custom opportunities failed due to API errors
        non_custom = [o for o in selected if not o.get("is_custom")]
        if patch_api_errors and len(patch_api_errors) >= len(non_custom) and not patches:
            if "quota" in patch_api_errors:
                error_msg = (
                    "OpenAI API quota exceeded. Your API key has no remaining credits. "
                    "Please check your billing at platform.openai.com or update your API key in Admin → API Keys."
                )
            else:
                error_msg = (
                    "OpenAI API key is invalid or expired. "
                    "Please update your API key in Admin → API Keys."
                )
            await session_repo.update_session(session_id, {
                "status": SessionStatus.ERROR.value,
                "error_message": error_msg,
            })
            logger.error("Patch generation aborted for session %s: %s", session_id, error_msg)
            return

        await session_repo.create_patches(patches)
        await session_repo.update_session(session_id, {
            "status": SessionStatus.PATCHES_GENERATED.value,
        })

        logger.info("Patch generation complete for session %s: %d patches", session_id, len(patches))

    except Exception as e:
        logger.error("Patch generation failed for session %s: %s", session_id, e)
        error_str = str(e)
        if "insufficient_quota" in error_str or ("429" in error_str and "quota" in error_str.lower()):
            error_msg = (
                "OpenAI API quota exceeded. Your API key has no remaining credits. "
                "Please check your billing at platform.openai.com or update your API key in Admin → API Keys."
            )
        elif "invalid_api_key" in error_str or "401" in error_str:
            error_msg = (
                "OpenAI API key is invalid or expired. "
                "Please update your API key in Admin → API Keys."
            )
        else:
            error_msg = error_str
        await session_repo.update_session(session_id, {
            "status": SessionStatus.ERROR.value,
            "error_message": error_msg,
        })


@router.get("/{session_id}/patches")
async def get_patches(session_id: str, user=Depends(get_current_user_dep)):
    """Get all patches for review."""
    db = get_database()
    session_repo = SessionRepository(db)
    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    patches = await session_repo.find_patches(session_id)
    return {
        "session_id": session_id,
        "patches": patches,
        "total": len(patches),
        "pending": sum(1 for p in patches if p.get("status") == "pending"),
        "approved": sum(1 for p in patches if p.get("status") in ("approved", "edited")),
        "rejected": sum(1 for p in patches if p.get("status") == "rejected"),
    }


@router.put("/{session_id}/patches/{patch_id}")
async def review_patch(
    session_id: str,
    patch_id: str,
    req: PatchReviewRequest,
    user=Depends(get_current_user_dep),
):
    """Approve, reject, or edit a patch."""
    db = get_database()
    session_repo = SessionRepository(db)

    fields = {"reviewed_at": datetime.utcnow()}
    if req.action == "approve":
        fields["status"] = PatchStatus.APPROVED.value
    elif req.action == "reject":
        fields["status"] = PatchStatus.REJECTED.value
    elif req.action == "edit":
        if not req.editor_revision:
            raise HTTPException(400, "editor_revision required for edit action")
        fields["status"] = PatchStatus.EDITED.value
        fields["editor_revision"] = req.editor_revision
    else:
        raise HTTPException(400, "Invalid action. Use 'approve', 'reject', or 'edit'.")

    updated = await session_repo.update_patch(patch_id, fields)
    logger.info("Patch %s reviewed: action=%s, status=%s, updated=%s", patch_id, req.action, fields["status"], updated)
    return {
        "patch_id": patch_id,
        "status": fields["status"],
        "updated": updated,
    }


@router.get("/{session_id}/section-paragraphs")
async def get_section_paragraphs(
    session_id: str,
    section_id: str,
    user=Depends(get_current_user_dep),
):
    """Get all paragraphs within a section (between this heading and the next)."""
    db = get_database()
    session_repo = SessionRepository(db)
    doc_repo = DocumentRepository(db)

    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    outline = session.get("outline", [])
    target = next((item for item in outline if item.get("id") == section_id), None)
    if not target:
        raise HTTPException(404, "Section not found in outline")

    # Find the next heading's paragraph_index to bound the section
    target_para_idx = target.get("paragraph_index", 0)
    next_para_idx = None
    found_target = False
    for item in outline:
        if found_target:
            next_para_idx = item.get("paragraph_index")
            break
        if item.get("id") == section_id:
            found_target = True

    # ── Load paragraph metadata from DB (no DOCX re-read) ──────────
    doc = await doc_repo.find_with_paragraphs(session["document_id"])
    if not doc:
        raise HTTPException(404, "Document not found")

    db_paragraphs = doc.get("paragraphs", [])
    paragraphs = []

    # Build a set of paragraph indices that ARE outline headings, to skip them
    outline_para_indices = {item.get("paragraph_index") for item in outline if item.get("paragraph_index") is not None}

    import re as _re

    # Build a lookup dict for fast access: idx -> paragraph data
    para_lookup = {p["idx"]: p for p in db_paragraphs}

    end_idx = next_para_idx if next_para_idx else (max(p["idx"] for p in db_paragraphs) + 1 if db_paragraphs else 0)

    for idx in range(target_para_idx + 1, end_idx):
        p = para_lookup.get(idx)
        if not p:
            continue
        text = p.get("text", "").strip()
        if not text:
            continue
        style_name = p.get("style", "")
        # Stop at next heading style
        if style_name.lower().startswith("heading"):
            break
        # Skip outline headings
        if idx in outline_para_indices:
            continue
        # Skip header/footer/caption/toc styles
        if any(skip in style_name.lower() for skip in ("header", "footer", "caption", "toc")):
            continue
        # Skip bare numbers / page numbers
        if _re.match(r'^\d+(\.\d+)*$', text):
            continue
        # Skip roman numeral page numbers
        if _re.match(r'^[ivxlcdm]+$', text.lower()):
            continue
        # Skip short numeric-only lines
        if len(text) < 15 and _re.match(r'^[\d\s\.\-\u2013\u2014]+$', text):
            continue
        # Skip horizontal lines / separators / dividers
        if _re.match(r'^[\s_\-\u2013\u2014=~*#\.]{2,}$', text):
            continue
        paragraphs.append({
            "index": idx,
            "text": text[:200] + ("..." if len(text) > 200 else ""),
            "full_text": text,
        })

    return {
        "session_id": session_id,
        "section_id": section_id,
        "section_text": target.get("text", ""),
        "paragraphs": paragraphs,
        "total": len(paragraphs),
    }


class AskAIRequest(BaseModel):
    prompt: str
    section_id: str  # outline item id
    section_text: str = ""  # the heading text of the selected section
    position: str = "replace"  # "before", "after", "replace"
    # Paragraph-level targeting (optional — if set, targets a specific paragraph)
    paragraph_index: int = -1  # -1 means section-level (legacy)
    paragraph_text: str = ""  # the paragraph text for matching


@router.post("/{session_id}/ask-ai")
async def ask_ai_patch(
    session_id: str,
    req: AskAIRequest,
    user=Depends(get_current_user_dep),
):
    """Use AI to generate a new patch based on user prompt, section, and position."""
    from openai import AsyncOpenAI
    import json as _json

    if req.position not in ("before", "after", "replace"):
        raise HTTPException(400, "position must be 'before', 'after', or 'replace'")

    db = get_database()
    session_repo = SessionRepository(db)
    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    rules = session.get("rules", {}) or {}
    citation_style = rules.get("citation_style", "inline")
    voice_preservation = rules.get("voice_preservation", True)

    style_instructions = ""
    if voice_preservation:
        style_instructions = "Preserve the original author's writing style, tone, and voice. "
    style_instructions += f"Citation style: {citation_style}. "

    # Determine if this is paragraph-level or section-level targeting
    is_paragraph_level = req.paragraph_index >= 0 and req.paragraph_text

    position_label = {
        "before": "INSERT NEW CONTENT BEFORE",
        "after": "INSERT NEW CONTENT AFTER",
        "replace": "REPLACE THE CONTENT OF",
    }[req.position]

    if is_paragraph_level:
        target_desc = f'the paragraph: "{req.paragraph_text[:200]}"'
        target_context = f'This paragraph is in section "{req.section_text}".'
    else:
        target_desc = f'the section "{req.section_text}"'
        target_context = ""

    prompt_text = f"""You are a technical editor. The user wants to {position_label.lower()} {target_desc} in a document.
{target_context}

User instruction: "{req.prompt}"
Position: {position_label} this {"paragraph" if is_paragraph_level else "section"}

{style_instructions}

IMPORTANT: Keep the generated content concise and accurate — maximum 4 to 5 lines. Be direct and to the point. Do not write lengthy paragraphs.

Return ONLY valid JSON with:
- "revised_sentence": the generated content (concise, 4-5 lines max)
- "citation": citation if applicable (empty string if none)
- "rationale": brief explanation of what was generated and why (1 sentence)
- "confidence": float 0.0-1.0

Return ONLY valid JSON. No other text."""

    try:
        client = AsyncOpenAI(api_key=(await _get_openai_key(db)))
        response = await client.chat.completions.create(
            model=settings.GPT_MODEL,
            messages=[
                {"role": "system", "content": "You are a precise technical editor. Generate accurate, well-cited content that matches the document's style. CRITICAL: Never spell out Greek letters — always use the original symbols (Δ, α, β, θ, Σ, π, etc.). Never convert subscripts/superscripts to plain text (keep v₁ not v1, Δv not Delta v)."},
                {"role": "user", "content": prompt_text},
            ],
            temperature=0.3,
            max_tokens=500,
        )

        raw = response.choices[0].message.content.strip()
        if raw.startswith("```"):
            raw = raw.split("\n", 1)[1].rsplit("```", 1)[0]
        ai_data = _json.loads(raw)

        position_labels_short = {"before": "Before", "after": "After", "replace": "Replace"}

        if is_paragraph_level:
            short_para = req.paragraph_text[:60] + ("..." if len(req.paragraph_text) > 60 else "")
            section_ref = f"{position_labels_short[req.position]} paragraph in section: {req.section_text}"
            original = f"[{position_labels_short[req.position]} paragraph: {short_para}]" if req.position != "replace" else req.paragraph_text
        else:
            section_ref = f"{position_labels_short[req.position]} section: {req.section_text}"
            original = f"[{position_labels_short[req.position]} section: {req.section_text}]" if req.position != "replace" else req.section_text

        ask_ai_meta = {
            "section_id": req.section_id,
            "section_text": req.section_text,
            "position": req.position,
            "user_prompt": req.prompt,
        }
        # Add paragraph-level targeting info
        if is_paragraph_level:
            ask_ai_meta["paragraph_index"] = req.paragraph_index
            ask_ai_meta["paragraph_text"] = req.paragraph_text

        patch = {
            "patch_id": str(uuid.uuid4())[:8],
            "opportunity_id": "ask-ai",
            "session_id": session_id,
            "original_sentence": original,
            "revised_sentence": ai_data.get("revised_sentence", ""),
            "citation": ai_data.get("citation", ""),
            "rationale": ai_data.get("rationale", ""),
            "confidence": ai_data.get("confidence", 0.5),
            "change_pct": 100.0,
            "status": PatchStatus.PENDING.value,
            "editor_revision": None,
            "reviewed_at": None,
            "section_ref": section_ref,
            "ask_ai_meta": ask_ai_meta,
        }

        await session_repo.create_patches([patch])
        # Remove MongoDB _id if present (not JSON-serializable)
        patch.pop("_id", None)
        logger.info("Ask AI patch created for session %s: %s (section=%s, position=%s)",
                     session_id, patch["patch_id"], req.section_text, req.position)

        return {
            "patch": patch,
            "message": "AI-generated patch created successfully.",
        }
    except Exception as e:
        logger.error("Ask AI failed for session %s: %s", session_id, str(e))
        raise HTTPException(500, f"AI generation failed: {str(e)}")


# ══════════════════════════════════════════════════════════════════════════════
# INSERT NEW FIGURE / TABLE (with auto-renumbering at apply time)
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/{session_id}/insert-media")
async def insert_media_patch(
    session_id: str,
    file: UploadFile = File(...),
    media_type: str = Form(...),  # "figure" or "table"
    section_id: str = Form(...),
    section_text: str = Form(""),
    position: str = Form("after"),  # "before", "after", "replace"
    paragraph_index: int = Form(-1),  # -1 = section-level
    paragraph_text: str = Form(""),
    caption: Optional[str] = Form(None),
    user=Depends(get_current_user_dep),
):
    """Accept a user-uploaded figure or table and queue it as a patch
    to be inserted into the document at apply-patches time. Downstream
    figure/table numbers will auto-renumber."""
    if media_type not in ("figure", "table"):
        raise HTTPException(400, "media_type must be 'figure' or 'table'")
    if position not in ("before", "after", "replace"):
        raise HTTPException(400, "position must be 'before', 'after', or 'replace'")

    db = get_database()
    session_repo = SessionRepository(db)
    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    # Persist the uploaded file
    upload_dir = os.path.join(settings.PROCESSING_DIR, session_id, "media_inserts")
    os.makedirs(upload_dir, exist_ok=True)
    safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", file.filename or "upload.bin")
    stored_name = f"{uuid.uuid4().hex[:8]}_{safe_name}"
    stored_path = os.path.join(upload_dir, stored_name)

    try:
        content = await file.read()
        with open(stored_path, "wb") as fh:
            fh.write(content)
    except Exception as e:
        logger.error("Failed to save uploaded media for session %s: %s", session_id, e)
        raise HTTPException(500, "Failed to save upload")

    file_size = len(content)

    pos_label = {"before": "Before", "after": "After", "replace": "Replace"}[position]
    label_type = "Figure" if media_type == "figure" else "Table"

    insert_media_meta = {
        "media_type": media_type,
        "file_path": stored_path,
        "file_name": safe_name,
        "file_size": file_size,
        "caption": (caption or "").strip(),
        "section_id": section_id,
        "section_text": section_text,
        "position": position,
        "paragraph_index": paragraph_index,
        "paragraph_text": paragraph_text,
    }

    patch = {
        "patch_id": str(uuid.uuid4())[:8],
        "opportunity_id": "insert-media",
        "session_id": session_id,
        "original_sentence": f"[Insert {label_type} {pos_label.lower()} section: {section_text}]",
        "revised_sentence": f"[New {label_type}: {safe_name}"
                            + (f" — {caption.strip()}" if caption and caption.strip() else "")
                            + "]",
        "citation": "",
        "rationale": f"User-uploaded {media_type} inserted {pos_label.lower()} "
                     f"section '{section_text}'. Downstream {media_type}s are "
                     f"auto-renumbered.",
        "confidence": 1.0,
        "change_pct": 100.0,
        # PENDING: the uploaded file appears as a preview card in the Figures /
        # Tables tab. The user must explicitly approve it before it is applied
        # to the clean document and tracked-changes exports.
        "status": PatchStatus.PENDING.value,
        "editor_revision": None,
        "reviewed_at": None,
        "section_ref": f"{pos_label}: {section_text}",
        "insert_media_meta": insert_media_meta,
    }

    await session_repo.create_patches([patch])
    patch.pop("_id", None)
    logger.info(
        "Insert-media patch created for session %s: %s (%s, section=%s, position=%s, file=%s)",
        session_id, patch["patch_id"], media_type, section_text, position, safe_name,
    )

    return {
        "patch": patch,
        "message": f"{label_type} uploaded. Approve the preview card to insert it.",
    }


@router.get("/{session_id}/insert-media-patches")
async def list_insert_media_patches(
    session_id: str,
    user=Depends(get_current_user_dep),
):
    """Return all insert-media patches (pending/approved/rejected) for a session
    so the frontend can render preview cards in the Figures/Tables tabs."""
    db = get_database()
    session_repo = SessionRepository(db)
    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    all_patches = await session_repo.find_patches(session_id)
    items = []
    for p in all_patches:
        meta = p.get("insert_media_meta")
        if not meta:
            continue
        items.append({
            "patch_id": p.get("patch_id"),
            "status": p.get("status"),
            "media_type": meta.get("media_type"),
            "file_name": meta.get("file_name"),
            "file_size": meta.get("file_size"),
            "caption": meta.get("caption") or "",
            "section_text": meta.get("section_text") or "",
            "position": meta.get("position") or "after",
            "paragraph_index": meta.get("paragraph_index", -1),
            "paragraph_text": meta.get("paragraph_text") or "",
            "preview_url": f"/sessions/{session_id}/insert-media-preview/{p.get('patch_id')}",
            "created_at": p.get("created_at"),
        })
    return {"patches": items}


@router.get("/{session_id}/insert-media-preview/{patch_id}")
async def insert_media_preview(
    session_id: str,
    patch_id: str,
    token: Optional[str] = None,
    request: Request = None,
):
    """Serve the uploaded figure/table file so the frontend can show a preview.
    Auth accepts either an Authorization header (fetch) or a ?token= query
    param so <img src=...> tags can load the file directly."""
    from fastapi.responses import FileResponse

    # Auth: query-param token for <img> tags, or Authorization header for fetch
    user = None
    if token:
        user = verify_download_token(token)
        if not user:
            # Fallback: some frontends pass the raw JWT directly as the token
            try:
                payload = decode_token(token)
                user = {"email": payload.get("sub"), "role": payload.get("role", "user")}
            except Exception:
                user = None
    else:
        auth_header = request.headers.get("authorization", "") if request else ""
        if auth_header.startswith("Bearer "):
            jwt_token = auth_header[7:]
            try:
                payload = decode_token(jwt_token)
                user = {"email": payload.get("sub"), "role": payload.get("role", "user")}
            except Exception:
                user = None
    if not user:
        raise HTTPException(401, "Authentication required")

    db = get_database()
    session_repo = SessionRepository(db)
    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    all_patches = await session_repo.find_patches(session_id)
    patch = next((p for p in all_patches if p.get("patch_id") == patch_id), None)
    if not patch or not patch.get("insert_media_meta"):
        raise HTTPException(404, "Insert-media patch not found")

    file_path = patch["insert_media_meta"].get("file_path")
    if not file_path or not os.path.exists(file_path):
        raise HTTPException(404, "Uploaded file not found on disk")

    return FileResponse(file_path, filename=patch["insert_media_meta"].get("file_name"))


# ══════════════════════════════════════════════════════════════════════════════
# STAGE 8: APPLY PATCHES
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/{session_id}/apply-patches")
async def apply_patches(
    session_id: str,
    user=Depends(get_current_user_dep),
):
    """Apply approved patches to a copy of the original document."""
    db = get_database()
    session_repo = SessionRepository(db)
    doc_repo = DocumentRepository(db)

    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    doc = await doc_repo.find_by_id(session["document_id"])
    if not doc:
        raise HTTPException(404, "Document not found")

    approved_patches = await session_repo.find_approved_patches(session_id)
    ask_ai_count = sum(1 for p in approved_patches if p.get("ask_ai_meta"))
    insert_media_count = sum(1 for p in approved_patches if p.get("insert_media_meta"))
    regular_count = len(approved_patches) - ask_ai_count - insert_media_count
    logger.info("Session %s: %d approved patches (%d regular, %d ask-ai, %d insert-media)",
                session_id, len(approved_patches), regular_count, ask_ai_count, insert_media_count)

    if not approved_patches:
        # No approved patches — skip application, just advance the session
        await session_repo.update_session(session_id, {
            "status": SessionStatus.EDITS_APPLIED.value,
        })
        return {
            "session_id": session_id,
            "status": SessionStatus.EDITS_APPLIED.value,
            "applied": 0,
            "skipped": 0,
            "message": "No approved patches to apply. Skipped to next stage.",
        }

    original_path = _resolve_file_path(doc.get("file_path", ""))
    applied = 0
    skipped = 0
    working_path = None

    if os.path.exists(original_path):
        # Copy original to working path and apply patches in DOCX
        working_dir = os.path.join(settings.PROCESSING_DIR, session_id)
        os.makedirs(working_dir, exist_ok=True)
        working_path = os.path.join(working_dir, f"working_{doc.get('original_filename', 'document.docx')}")
        shutil.copy2(original_path, working_path)

        try:
            from docx import Document as DocxDocument
            from copy import deepcopy
            docx_doc = DocxDocument(working_path)

            para_count_before = len(docx_doc.paragraphs)
            logger.info("DOCX loaded: %d paragraphs", para_count_before)

            for patch in approved_patches:
                final_text = patch.get("editor_revision") or patch.get("revised_sentence", "")
                ask_ai_meta = patch.get("ask_ai_meta")
                insert_media_meta = patch.get("insert_media_meta")
                research_insert_meta = patch.get("research_insert_meta")

                # --- Research-insert patches: drop drafted content at chosen
                # section/paragraph position ---
                if research_insert_meta:
                    try:
                        ok = _apply_research_insert_patch(
                            docx_doc, research_insert_meta,
                            outline=session.get("outline", []),
                        )
                        if ok:
                            applied += 1
                        else:
                            skipped += 1
                    except Exception as ri_err:
                        logger.error("Research-insert patch failed: %s", ri_err)
                        skipped += 1
                    continue

                # --- Insert Media patches: insert uploaded figure/table + renumber ---
                if insert_media_meta:
                    try:
                        inserted = _apply_insert_media_patch(docx_doc, insert_media_meta)
                        if inserted:
                            applied += 1
                        else:
                            skipped += 1
                    except Exception as im_err:
                        logger.error("Insert-media patch failed: %s", im_err)
                        skipped += 1
                    continue

                if not final_text:
                    skipped += 1
                    continue

                # --- Ask AI patches: insert/replace relative to a paragraph or section ---
                if ask_ai_meta:
                    section_text = ask_ai_meta.get("section_text", "").strip()
                    position = ask_ai_meta.get("position", "after")
                    para_text = ask_ai_meta.get("paragraph_text", "").strip()
                    para_idx = ask_ai_meta.get("paragraph_index", -1)

                    # Paragraph-level targeting: find the specific paragraph
                    if para_text and para_idx >= 0:
                        target_idx = None
                        # Strategy 1: exact paragraph index if text matches
                        if para_idx < len(docx_doc.paragraphs):
                            if para_text[:50] in docx_doc.paragraphs[para_idx].text:
                                target_idx = para_idx
                        # Strategy 2: search all paragraphs for matching text
                        if target_idx is None:
                            for search_idx, p in enumerate(docx_doc.paragraphs):
                                if para_text[:50] in p.text:
                                    target_idx = search_idx
                                    break

                        if target_idx is not None:
                            logger.info("Ask AI patch (paragraph-level): found paragraph at idx %d (position=%s)",
                                        target_idx, position)
                            if position == "before":
                                _insert_text_near_paragraph(docx_doc, target_idx, final_text, before=True)
                                applied += 1
                            elif position == "after":
                                _insert_text_near_paragraph(docx_doc, target_idx, final_text, before=False)
                                applied += 1
                            elif position == "replace":
                                _replace_text_in_paragraph(docx_doc.paragraphs[target_idx],
                                                          docx_doc.paragraphs[target_idx].text, final_text)
                                applied += 1
                            else:
                                skipped += 1
                            continue
                        else:
                            logger.warning("Ask AI patch: paragraph '%s...' not found, falling back to section",
                                           para_text[:40])

                    # Section-level targeting (fallback or legacy)
                    if not section_text:
                        skipped += 1
                        continue

                    heading_idx = _find_heading_in_docx(docx_doc, section_text)

                    if heading_idx is None:
                        logger.warning("Ask AI patch: heading '%s' not found in DOCX (patch_id=%s)",
                                       section_text, patch.get("patch_id"))
                        skipped += 1
                        continue

                    logger.info("Ask AI patch: found heading at paragraph %d for '%s' (position=%s)",
                                heading_idx, section_text, position)

                    if position == "before":
                        _insert_text_near_paragraph(docx_doc, heading_idx, final_text, before=True)
                        applied += 1
                    elif position == "after":
                        _insert_text_near_paragraph(docx_doc, heading_idx, final_text, before=False)
                        applied += 1
                    elif position == "replace":
                        _replace_section_content(docx_doc, heading_idx, final_text)
                        applied += 1
                    else:
                        skipped += 1
                    continue

                # --- Regular patches: find-and-replace ---
                original_text = patch.get("original_sentence", "")

                if not original_text:
                    skipped += 1
                    continue

                # Custom patches with whole_document scope: replace ALL occurrences
                replace_all = patch.get("is_custom") and patch.get("scope") == "whole_document"

                found = False
                replace_count = 0

                # Helper: normalize whitespace & common unicode variants for matching
                import re as _re_ws

                def _norm_ws(s: str) -> str:
                    # Normalize various dashes/hyphens to standard hyphen
                    s = s.replace('\u2013', '-').replace('\u2014', '-').replace('\u2012', '-')
                    # Normalize various quotes
                    s = s.replace('\u201c', '"').replace('\u201d', '"')
                    s = s.replace('\u2018', "'").replace('\u2019', "'")
                    # Collapse all whitespace (newlines, tabs, multiple spaces) to single space
                    return _re_ws.sub(r'\s+', ' ', s).strip()

                norm_orig = _norm_ws(original_text)

                for para in docx_doc.paragraphs:
                    para_text = para.text
                    # Try exact match first
                    if original_text in para_text:
                        _replace_text_in_paragraph(para, original_text, final_text, replace_all_in_para=replace_all)
                        found = True
                        replace_count += 1
                        if not replace_all:
                            applied += 1
                            break
                    # Fallback: normalized whitespace/unicode match
                    elif not found and norm_orig in _norm_ws(para_text):
                        # The original_text is in the paragraph but with different whitespace/unicode chars.
                        # Directly rewrite the runs since _replace_text_in_paragraph needs exact match.
                        if para.runs:
                            full_run_text = "".join(r.text for r in para.runs)
                            # If original is ~the whole paragraph, replace entirely
                            if len(norm_orig) >= len(_norm_ws(full_run_text)) * 0.7:
                                para.runs[0].text = final_text
                                for r in para.runs[1:]:
                                    r.text = ""
                            else:
                                # Substring case: replace the best-matching window
                                para.runs[0].text = full_run_text.replace(original_text, final_text, 1) if original_text in full_run_text else final_text
                                for r in para.runs[1:]:
                                    r.text = ""
                        found = True
                        replace_count += 1
                        if not replace_all:
                            applied += 1
                            logger.info("Patch applied via normalized match: '%s...'", original_text[:50])
                            break

                # Fallback 2: if still not found, try matching first N significant words
                if not found and not replace_all:
                    orig_words = _norm_ws(original_text).split()
                    if len(orig_words) >= 6:
                        # Match on first 6+ words as a prefix search
                        prefix_phrase = " ".join(orig_words[:6])
                        for para in docx_doc.paragraphs:
                            if prefix_phrase in _norm_ws(para.text):
                                # Direct run replacement — safer than _replace_text_in_paragraph
                                if para.runs:
                                    para.runs[0].text = final_text
                                    for r in para.runs[1:]:
                                        r.text = ""
                                found = True
                                applied += 1
                                logger.info("Patch applied via prefix match (%d words): '%s...'",
                                            len(orig_words[:6]), prefix_phrase[:50])
                                break

                # Fallback 3: difflib similarity match (>80% similar)
                if not found and not replace_all:
                    from difflib import SequenceMatcher
                    best_ratio = 0.0
                    best_para = None
                    for para in docx_doc.paragraphs:
                        pt = para.text.strip()
                        if not pt or len(pt) < 20:
                            continue
                        ratio = SequenceMatcher(None, _norm_ws(original_text), _norm_ws(pt)).ratio()
                        if ratio > best_ratio:
                            best_ratio = ratio
                            best_para = para
                    if best_ratio >= 0.80 and best_para and best_para.runs:
                        best_para.runs[0].text = final_text
                        for r in best_para.runs[1:]:
                            r.text = ""
                        found = True
                        applied += 1
                        logger.info("Patch applied via similarity match (%.0f%%): '%s...'",
                                    best_ratio * 100, original_text[:50])

                # Also search in table cells for whole-document custom patches
                if replace_all:
                    for table in docx_doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    if original_text in para.text:
                                        _replace_text_in_paragraph(para, original_text, final_text, replace_all_in_para=True)
                                        found = True
                                        replace_count += 1

                if replace_all and found:
                    applied += 1
                    logger.info("Custom replace-all: '%s' → '%s' in %d paragraphs",
                                original_text[:40], final_text[:40], replace_count)

                if not found:
                    logger.warning("Patch SKIPPED — text not found: '%s...' (patch_id=%s)",
                                   original_text[:60], patch.get("patch_id", "?"))
                    skipped += 1

            # --- Append new references to the References/Bibliography section ---
            refs_added = _append_references_to_docx(docx_doc, approved_patches, tracked=False)
            if refs_added:
                logger.info("Added %d new reference entries to References section", refs_added)

            para_count_after = len(docx_doc.paragraphs)
            logger.info("DOCX after patches: %d paragraphs (was %d, diff=%d)",
                        para_count_after, para_count_before, para_count_after - para_count_before)

            docx_doc.save(working_path)
            logger.info("DOCX saved to %s", working_path)

            # ── Apply approved media patches (figures, equations, tables) ────
            try:
                approved_media = await db.media_patches.find({
                    "session_id": session_id,
                    "status": "approved",
                }).to_list(length=500)

                if approved_media:
                    # Reload the saved doc for media edits
                    docx_doc = DocxDocument(working_path)
                    media_applied = 0

                    # ── Figures: replace images + update captions + add source URLs ──
                    approved_figures = [m for m in approved_media if m.get("type") == "figure"]
                    if approved_figures:
                        fig_count = await _apply_figure_replacements_to_docx(docx_doc, approved_figures)
                        media_applied += fig_count
                        cap_count = _update_figure_captions_in_docx(docx_doc, approved_figures)
                        if cap_count:
                            logger.info("Updated %d figure captions", cap_count)
                        src_count = _add_figure_source_urls(docx_doc, approved_figures)
                        if src_count:
                            logger.info("Added %d figure source URLs", src_count)

                    # ── Equations: apply OMML fixes ──
                    approved_equations = [m for m in approved_media if m.get("type") == "equation"]
                    if approved_equations:
                        eq_count = _apply_equation_replacements_to_docx(docx_doc, approved_equations)
                        media_applied += eq_count
                        logger.info("Applied %d equation replacements", eq_count)

                    # ── Tables: apply cell updates + update captions ──
                    approved_tables = [m for m in approved_media if m.get("type") == "table"]
                    if approved_tables:
                        tbl_count = _apply_table_updates_to_docx(docx_doc, approved_tables, highlight=False)
                        media_applied += tbl_count
                        logger.info("Applied %d table updates (clean, no highlight)", tbl_count)
                        tcap_count = _update_table_captions_in_docx(docx_doc, approved_tables)
                        if tcap_count:
                            logger.info("Updated %d table captions", tcap_count)

                    # Headers/footers: preserve from original
                    try:
                        original_doc_for_hf = DocxDocument(original_path)
                        _preserve_headers_footers(original_doc_for_hf, docx_doc)
                    except Exception as hf_err:
                        logger.debug("Headers/footers preservation skipped: %s", hf_err)

                    if media_applied:
                        docx_doc.save(working_path)
                        logger.info("Saved DOCX after %d media patches", media_applied)

            except Exception as media_err:
                logger.warning("Media patch application phase failed: %s", media_err)

            # Verify: reload and check paragraph count
            verify_doc = DocxDocument(working_path)
            verify_count = len(verify_doc.paragraphs)
            logger.info("VERIFY: Reloaded saved DOCX has %d paragraphs (expected %d)",
                        verify_count, para_count_after)
            # Log first 5 new paragraphs for debugging
            if verify_count > para_count_before:
                for i in range(min(5, verify_count)):
                    txt = verify_doc.paragraphs[i].text[:80] if verify_doc.paragraphs[i].text else "(empty)"
                    logger.info("  VERIFY para[%d]: %s", i, txt)

        except Exception as e:
            logger.error("Patch application failed for session %s: %s", session_id, e)
            raise HTTPException(500, f"Failed to apply patches: {str(e)}")
    else:
        # File not available locally — record patches as applied without DOCX modification
        logger.warning("Original file not found at %s — marking patches as applied without DOCX edit", original_path)
        applied = len(approved_patches)

    await session_repo.update_session(session_id, {
        "working_doc_path": working_path or "",
        "status": SessionStatus.EDITS_APPLIED.value,
    })

    return {
        "session_id": session_id,
        "status": SessionStatus.EDITS_APPLIED.value,
        "applied": applied,
        "skipped": skipped,
        "message": f"{applied} patches applied, {skipped} skipped.",
    }


def _replace_text_in_paragraph(paragraph, old_text: str, new_text: str, replace_all_in_para: bool = False):
    """Replace text in a paragraph while preserving run formatting.

    When replace_all_in_para=True, replaces ALL occurrences within this paragraph.
    Otherwise, only replaces the first occurrence.
    """
    # Join all runs to get full text
    full_text = "".join(run.text for run in paragraph.runs)
    if old_text not in full_text:
        return

    if replace_all_in_para:
        new_full = full_text.replace(old_text, new_text)
    else:
        new_full = full_text.replace(old_text, new_text, 1)

    # Put all text in first run, clear the rest
    if paragraph.runs:
        paragraph.runs[0].text = new_full
        for run in paragraph.runs[1:]:
            run.text = ""


def _normalize_text(text: str) -> str:
    """Normalize whitespace and case for fuzzy heading matching."""
    import re
    return re.sub(r'\s+', ' ', text.strip()).lower()


def _find_heading_in_docx(docx_doc, section_text: str):
    """Find a heading paragraph index using multiple matching strategies.

    Prioritises paragraphs whose style starts with 'Heading' so that body
    paragraphs that merely *reference* a section title are not picked up
    instead of the real heading.

    Handles field-code numbering: Word auto-numbered headings may have
    para.text = "Step 2: Title" while section_text = "1.4 Step 2: Title".
    We strip the leading number from BOTH sides for robust matching.
    """
    import re

    normalized_target = _normalize_text(section_text)
    # Strip leading section number (e.g. "1.4 " or "1.4.1 ") from the search text
    core_text = re.sub(r'^[\d.]+\s*', '', section_text).strip()
    core_lower = core_text.lower() if core_text and len(core_text) > 3 else ""
    target_words = set(normalized_target.split())
    # Also build word set from core text (without section number)
    core_words = set(_normalize_text(core_text).split()) if core_lower else set()

    # Log all heading-styled paragraphs for debugging
    heading_paras = []
    for idx, para in enumerate(docx_doc.paragraphs):
        style_name = para.style.name if para.style else ""
        if style_name.lower().startswith("heading"):
            heading_paras.append((idx, style_name, (para.text or "")[:80]))
    logger.info("_find_heading_in_docx: searching for %r", section_text[:80])
    logger.info("_find_heading_in_docx: document has %d heading-styled paragraphs", len(heading_paras))
    for idx, style, txt in heading_paras[:30]:
        logger.info("  heading[%d] style=%s text=%r", idx, style, txt)

    def _is_heading(para) -> bool:
        style_name = para.style.name if para.style else ""
        return style_name.lower().startswith("heading")

    # ---------- helpers for each strategy ----------
    def _exact_match(para) -> bool:
        return section_text in para.text

    def _normalized_match(para) -> bool:
        return normalized_target in _normalize_text(para.text)

    def _core_match(para) -> bool:
        """Match by stripping leading numbers from BOTH target and paragraph.
        Handles field-code numbering where para.text lacks section numbers."""
        if not core_lower:
            return False
        para_core = re.sub(r'^[\d.]+\s*', '', para.text).strip().lower()
        if not para_core:
            return False
        return core_lower in para_core or para_core in core_lower

    def _core_startswith_match(para) -> bool:
        """Looser match: paragraph text (with numbers stripped) starts with
        a significant prefix of the core target text, or vice-versa.
        Useful when the outline text has slight variations."""
        if not core_lower or len(core_lower) < 10:
            return False
        para_core = re.sub(r'^[\d.]+\s*', '', para.text).strip().lower()
        if not para_core or len(para_core) < 5:
            return False
        # Check if either starts with the other's first 30 chars
        prefix_len = min(30, len(core_lower), len(para_core))
        return (core_lower[:prefix_len] == para_core[:prefix_len])

    # ---------- Pass 1: only heading-styled paragraphs ----------
    for strategy_name, strategy in [
        ("exact", _exact_match),
        ("normalized", _normalized_match),
        ("core", _core_match),
        ("core_startswith", _core_startswith_match),
    ]:
        for idx, para in enumerate(docx_doc.paragraphs):
            if _is_heading(para) and strategy(para):
                logger.info("_find_heading_in_docx: MATCHED heading idx %d via '%s' (style=%s) text=%r",
                            idx, strategy_name, para.style.name if para.style else "?", para.text[:80])
                return idx

    # Word-overlap strategy — heading-styled only
    if len(core_words) >= 3:
        best_idx, best_overlap = None, 0
        for idx, para in enumerate(docx_doc.paragraphs):
            if not _is_heading(para):
                continue
            para_text = para.text.strip()
            if not para_text:
                continue
            # Compare using core words (without section number) for robustness
            para_words = set(_normalize_text(para_text).split())
            overlap = len(core_words & para_words)
            ratio = overlap / max(len(core_words), 1)
            if ratio > 0.6 and overlap > best_overlap:
                best_overlap = overlap
                best_idx = idx
        if best_idx is not None:
            logger.info("_find_heading_in_docx: word-overlap heading idx %d (overlap=%d/%d)",
                        best_idx, best_overlap, len(core_words))
            return best_idx

    # ---------- Pass 2: fallback — any paragraph ----------
    logger.warning("_find_heading_in_docx: no heading-styled match for %r, falling back to all paragraphs",
                   section_text[:60])
    for strategy_name, strategy in [
        ("exact", _exact_match),
        ("normalized", _normalized_match),
        ("core", _core_match),
        ("core_startswith", _core_startswith_match),
    ]:
        for idx, para in enumerate(docx_doc.paragraphs):
            if strategy(para):
                logger.info("_find_heading_in_docx: fallback matched idx %d via '%s' text=%r",
                            idx, strategy_name, para.text[:80])
                return idx

    if len(target_words) >= 3:
        best_idx, best_overlap = None, 0
        for idx, para in enumerate(docx_doc.paragraphs):
            para_text = para.text.strip()
            if not para_text:
                continue
            para_words = set(_normalize_text(para_text).split())
            overlap = len(target_words & para_words)
            ratio = overlap / max(len(target_words), 1)
            if ratio > 0.7 and overlap > best_overlap:
                best_overlap = overlap
                best_idx = idx
        if best_idx is not None:
            logger.info("_find_heading_in_docx: fallback word-overlap idx %d", best_idx)
            return best_idx

    return None


def _is_body_text_paragraph(para) -> bool:
    """Check if a paragraph looks like real body text (not a heading, caption,
    page header, or short label). Used to find good reference paragraphs
    for formatting.
    """
    style_name = para.style.name if para.style else ""
    style_lower = style_name.lower()
    # Skip headings, captions, TOC, header/footer styles
    if any(kw in style_lower for kw in ("heading", "caption", "toc", "header", "footer", "title")):
        return False
    text = para.text.strip()
    if not text:
        return False
    # Skip very short paragraphs (page numbers, labels, figure refs)
    if len(text) < 40:
        return False
    return True


def _find_reference_body_paragraph(docx_doc, heading_idx: int):
    """Find the nearest body (non-heading) paragraph to copy formatting from.

    Searches BEFORE the heading first (same visual section), then AFTER.
    Prefers paragraphs that are genuine body text (long enough, not captions
    or page headers). This ensures the inserted content matches the
    surrounding body text font, size, and spacing.
    """
    import re as _re

    # Prefer a real body paragraph BEFORE (same section context)
    for idx in range(heading_idx - 1, -1, -1):
        para = docx_doc.paragraphs[idx]
        style_name = para.style.name if para.style else ""
        # Stop if we hit a heading (don't cross sections) — check both
        # style-based headings and numbered headings like "1.2 Title"
        if style_name.lower().startswith("heading"):
            break
        if _is_body_text_paragraph(para):
            return para

    # Fallback: search AFTER the heading
    for idx in range(heading_idx + 1, min(heading_idx + 20, len(docx_doc.paragraphs))):
        para = docx_doc.paragraphs[idx]
        if _is_body_text_paragraph(para):
            return para

    # Wider fallback: any body text paragraph before (cross section boundaries)
    for idx in range(heading_idx - 1, -1, -1):
        para = docx_doc.paragraphs[idx]
        if _is_body_text_paragraph(para):
            return para

    # Last resort: any non-empty non-heading paragraph
    for idx in range(heading_idx - 1, -1, -1):
        para = docx_doc.paragraphs[idx]
        style_name = para.style.name if para.style else ""
        if not style_name.lower().startswith("heading") and para.text.strip():
            return para
    for idx in range(heading_idx + 1, len(docx_doc.paragraphs)):
        para = docx_doc.paragraphs[idx]
        style_name = para.style.name if para.style else ""
        if not style_name.lower().startswith("heading") and para.text.strip():
            return para
    return None


def _build_formatted_paragraph(docx_doc, ref_para, text_content: str):
    """Create a new paragraph that matches the formatting of ref_para using python-docx API."""
    from copy import deepcopy
    from docx.oxml.ns import qn

    # Use add_paragraph with the same style
    style_name = None
    if ref_para and ref_para.style:
        style_name = ref_para.style.name

    new_para = docx_doc.add_paragraph("", style=style_name)

    # Copy paragraph-level XML properties (alignment, spacing, indentation)
    if ref_para:
        source_pPr = ref_para._element.find(qn("w:pPr"))
        if source_pPr is not None:
            existing_pPr = new_para._element.find(qn("w:pPr"))
            if existing_pPr is not None:
                new_para._element.remove(existing_pPr)
            new_para._element.insert(0, deepcopy(source_pPr))

    # Add the text as a run, copying run formatting from reference
    run = new_para.add_run(text_content)

    if ref_para and ref_para.runs:
        src_run = ref_para.runs[0]
        # Copy font properties
        if src_run.font.name:
            run.font.name = src_run.font.name
        if src_run.font.size:
            run.font.size = src_run.font.size
        run.font.bold = src_run.font.bold
        run.font.italic = src_run.font.italic
        if src_run.font.color and src_run.font.color.rgb:
            run.font.color.rgb = src_run.font.color.rgb

        # Also copy raw rPr XML for properties not covered by python-docx API
        source_rPr = src_run._element.find(qn("w:rPr"))
        if source_rPr is not None:
            existing_rPr = run._element.find(qn("w:rPr"))
            if existing_rPr is not None:
                run._element.remove(existing_rPr)
            run._element.insert(0, deepcopy(source_rPr))

    # If font name/size still not set (run inherited from style), resolve from style hierarchy
    if ref_para and not run.font.name:
        try:
            style = ref_para.style
            checked = set()
            while style and style.name not in checked:
                checked.add(style.name)
                if style.font and style.font.name:
                    run.font.name = style.font.name
                    break
                style = style.base_style
        except Exception:
            pass
    if ref_para and not run.font.size:
        try:
            style = ref_para.style
            checked = set()
            while style and style.name not in checked:
                checked.add(style.name)
                if style.font and style.font.size:
                    run.font.size = style.font.size
                    break
                style = style.base_style
        except Exception:
            pass

    # Copy paragraph alignment
    if ref_para:
        new_para.alignment = ref_para.alignment

    return new_para


def _insert_text_near_paragraph(docx_doc, para_idx: int, text: str, before: bool = False):
    """Insert text content before or after a paragraph, matching document formatting."""
    target_para = docx_doc.paragraphs[para_idx]
    target_element = target_para._element

    ref_para = _find_reference_body_paragraph(docx_doc, para_idx)
    if ref_para:
        logger.info("  Using reference style: '%s', font: %s, size: %s",
                     ref_para.style.name if ref_para.style else "?",
                     ref_para.runs[0].font.name if ref_para.runs else "?",
                     ref_para.runs[0].font.size if ref_para.runs else "?")

    lines = [line.strip() for line in text.split("\n") if line.strip()]

    if before:
        # Each addprevious() inserts directly before target_element, so later
        # inserts land between earlier inserts and target. Iterating in
        # original order produces correct top-to-bottom reading order.
        for line in lines:
            new_para = _build_formatted_paragraph(docx_doc, ref_para, line)
            target_element.addprevious(new_para._element)
            logger.info("  Inserted paragraph BEFORE idx %d: '%s...'", para_idx, line[:60])
    else:
        insert_after = target_element
        for line in lines:
            new_para = _build_formatted_paragraph(docx_doc, ref_para, line)
            insert_after.addnext(new_para._element)
            logger.info("  Inserted paragraph AFTER idx %d: '%s...'", para_idx, line[:60])
            insert_after = new_para._element


def _replace_section_content(docx_doc, heading_idx: int, text: str):
    """Replace content between this heading and the next heading with new text."""
    heading_para = docx_doc.paragraphs[heading_idx]

    body_indices = []
    for idx in range(heading_idx + 1, len(docx_doc.paragraphs)):
        para = docx_doc.paragraphs[idx]
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            break
        body_indices.append(idx)

    ref_para = _find_reference_body_paragraph(docx_doc, heading_idx)

    for idx in reversed(body_indices):
        para = docx_doc.paragraphs[idx]
        para._element.getparent().remove(para._element)

    logger.info("  Removed %d body paragraphs for replace", len(body_indices))
    _insert_text_near_paragraph(docx_doc, heading_idx, text, before=False)


def _find_references_section_idx(docx_doc) -> int | None:
    """Find the paragraph index of the 'References' heading in the DOCX."""
    ref_keywords = ("references", "bibliography", "works cited")
    for idx, para in enumerate(docx_doc.paragraphs):
        text = para.text.strip().lower()
        style_name = (para.style.name if para.style else "").lower()
        if text in ref_keywords:
            # Prefer heading-styled paragraphs
            if "heading" in style_name or not text.endswith("."):
                return idx
    # Looser fallback: any paragraph starting with "References"
    for idx, para in enumerate(docx_doc.paragraphs):
        text = para.text.strip()
        if text.lower().startswith("references"):
            return idx
    return None


def _find_last_reference_number(docx_doc, ref_heading_idx: int) -> int:
    """Scan paragraphs after the References heading to find the highest [N] reference number."""
    import re as _re
    max_num = 0
    for idx in range(ref_heading_idx + 1, len(docx_doc.paragraphs)):
        para = docx_doc.paragraphs[idx]
        text = para.text.strip()
        if not text:
            continue
        # Stop if we hit another heading (next section)
        style_name = (para.style.name if para.style else "").lower()
        if "heading" in style_name and text:
            break
        # Look for [N] at start of line or standalone
        matches = _re.findall(r"\[(\d+)\]", text)
        for m in matches:
            num = int(m)
            if num > max_num:
                max_num = num
    return max_num


def _is_valid_reference_text(text: str) -> bool:
    """Check if a reference entry is valid readable text (not garbled binary data)."""
    if not text or len(text) < 10:
        return False
    # Count printable ASCII + common Unicode letters vs total chars
    import unicodedata
    readable = 0
    total = 0
    for ch in text:
        total += 1
        cat = unicodedata.category(ch)
        # Letters, numbers, punctuation, spaces, symbols
        if cat.startswith(("L", "N", "P", "Z", "S")):
            readable += 1
    if total == 0:
        return False
    ratio = readable / total
    # Reject if less than 80% readable characters
    if ratio < 0.80:
        return False
    # Reject if too many backslashes, pipes, or braces (binary artifacts)
    junk_chars = sum(1 for c in text if c in "\\|{}[]^~`@#$%&*_=<>")
    if len(text) > 0 and junk_chars / len(text) > 0.15:
        return False
    return True


def _normalize_reference_key(text: str) -> str:
    """Create an aggressively normalized key for deduplication.
    Strips URLs, years, punctuation, publisher names, and common words
    to match near-duplicate references that differ only in formatting."""
    import re as _re
    key = text.lower().strip()
    # Remove URLs
    key = _re.sub(r'https?://\S+', '', key)
    # Remove years
    key = _re.sub(r'\b(19|20)\d{2}\b', '', key)
    # Remove common publisher/source labels
    for noise in ("nasa", "url", "faa", "page", "retrieved", "accessed",
                   "available at", "available from", "space.com", "spacenews"):
        key = key.replace(noise, "")
    # Remove all punctuation
    key = _re.sub(r'[^\w\s]', '', key)
    # Collapse whitespace and sort words for order-independent matching
    words = sorted(set(key.split()))
    return " ".join(words)


def _split_reference_parts(entry: str) -> list:
    """Split a reference entry into segments: normal and italic (quoted titles).
    Returns list of (text, is_italic) tuples.
    E.g. 'Smith. 2024. "Article Title." NASA.' →
         [('Smith. 2024. ', False), ('Article Title.', True), (' NASA.', False)]
    """
    import re as _re
    parts = []
    # Match text within quotes (both straight and curly)
    pattern = _re.compile(r'["\u201c]([^"\u201d]+)["\u201d]')
    last_end = 0
    for m in pattern.finditer(entry):
        # Text before the quoted part
        if m.start() > last_end:
            parts.append((entry[last_end:m.start()], False))
        # The quoted text (italic)
        parts.append((m.group(1), True))
        last_end = m.end()
    # Remaining text after last quote
    if last_end < len(entry):
        parts.append((entry[last_end:], False))
    # If no quotes found, return the whole thing as non-italic
    if not parts:
        parts = [(entry, False)]
    return parts


def _build_reference_paragraph_xml(ref_para, entry: str, highlight_color: str):
    """Build a reference paragraph with italic titles and optional highlight (for tracked export)."""
    from copy import deepcopy
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    new_p = OxmlElement("w:p")

    # Copy paragraph properties from reference (includes hanging indent, spacing, etc.)
    pPr_copied = False
    if ref_para:
        source_pPr = ref_para._element.find(qn("w:pPr"))
        if source_pPr is not None:
            new_pPr = deepcopy(source_pPr)
            # Remove any style reference that might conflict — keep only formatting
            pStyle = new_pPr.find(qn("w:pStyle"))
            # Keep the style — it carries the hanging indent and spacing
            new_p.append(new_pPr)
            pPr_copied = True

    # If no pPr was copied, create a basic one with hanging indent
    if not pPr_copied:
        pPr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")       # 0.5 inch left margin
        ind.set(qn("w:hanging"), "720")    # 0.5 inch hanging indent
        pPr.append(ind)
        new_p.append(pPr)

    # Get base run properties from reference paragraph
    base_rPr = OxmlElement("w:rPr")
    if ref_para:
        source_runs = ref_para._element.findall(qn("w:r"))
        if source_runs:
            source_rPr = source_runs[0].find(qn("w:rPr"))
            if source_rPr is not None:
                base_rPr = deepcopy(source_rPr)

        # Resolve font from style if needed
        has_font = base_rPr.find(qn("w:rFonts")) is not None
        has_size = base_rPr.find(qn("w:sz")) is not None
        if not has_font or not has_size:
            try:
                style = ref_para.style
                checked = set()
                font_name, font_size = None, None
                while style and style.name not in checked:
                    checked.add(style.name)
                    if style.font:
                        if not font_name and style.font.name:
                            font_name = style.font.name
                        if not font_size and style.font.size:
                            font_size = style.font.size
                    if font_name and font_size:
                        break
                    style = style.base_style
                if font_name and not has_font:
                    rFonts = OxmlElement("w:rFonts")
                    rFonts.set(qn("w:ascii"), font_name)
                    rFonts.set(qn("w:hAnsi"), font_name)
                    rFonts.set(qn("w:cs"), font_name)
                    base_rPr.insert(0, rFonts)
                if font_size and not has_size:
                    half_pts = str(int(font_size / 6350))
                    sz = OxmlElement("w:sz")
                    sz.set(qn("w:val"), half_pts)
                    base_rPr.append(sz)
                    szCs = OxmlElement("w:szCs")
                    szCs.set(qn("w:val"), half_pts)
                    base_rPr.append(szCs)
            except Exception:
                pass

    # Split entry into normal and italic parts
    parts = _split_reference_parts(entry)

    for text, is_italic in parts:
        new_r = OxmlElement("w:r")
        rPr = deepcopy(base_rPr)

        # Add italic for quoted titles
        if is_italic:
            i_elem = OxmlElement("w:i")
            i_elem.set(qn("w:val"), "true")
            rPr.append(i_elem)
            iCs = OxmlElement("w:iCs")
            iCs.set(qn("w:val"), "true")
            rPr.append(iCs)

        # Add highlight if tracked
        if highlight_color:
            _add_highlight_to_rPr(rPr, highlight_color)

        new_r.append(rPr)

        new_t = OxmlElement("w:t")
        new_t.set(qn("xml:space"), "preserve")
        new_t.text = text
        new_r.append(new_t)

        new_p.append(new_r)

    return new_p


def _build_reference_paragraph_clean(docx_doc, ref_para, entry: str):
    """Build a reference paragraph with italic titles (for clean export)."""
    from copy import deepcopy
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    style_name = None
    if ref_para and ref_para.style:
        style_name = ref_para.style.name

    new_para = docx_doc.add_paragraph("", style=style_name)

    # Copy paragraph-level XML properties (hanging indent, spacing, etc.)
    pPr_copied = False
    if ref_para:
        source_pPr = ref_para._element.find(qn("w:pPr"))
        if source_pPr is not None:
            existing_pPr = new_para._element.find(qn("w:pPr"))
            if existing_pPr is not None:
                new_para._element.remove(existing_pPr)
            new_para._element.insert(0, deepcopy(source_pPr))
            pPr_copied = True

    # If no pPr was copied, add hanging indent to match reference style
    if not pPr_copied:
        pPr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "720")
        pPr.append(ind)
        new_para._element.insert(0, pPr)

    # Copy paragraph alignment
    if ref_para:
        new_para.alignment = ref_para.alignment

    # Split entry into parts and add runs
    parts = _split_reference_parts(entry)

    for text, is_italic in parts:
        run = new_para.add_run(text)

        # Copy font from reference
        if ref_para and ref_para.runs:
            src_run = ref_para.runs[0]
            if src_run.font.name:
                run.font.name = src_run.font.name
            if src_run.font.size:
                run.font.size = src_run.font.size

        # Resolve from style if needed
        if ref_para and not run.font.name:
            try:
                style = ref_para.style
                checked = set()
                while style and style.name not in checked:
                    checked.add(style.name)
                    if style.font and style.font.name:
                        run.font.name = style.font.name
                        break
                    style = style.base_style
            except Exception:
                pass
        if ref_para and not run.font.size:
            try:
                style = ref_para.style
                checked = set()
                while style and style.name not in checked:
                    checked.add(style.name)
                    if style.font and style.font.size:
                        run.font.size = style.font.size
                        break
                    style = style.base_style
            except Exception:
                pass

        if is_italic:
            run.font.italic = True

    return new_para._element


def _append_references_to_docx(docx_doc, patches: list, tracked: bool = False,
                                highlight_color: str = "cyan"):
    """Append new reference entries to the References/Bibliography section.

    Deduplicates entries (same source only appears once), filters out
    garbled/corrupted text, and does NOT add [N] numbering.

    In tracked mode, entries are highlighted to show they are new additions.
    Returns the number of references appended.
    """

    # Collect reference entries from patches
    raw_entries = []
    for patch in patches:
        raw = patch.get("reference_entry", "")
        if not raw:
            continue
        # Support multiple entries separated by ||
        for entry in raw.split("||"):
            entry = entry.strip()
            if entry:
                raw_entries.append(entry)

    if not raw_entries:
        return 0

    # Filter out garbled/corrupted entries
    valid_entries = [e for e in raw_entries if _is_valid_reference_text(e)]
    filtered_count = len(raw_entries) - len(valid_entries)
    if filtered_count:
        logger.info("REFERENCES: Filtered out %d garbled/corrupted entries", filtered_count)

    # Deduplicate: keep first occurrence of each unique reference
    seen_keys = set()
    unique_entries = []
    for entry in valid_entries:
        key = _normalize_reference_key(entry)
        if key not in seen_keys:
            seen_keys.add(key)
            unique_entries.append(entry)
        else:
            logger.debug("REFERENCES: Skipping duplicate: %s", entry[:60])

    dedup_count = len(valid_entries) - len(unique_entries)
    if dedup_count:
        logger.info("REFERENCES: Deduplicated %d duplicate entries", dedup_count)

    if not unique_entries:
        return 0

    # Also check against existing references in the document to avoid
    # duplicating references that are already present
    ref_heading_idx = _find_references_section_idx(docx_doc)
    if ref_heading_idx is None:
        logger.info("REFERENCES: No References section found — skipping reference append")
        return 0

    logger.info("REFERENCES: Found References heading at paragraph %d", ref_heading_idx)

    # Collect existing reference text for deduplication
    existing_ref_keys = set()
    last_ref_idx = ref_heading_idx
    for idx in range(ref_heading_idx + 1, len(docx_doc.paragraphs)):
        para = docx_doc.paragraphs[idx]
        style_name = (para.style.name if para.style else "").lower()
        if "heading" in style_name and para.text.strip():
            break
        text = para.text.strip()
        if text:
            last_ref_idx = idx
            existing_ref_keys.add(_normalize_reference_key(text))

    # Remove entries that already exist in the References section
    new_entries = []
    for entry in unique_entries:
        key = _normalize_reference_key(entry)
        if key not in existing_ref_keys:
            new_entries.append(entry)
        else:
            logger.debug("REFERENCES: Already in document: %s", entry[:60])

    if not new_entries:
        logger.info("REFERENCES: All entries already exist in document — nothing to add")
        return 0

    # Find an existing reference paragraph to copy formatting from.
    # Look for paragraphs that look like real references (contain a year + author pattern).
    import re as _ref_re
    ref_para = None
    for idx in range(ref_heading_idx + 1, len(docx_doc.paragraphs)):
        para = docx_doc.paragraphs[idx]
        style_name = (para.style.name if para.style else "").lower()
        if "heading" in style_name and para.text.strip():
            break
        text = para.text.strip()
        # Look for reference-like text: has a year (19xx or 20xx) and is long enough
        if text and len(text) > 30 and _ref_re.search(r'\b(19|20)\d{2}\b', text):
            ref_para = para
            logger.info("REFERENCES: Using formatting from existing ref: '%s'", text[:80])
            break
    if not ref_para:
        # Fallback to any non-empty paragraph in references section
        for idx in range(last_ref_idx, ref_heading_idx, -1):
            para = docx_doc.paragraphs[idx]
            if para.text.strip() and len(para.text.strip()) > 10:
                ref_para = para
                break
    if not ref_para:
        ref_para = _find_reference_body_paragraph(docx_doc, ref_heading_idx)

    # Insert each new reference entry (no [N] numbering)
    # Format with italic title to match existing reference style
    insert_after = docx_doc.paragraphs[last_ref_idx]._element
    count = 0

    for entry in new_entries:
        if tracked:
            new_p = _build_reference_paragraph_xml(ref_para, entry, highlight_color)
        else:
            new_p = _build_reference_paragraph_clean(docx_doc, ref_para, entry)

        insert_after.addnext(new_p)
        insert_after = new_p
        count += 1
        logger.info("REFERENCES: Appended '%s'", entry[:80])

    logger.info("REFERENCES: Appended %d new reference entries (from %d raw, %d valid, %d unique)",
                count, len(raw_entries), len(valid_entries), len(unique_entries))
    return count


def _scan_existing_media_numbers(docx_doc, media_type: str) -> list:
    """Scan the doc body text for existing 'Figure X-Y' / 'Table X-Y' numbers.
    Handles Unicode dashes (U+2010–U+2015, U+2212) and normalizes to ASCII '-'.
    Walks ALL paragraphs including those nested in tables via body.iter().
    """
    from docx.oxml.ns import qn as _qn
    DASHES = "\u002d\u2010\u2011\u2012\u2013\u2014\u2015\u2212"
    dash_class = f"[{DASHES}.]"
    if media_type == "figure":
        pat = re.compile(
            rf"\b(?:Figure|Fig\.?)\s+(\d+)({dash_class})(\d+)(?![0-9])", re.IGNORECASE,
        )
    else:
        pat = re.compile(
            rf"\b(?:Table|Tbl\.?)\s+(\d+)({dash_class})(\d+)(?![0-9])", re.IGNORECASE,
        )
    nums: set = set()
    P_TAG = _qn("w:p")
    body = docx_doc.element.body
    for p_el in body.iter(P_TAG):
        try:
            text = "".join(p_el.itertext())
        except Exception:
            continue
        for m in pat.finditer(text):
            norm_sep = "." if m.group(2) == "." else "-"
            nums.add(f"{m.group(1)}{norm_sep}{m.group(3)}")
    return sorted(nums, key=lambda n: [int(x) for x in re.split(r"[-.]", n)])


def _pick_next_media_number(existing: list, chapter_hint: str = "1") -> str:
    """(Legacy) Pick next sequential number — kept for backwards compat."""
    if not existing:
        return f"{chapter_hint}-1"
    last = existing[-1]
    m = re.match(r"^(\d+)([\-\.])(\d+)$", last)
    if not m:
        return f"{chapter_hint}-1"
    chap, sep, seq = m.group(1), m.group(2), int(m.group(3))
    return f"{chap}{sep}{seq + 1}"


def _determine_inserted_number(
    target_el, chapter_hint: str, media_type: str, position: str, docx_doc=None,
) -> tuple:
    """Determine the correct new figure/table number based on the insertion
    POSITION in the document body (not just "last + 1").

    Walks ALL <w:p> descendants (including ones nested inside tables, which
    is common in books where figure captions sit in layout table cells),
    top-to-bottom, and finds the highest caption number in the target
    chapter that appears BEFORE the insertion point. The new number is
    that max + 1. Also returns the full set of existing numbers (used
    later by RenumberingService.renumber_after_insertion to shift every
    downstream item + in-text reference).

    Args:
        target_el: lxml element the new media will be inserted next to.
        chapter_hint: chapter prefix from the section heading (e.g. "1").
        media_type: "figure" or "table".
        position: "before" or "after" — changes whether target_el's own
                  caption counts as "before" or "at/after" the new number.
        docx_doc: the Document object, used to fall back to document.body
                  if the target element's parent isn't directly body.
    """
    from docx.oxml.ns import qn as _qn
    P_TAG = _qn("w:p")

    if media_type == "figure":
        prefix = r"(?:Figure|Fig\.?)"
    else:
        prefix = r"(?:Table|Tbl\.?)"
    # Match ALL Unicode dash variants — SMAD books use non-breaking hyphen
    # (U+2011) inside caption numbers so they don't line-break.
    DASH_CLASS = "[\\-\u2010\u2011\u2012\u2013\u2014\u2015\u2212.]"
    # Caption = pattern at the very start of a paragraph's text
    cap_pat = re.compile(rf"^\s*{prefix}\s+(\d+)({DASH_CLASS})(\d+)(?![0-9])", re.IGNORECASE)
    # Any mention anywhere — used to collect the full "existing numbers" list
    any_pat = re.compile(rf"\b{prefix}\s+(\d+)({DASH_CLASS})(\d+)(?![0-9])", re.IGNORECASE)

    # Walk up to <w:body> so we can iterate every paragraph in the document
    # in reading order, including those nested inside tables.
    body = target_el.getparent()
    while body is not None and body.tag != _qn("w:body"):
        body = body.getparent()
    if body is None and docx_doc is not None:
        body = docx_doc.element.body
    if body is None:
        return f"{chapter_hint}-1", []

    paragraphs_in_order = list(body.iter(P_TAG))
    try:
        target_pos = paragraphs_in_order.index(target_el)
    except ValueError:
        target_pos = len(paragraphs_in_order)

    # "Before the new figure" cutoff:
    #   position=before → new lands AT target_pos → count i < target_pos
    #   position=after  → new lands AT target_pos+1 → count i <= target_pos
    cutoff = target_pos if position == "before" else target_pos + 1

    max_seq_before = 0
    sep = "-"
    all_numbers: set = set()

    # Also build a looser pattern that searches ANYWHERE in the paragraph,
    # not just at the start. Captions in SMAD books sometimes have invisible
    # field codes (SEQ, STYLEREF) prepended to the visible text, making
    # cap_pat.match() fail even though the paragraph IS a caption.
    # We use this as a fallback for positional counting.
    cap_search_pat = re.compile(
        rf"(?:^|\s)(?:{prefix})\s+(\d+)({DASH_CLASS})(\d+)(?![0-9])", re.IGNORECASE,
    )

    captions_before: list = []  # for logging
    captions_after: list = []

    for i, p_el in enumerate(paragraphs_in_order):
        try:
            text = "".join(p_el.itertext())
        except Exception:
            continue
        if not text:
            continue

        # Caption position check — find the FIRST figure/table mention that
        # looks like a caption. Try cap_pat.match() first (strict: starts at
        # paragraph beginning), fall back to cap_search_pat.search() for
        # paragraphs with field-code or layout-table prefixes.
        cm = cap_pat.match(text)
        if not cm:
            cm = cap_search_pat.search(text)
        if cm and cm.group(1) == chapter_hint:
            seq = int(cm.group(3))
            raw_sep = cm.group(2)
            norm_sep = "." if raw_sep == "." else "-"
            num_label = f"{cm.group(1)}{norm_sep}{cm.group(3)}"
            if i < cutoff:
                if seq > max_seq_before:
                    max_seq_before = seq
                    sep = norm_sep
                captions_before.append(num_label)
            else:
                captions_after.append(num_label)

        # Collect all mentions in the whole doc (captions + in-text refs).
        # Always store with ASCII "-" so number_map keys are consistent.
        for m in any_pat.finditer(text):
            raw_sep = m.group(2)
            norm_sep = "." if raw_sep == "." else "-"
            all_numbers.add(f"{m.group(1)}{norm_sep}{m.group(3)}")

    new_number = f"{chapter_hint}{sep}{max_seq_before + 1}"

    sorted_all = sorted(
        all_numbers, key=lambda n: [int(x) for x in re.split(r"[-.]", n)]
    )
    logger.info(
        "Insert-number: chapter=%s, captions_before=%s, captions_after=%s, "
        "max_seq_before=%d, new_number=%s, all_numbers=%s",
        chapter_hint, captions_before, captions_after,
        max_seq_before, new_number, sorted_all,
    )
    return new_number, sorted_all


def _find_heading_from_outline(outline: list, section_text: str, total_paras: int) -> int | None:
    """Look up a heading's paragraph_index from the stored outline data.

    The outline was built during extract-outline and stores the correct
    paragraph indices from the original DOCX. This is far more reliable than
    re-searching the document text (which can fail due to field codes, tabs,
    duplicate text, or non-standard styles).
    """
    if not outline or not section_text:
        return None

    section_lower = _normalize_text(section_text)
    import re
    section_core = re.sub(r'^[\d.]+\s*', '', section_text).strip().lower()

    for item in outline:
        item_text = item.get("text", "")
        item_idx = item.get("paragraph_index")
        if item_idx is None:
            continue
        # Must be a valid index
        if item_idx < 0 or item_idx >= total_paras:
            continue

        # Strategy 1: exact text match
        if item_text == section_text:
            logger.info("Research-insert: outline exact match -> idx %d text=%r", item_idx, item_text[:60])
            return item_idx

        # Strategy 2: normalized match
        item_lower = _normalize_text(item_text)
        if section_lower == item_lower or section_lower in item_lower or item_lower in section_lower:
            logger.info("Research-insert: outline normalized match -> idx %d text=%r", item_idx, item_text[:60])
            return item_idx

        # Strategy 3: core match (strip section numbers)
        item_core = re.sub(r'^[\d.]+\s*', '', item_text).strip().lower()
        if section_core and item_core and (section_core in item_core or item_core in section_core):
            logger.info("Research-insert: outline core match -> idx %d text=%r", item_idx, item_text[:60])
            return item_idx

    return None


def _apply_research_insert_patch(docx_doc, meta: dict, tracked: bool = False,
                                  outline: list | None = None) -> bool:
    """Insert user-approved drafted research content at the chosen location.

    meta fields:
        content:          final edited prose (may be multi-paragraph, \\n\\n separated)
        section_text:     heading text of target section
        paragraph_index:  -1 for section-level; otherwise index hint
        paragraph_text:   for fuzzy paragraph lookup (robustness across edits)
        position:         "before" | "after" | "replace"
        title, citations: metadata (not used at apply-time)

    outline:  the session's stored outline items (from extract-outline step).
              Used to find heading paragraph indices reliably, since the
              outline already resolved the correct indices during extraction.

    When `tracked=True`, inserted paragraphs are highlighted yellow so the
    editor can see them in the tracked-changes export (matches ask-ai style).
    """
    content = (meta.get("content") or "").strip()
    if not content:
        logger.warning("Research-insert: empty content")
        return False

    section_text = (meta.get("section_text") or "").strip()
    position = meta.get("position", "after")
    para_idx_hint = meta.get("paragraph_index", -1)
    para_text = (meta.get("paragraph_text") or "").strip()

    # Split into paragraphs on double newlines first, then single newlines as fallback.
    raw_paras = [p.strip() for p in content.split("\n\n") if p.strip()]
    if not raw_paras:
        raw_paras = [line.strip() for line in content.split("\n") if line.strip()]
    if not raw_paras:
        return False

    total_paras = len(docx_doc.paragraphs)

    # Determine target paragraph index
    target_idx = None
    if para_idx_hint is not None and para_idx_hint >= 0 and para_text:
        if para_idx_hint < total_paras:
            if para_text[:50] in (docx_doc.paragraphs[para_idx_hint].text or ""):
                target_idx = para_idx_hint
        if target_idx is None:
            for search_idx, p in enumerate(docx_doc.paragraphs):
                if para_text[:50] in (p.text or ""):
                    target_idx = search_idx
                    break

    if target_idx is None:
        # Section-level fallback: insert near the section heading
        if not section_text:
            logger.warning("Research-insert: no paragraph or section target")
            return False

        # Primary: use stored outline indices (most reliable)
        heading_idx = _find_heading_from_outline(outline or [], section_text, total_paras)

        # Fallback: search the DOCX paragraphs directly
        if heading_idx is None:
            logger.info("Research-insert: outline lookup failed, falling back to _find_heading_in_docx")
            heading_idx = _find_heading_in_docx(docx_doc, section_text)
        if heading_idx is None:
            logger.warning("Research-insert: heading '%s' not found", section_text[:60])
            return False
        target_idx = heading_idx
        # For section-level insert, "replace" means replace section content
        if position == "replace":
            joined = "\n\n".join(raw_paras)
            _replace_section_content(docx_doc, heading_idx, joined)
            logger.info(
                "Research-insert: replaced section '%s' content (%d paragraphs)",
                section_text[:40], len(raw_paras),
            )
            return True
        # Otherwise fall through to paragraph-level insert

    # Paragraph-level insertion
    if tracked:
        target_para = docx_doc.paragraphs[target_idx]
        target_el = target_para._element
        ref_para = _find_reference_body_paragraph(docx_doc, target_idx) or target_para

        if position == "replace":
            # Strike through target and insert highlighted paragraphs after
            for run in target_para.runs:
                rPr = run._element.find(_qn_w("rPr"))
                if rPr is None:
                    from docx.oxml import OxmlElement
                    rPr = OxmlElement("w:rPr")
                    run._element.insert(0, rPr)
                _add_strikethrough_to_rPr(rPr)
                _add_font_color_to_rPr(rPr, "FF0000")
            insert_after = target_el
            for para_text_line in raw_paras:
                new_p = _make_highlighted_paragraph_xml(ref_para, para_text_line, "yellow")
                insert_after.addnext(new_p)
                insert_after = new_p
        elif position == "before":
            # Each addprevious() inserts directly before target_el, so later
            # inserts land between earlier inserts and target. Iterating in
            # original order produces correct top-to-bottom reading order.
            for para_text_line in raw_paras:
                new_p = _make_highlighted_paragraph_xml(ref_para, para_text_line, "yellow")
                target_el.addprevious(new_p)
        else:  # after (default)
            insert_after = target_el
            for para_text_line in raw_paras:
                new_p = _make_highlighted_paragraph_xml(ref_para, para_text_line, "yellow")
                insert_after.addnext(new_p)
                insert_after = new_p
        logger.info(
            "Research-insert (tracked): %d paragraph(s) %s idx %d",
            len(raw_paras), position, target_idx,
        )
        return True

    # Clean export — plain inserted paragraphs, matching document formatting
    joined = "\n\n".join(raw_paras)
    if position == "replace":
        # Clear the target paragraph's text, then insert new content after it
        target_para = docx_doc.paragraphs[target_idx]
        target_el = target_para._element
        prev_sibling = target_el.getprevious()
        target_el.getparent().remove(target_el)
        # Re-index after removal: the previous sibling anchors insertion
        if prev_sibling is not None:
            # Walk paragraphs list again to locate the new index
            for i, p in enumerate(docx_doc.paragraphs):
                if p._element is prev_sibling:
                    _insert_text_near_paragraph(docx_doc, i, joined, before=False)
                    break
        else:
            _insert_text_near_paragraph(docx_doc, 0, joined, before=True)
    else:
        _insert_text_near_paragraph(docx_doc, target_idx, joined, before=(position == "before"))
    logger.info(
        "Research-insert (clean): %d paragraph(s) %s idx %d",
        len(raw_paras), position, target_idx,
    )
    return True


def _qn_w(tag: str) -> str:
    """Short helper for qn('w:tag') without importing at every use site."""
    from docx.oxml.ns import qn
    return qn(f"w:{tag}")


def _apply_insert_media_patch(docx_doc, meta: dict) -> bool:
    """Insert a user-uploaded figure/table into the DOCX at the chosen section,
    then auto-renumber downstream figures/tables and their in-text references.

    Returns True on success, False otherwise.
    """
    from docx.shared import Inches
    from app.services.renumbering_service import RenumberingService

    media_type = meta.get("media_type")
    file_path = meta.get("file_path", "")
    caption = (meta.get("caption") or "").strip()
    section_text = (meta.get("section_text") or "").strip()
    position = meta.get("position", "after")
    para_idx_hint = meta.get("paragraph_index", -1)
    para_text = (meta.get("paragraph_text") or "").strip()
    is_paragraph_level = para_idx_hint is not None and para_idx_hint >= 0 and para_text

    if not os.path.exists(file_path):
        logger.warning("Insert-media: file not found at %s", file_path)
        return False
    if not section_text and not is_paragraph_level:
        logger.warning("Insert-media: empty section_text")
        return False

    # Determine the target element (paragraph-level vs section-level)
    target_el = None

    if is_paragraph_level:
        # Find the target paragraph by index hint, fall back to text search
        target_idx = None
        if para_idx_hint < len(docx_doc.paragraphs):
            if para_text[:50] in (docx_doc.paragraphs[para_idx_hint].text or ""):
                target_idx = para_idx_hint
        if target_idx is None:
            for search_idx, p in enumerate(docx_doc.paragraphs):
                if para_text[:50] in (p.text or ""):
                    target_idx = search_idx
                    break
        if target_idx is None:
            logger.warning(
                "Insert-media: paragraph '%s...' not found, falling back to section",
                para_text[:40],
            )
            is_paragraph_level = False
        else:
            logger.info(
                "Insert-media: found target paragraph at idx %d (position=%s)",
                target_idx, position,
            )
            target_el = docx_doc.paragraphs[target_idx]._element
            if position == "replace":
                # Remove target paragraph — insertion will land where it was
                parent = target_el.getparent()
                # Insert a placeholder empty p to hold the position, then remove later
                # Simpler: keep a reference to the previous sibling and insert after it
                prev_sibling = target_el.getprevious()
                parent.remove(target_el)
                target_el = prev_sibling  # insert after prev
                position = "after"
                if target_el is None:
                    # Edge case: target was the first element — insert at top of body
                    target_el = parent[0] if len(parent) > 0 else None
                    position = "before"

    if target_el is None:
        # Section-level targeting (fallback or explicit)
        heading_idx = _find_heading_in_docx(docx_doc, section_text)
        if heading_idx is None:
            logger.warning("Insert-media: heading '%s' not found", section_text)
            return False

        target_el = docx_doc.paragraphs[heading_idx]._element

        if position == "replace":
            # Remove body paragraphs under this heading, then insert after heading
            body_indices = []
            for idx in range(heading_idx + 1, len(docx_doc.paragraphs)):
                p = docx_doc.paragraphs[idx]
                sname = p.style.name if p.style else ""
                if sname.startswith("Heading"):
                    break
                body_indices.append(idx)
            for idx in reversed(body_indices):
                p = docx_doc.paragraphs[idx]
                p._element.getparent().remove(p._element)
            position = "after"

    # Determine the correct new number based on where in the body the
    # figure/table will actually land — NOT just "last existing + 1".
    # Chapter hint comes from the selected section heading:
    #   "1.3 Step 1: ..."  → "1"
    #   "Chapter 3"        → "3"
    chap_match = re.match(r"^\s*(?:Chapter\s+)?(\d+)\b", section_text)
    chapter_hint = chap_match.group(1) if chap_match else "1"
    new_number, existing_numbers = _determine_inserted_number(
        target_el, chapter_hint, media_type, position, docx_doc=docx_doc,
    )
    logger.info(
        "Insert-media: positional new number for %s = %s (chapter=%s, existing=%s)",
        media_type, new_number, chapter_hint, existing_numbers,
    )

    # Build new paragraphs: image/table + caption
    label = "Figure" if media_type == "figure" else "Table"
    full_caption = f"{label} {new_number}" + (f". {caption}" if caption else "")

    # Insert the media paragraph
    media_para = docx_doc.add_paragraph()
    try:
        if media_type == "figure" or file_path.lower().endswith(
            (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".tif", ".tiff")
        ):
            run = media_para.add_run()
            run.add_picture(file_path, width=Inches(5.5))
        else:
            # Table uploaded as docx/xlsx/csv — embed as plain text fallback
            media_para.add_run(f"[Embedded {label} file: {os.path.basename(file_path)}]")
    except Exception as pic_err:
        logger.error("Insert-media: failed to embed file: %s", pic_err)
        # Leave the empty paragraph; caption will still show the number
    media_el = media_para._element

    # Caption paragraph
    caption_para = docx_doc.add_paragraph()
    cap_run = caption_para.add_run(full_caption)
    cap_run.bold = True
    caption_el = caption_para._element

    # Move the two new elements from the end of doc to the target position
    parent = media_el.getparent()
    parent.remove(media_el)
    parent = caption_el.getparent()
    parent.remove(caption_el)

    if position == "before":
        target_el.addprevious(media_el)
        media_el.addnext(caption_el)
    else:
        # "after" (replace has been normalized to after)
        target_el.addnext(media_el)
        media_el.addnext(caption_el)

    logger.info(
        "Insert-media: inserted %s %s (%s) %s %s '%s'",
        media_type, new_number, os.path.basename(file_path), position,
        "paragraph in" if is_paragraph_level else "section",
        section_text,
    )

    # Auto-renumber downstream items and their in-text references
    if existing_numbers:
        full_text = "\n".join(p.text or "" for p in docx_doc.paragraphs)
        _, number_map = RenumberingService.renumber_after_insertion(
            full_text, existing_numbers, new_number, ref_type=media_type,
        )
        if number_map:
            # Walk the body paragraphs and rewrite runs for shifted numbers.
            # IMPORTANT: skip the newly-inserted caption so it doesn't get
            # swept up by the regex rename (which would corrupt its number).
            _apply_number_map_to_paragraphs(
                docx_doc, number_map, media_type, skip_elements={caption_el},
            )
            logger.info("Insert-media: renumbered %d downstream %s(s): %s",
                        len(number_map), media_type, number_map)
        else:
            logger.info("Insert-media: no downstream renumbering needed (number_map empty)")
    else:
        logger.info("Insert-media: no existing numbers found — skipping renumber")

    # Post-renumber verification: scan all captions after insertion and check
    # for duplicates or gaps (helps diagnose renumbering issues).
    post_numbers = _scan_existing_media_numbers(docx_doc, media_type)
    logger.info("Insert-media: post-renumber %s numbers in doc: %s", media_type, post_numbers)
    # Check for duplicates
    seen = set()
    for n in post_numbers:
        if n in seen:
            logger.warning("Insert-media: DUPLICATE %s number found after renumber: %s", media_type, n)
        seen.add(n)

    return True


def _apply_number_map_to_paragraphs(
    docx_doc, number_map: dict, media_type: str, skip_elements: set = None,
):
    """Rewrite in-text 'Figure X-Y' / 'Table X-Y' references using number_map.
    Order: renumber highest first to avoid collisions (3-3 -> 3-4 before 3-2 -> 3-3).

    Iterates ALL <w:p> descendants in the document body, including
    paragraphs nested inside tables (captions are often placed in layout
    table cells in books). python-docx's `Document.paragraphs` only
    returns top-level paragraphs, so we walk the XML tree directly.

    skip_elements: set of XML elements to exclude from rewrite (e.g. a newly
    inserted caption that should keep its fresh number).
    """
    from docx.oxml.ns import qn as _qn

    skip_elements = skip_elements or set()
    if media_type == "figure":
        prefix_pat = r"(?:Figure|Fig\.?)"
    else:
        prefix_pat = r"(?:Table|Tbl\.?)"

    P_TAG = _qn("w:p")
    T_TAG = _qn("w:t")

    # SMAD-style books use non-breaking hyphen (U+2011) or other Unicode
    # dashes inside caption numbers ("Fig. 1‑4") to prevent line-breaking.
    # We must match ALL dash variants, not just ASCII "-".
    DASH_CHARS = "\u002d\u2010\u2011\u2012\u2013\u2014\u2015\u2212"
    dash_class = f"[{DASH_CHARS}]"
    dash_split_re = re.compile(f"[{DASH_CHARS}]")

    # Sort: old_nums descending by sequence so renames don't collide
    def sort_key(item):
        m = re.match(r"^(\d+)[\-\.](\d+)$", item[0])
        return (int(m.group(1)), int(m.group(2))) if m else (0, 0)

    ordered = sorted(number_map.items(), key=sort_key, reverse=True)

    body = docx_doc.element.body
    all_paragraphs = list(body.iter(P_TAG))

    total_hits = 0
    for old_num, new_num in ordered:
        old_parts = dash_split_re.split(old_num)
        new_parts = dash_split_re.split(new_num)
        if len(old_parts) != 2 or len(new_parts) != 2:
            logger.warning("Renumber %s: skipping malformed %s → %s", media_type, old_num, new_num)
            continue
        # Build pattern that matches the old number with ANY unicode dash
        # between the chapter and sequence parts, and captures that dash so
        # replacement can preserve it (keeping non-breaking behaviour).
        pattern = re.compile(
            rf"\b({prefix_pat}\s+){re.escape(old_parts[0])}({dash_class}){re.escape(old_parts[1])}(?![0-9])",
            re.IGNORECASE,
        )
        repl = rf"\g<1>{new_parts[0]}\g<2>{new_parts[1]}"
        hits = 0
        for p_el in all_paragraphs:
            if p_el in skip_elements:
                continue
            # Collect ALL <w:t> descendants of this paragraph — this reaches
            # runs nested inside <w:hyperlink>, <w:smartTag>, <w:fldSimple>,
            # <mc:AlternateContent>, etc., which `findall('w:r')` would miss.
            # In SMAD-style books, figure captions often use REF/SEQ fields
            # or hyperlinks, so nested traversal is required.
            t_refs = list(p_el.iter(T_TAG))
            if not t_refs:
                continue
            full = "".join(t.text or "" for t in t_refs)
            if not pattern.search(full):
                continue
            new_full = pattern.sub(repl, full)
            if new_full == full:
                continue
            # Place the rewritten text in the first <w:t>, blank the rest.
            # This may merge differently-styled runs but keeps text correct.
            t_refs[0].text = new_full
            for t in t_refs[1:]:
                t.text = ""
            hits += 1
        total_hits += hits
        logger.info(
            "Renumber %s: %s → %s applied to %d paragraph(s)",
            media_type, old_num, new_num, hits,
        )
    logger.info(
        "Renumber %s: total %d paragraph rewrites across %d number map entries",
        media_type, total_hits, len(ordered),
    )


# ── Visual Diff / Tracked Changes Helpers ─────────────────────────────────

def _add_highlight_to_rPr(rPr_element, color: str):
    """Add highlight color to a run properties element.
    Colors: yellow, green, red, cyan, magenta, blue, darkYellow, etc.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Remove existing highlight
    existing = rPr_element.find(qn("w:highlight"))
    if existing is not None:
        rPr_element.remove(existing)

    highlight = OxmlElement("w:highlight")
    highlight.set(qn("w:val"), color)
    rPr_element.append(highlight)


def _add_strikethrough_to_rPr(rPr_element):
    """Add strikethrough to a run properties element."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    strike = OxmlElement("w:strike")
    strike.set(qn("w:val"), "true")
    rPr_element.append(strike)


def _add_font_color_to_rPr(rPr_element, hex_color: str):
    """Set font color on a run properties element. hex_color like 'FF0000' for red."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    existing = rPr_element.find(qn("w:color"))
    if existing is not None:
        rPr_element.remove(existing)

    color_elem = OxmlElement("w:color")
    color_elem.set(qn("w:val"), hex_color)
    rPr_element.append(color_elem)


def _is_junk_equation_text(latex: str, omml_xml: str = "") -> bool:
    """Return True if the equation text is not a real mathematical equation.

    Catches: symbol palettes, alphabet listings, Word UI elements, standalone
    LaTeX commands, placeholders, trivial single-char items, and tabular tables.
    """
    import re as _re

    stripped = latex.strip()
    if not stripped:
        return True

    # 0. Single character or very short non-equation
    if len(stripped) <= 1:
        return True

    # 1. Single LaTeX command with no arguments (e.g., just \partial, \alpha)
    if _re.match(r'^\\[a-zA-Z]+$', stripped):
        return True

    # 2. Placeholder text
    if stripped.lower() in ('type equation here', 'type equation here.'):
        return True

    # 3. LaTeX tabular environment (Word UI table, not equation)
    if '\\begin{tabular}' in stripped:
        return True

    # 4. Multiple \hline (table layout)
    if stripped.count('\\hline') >= 2:
        return True

    # 5. Word UI keywords
    lower = stripped.lower()
    ui_keywords = [
        'file', 'home', 'insert', 'design', 'layout', 'mailings',
        'review', 'view', 'add-ins', 'cover page', 'page break',
        'blank page', 'gallery', 'category', 'description', 'save in',
        'autotext', 'building block', 'search document',
        'insert content only', 'online pictures',
    ]
    if any(kw in lower for kw in ui_keywords):
        return True

    # 6. Alphabet listing / symbol palette (A,B,C,...Z or a,b,c,...z)
    clean = stripped.replace(',', '').replace(' ', '')
    if len(clean) >= 20:
        alpha_only = ''.join(c for c in clean if c.isalpha())
        if len(alpha_only) >= 20:
            unique = set(alpha_only.lower())
            if len(unique) >= 18:
                return True

    # 7. Comma-separated single characters (symbol reference)
    parts = [p.strip() for p in stripped.split(',')]
    if len(parts) >= 15 and all(len(p) <= 2 for p in parts):
        return True

    return False


async def _apply_figure_replacements_to_docx(docx_doc, approved_figures: list):
    """Download and replace images in a DOCX document for approved figure patches.

    Returns the number of figures successfully replaced.
    Works with any python-docx Document object (clean or tracked).
    """
    import httpx as _httpx
    from io import BytesIO
    from PIL import Image as PILImage

    _BROWSER_HEADERS = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/124.0.0.0 Safari/537.36"
        ),
        "Accept": "image/avif,image/webp,image/apng,image/svg+xml,image/*,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
    }

    replaced = 0
    for mp in approved_figures:
        sel = mp.get("selected_replacement")
        if not sel:
            continue

        r_embed = mp.get("r_embed")
        if not r_embed or r_embed not in docx_doc.part.rels:
            logger.warning("Figure replacement: r_embed=%s not found in DOCX rels", r_embed)
            continue

        # ── Handle user-uploaded local images ──
        if sel.get("is_user_upload") and sel.get("local_path"):
            local_path = sel["local_path"]
            if os.path.exists(local_path):
                try:
                    with open(local_path, "rb") as lf:
                        local_bytes = lf.read()

                    rel = docx_doc.part.rels[r_embed]
                    orig_content_type = rel.target_part.content_type or ""
                    target_format = "PNG"
                    if "jpeg" in orig_content_type or "jpg" in orig_content_type:
                        target_format = "JPEG"
                    elif "gif" in orig_content_type:
                        target_format = "GIF"

                    orig_cx = mp.get("original_cx", 0)
                    orig_cy = mp.get("original_cy", 0)

                    new_img = PILImage.open(BytesIO(local_bytes))
                    if orig_cx > 0 and orig_cy > 0:
                        dpi = 150
                        target_w = max(1, int(orig_cx / 914400 * dpi))
                        target_h = max(1, int(orig_cy / 914400 * dpi))
                        src_w, src_h = new_img.size
                        if src_w > target_w or src_h > target_h:
                            new_img = new_img.resize((target_w, target_h), PILImage.LANCZOS)

                    if new_img.mode in ("RGBA", "P") and target_format == "JPEG":
                        new_img = new_img.convert("RGB")

                    buf = BytesIO()
                    new_img.save(buf, format=target_format, quality=95)
                    buf.seek(0)
                    rel.target_part._blob = buf.read()
                    replaced += 1
                    logger.info("Figure replacement (user upload): r_embed=%s replaced from %s", r_embed, local_path)
                    continue
                except Exception as local_err:
                    logger.warning("Figure replacement: failed to read local file %s: %s", local_path, local_err)
            else:
                logger.warning("Figure replacement: local file not found: %s", local_path)

        # Build ordered list of URLs to try: selected url, thumbnail, then other candidates
        urls_to_try = []
        primary_url = sel.get("url") or sel.get("thumbnail_url")
        if primary_url:
            urls_to_try.append(primary_url)
        thumb_url = sel.get("thumbnail_url") or sel.get("url")
        if thumb_url and thumb_url not in urls_to_try:
            urls_to_try.append(thumb_url)

        # Also try other replacement candidates if the selected one fails
        other_candidates = mp.get("replacement_candidates", [])
        for cand in other_candidates:
            for key in ("url", "thumbnail_url"):
                cand_url = cand.get(key)
                if cand_url and cand_url not in urls_to_try:
                    urls_to_try.append(cand_url)

        if not urls_to_try:
            continue

        try:
            img_resp = None
            download_url = None
            async with _httpx.AsyncClient(
                follow_redirects=True, headers=_BROWSER_HEADERS
            ) as _http:
                for attempt_url in urls_to_try:
                    try:
                        resp = await _http.get(attempt_url, timeout=30)
                        if resp.status_code == 200:
                            img_resp = resp
                            download_url = attempt_url
                            break
                        else:
                            logger.info(
                                "Figure replacement: URL returned %d for r_embed=%s — %s",
                                resp.status_code, r_embed, attempt_url[:150],
                            )
                    except Exception as url_err:
                        logger.info(
                            "Figure replacement: URL error for r_embed=%s — %s: %s",
                            r_embed, attempt_url[:150], url_err,
                        )

            if img_resp is None or img_resp.status_code != 200:
                logger.warning(
                    "Figure replacement: all %d URLs failed for r_embed=%s. "
                    "Tried: %s",
                    len(urls_to_try), r_embed,
                    [u[:80] for u in urls_to_try[:4]],
                )
                continue

            rel = docx_doc.part.rels[r_embed]

            # ── Determine original image format from rel ──
            orig_content_type = rel.target_part.content_type or ""
            target_format = "PNG"
            if "jpeg" in orig_content_type or "jpg" in orig_content_type:
                target_format = "JPEG"
            elif "gif" in orig_content_type:
                target_format = "GIF"

            # ── Get original dimensions from patch ──
            orig_cx = mp.get("original_cx", 0)
            orig_cy = mp.get("original_cy", 0)

            # ── Resize replacement image to match original dims (high quality) ──
            try:
                new_img = PILImage.open(BytesIO(img_resp.content))

                if orig_cx > 0 and orig_cy > 0:
                    # Use 150 DPI for better quality (1 inch = 914400 EMU)
                    dpi = 150
                    target_w = max(1, int(orig_cx / 914400 * dpi))
                    target_h = max(1, int(orig_cy / 914400 * dpi))
                    # Only downscale if needed; never upscale a small image
                    src_w, src_h = new_img.size
                    if src_w > target_w or src_h > target_h:
                        new_img = new_img.resize((target_w, target_h), PILImage.LANCZOS)
                    # If source is smaller, let it be — better than stretching a tiny image

                # Convert to RGB if saving as JPEG (no alpha)
                if target_format == "JPEG" and new_img.mode in ("RGBA", "P", "LA"):
                    background = PILImage.new("RGB", new_img.size, (255, 255, 255))
                    if new_img.mode == "P":
                        new_img = new_img.convert("RGBA")
                    background.paste(new_img, mask=new_img.split()[-1])
                    new_img = background

                buf = BytesIO()
                new_img.save(buf, format=target_format, quality=98)
                image_bytes = buf.getvalue()
            except Exception as resize_err:
                logger.warning("Image resize failed for r_embed=%s, using raw bytes: %s",
                               r_embed, resize_err)
                image_bytes = img_resp.content

            # ── Replace image blob ──
            rel.target_part._blob = image_bytes

            # ── Preserve DOCX XML extents (sizing) ──
            if orig_cx > 0 and orig_cy > 0:
                try:
                    WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
                    A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
                    R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
                    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

                    for para in docx_doc.paragraphs:
                        drawings = para._element.findall(f'.//{{{W_NS}}}drawing')
                        for drawing in drawings:
                            blip = drawing.find(f'.//{{{A_NS}}}blip')
                            if blip is not None and blip.get(f'{{{R_NS}}}embed') == r_embed:
                                extent = drawing.find(f'.//{{{WP_NS}}}extent')
                                if extent is not None:
                                    extent.set('cx', str(orig_cx))
                                    extent.set('cy', str(orig_cy))
                                for ext_el in drawing.findall(f'.//{{{A_NS}}}ext'):
                                    ext_el.set('cx', str(orig_cx))
                                    ext_el.set('cy', str(orig_cy))
                except Exception as extent_err:
                    logger.warning("Extent fix failed for r_embed=%s: %s", r_embed, extent_err)

            replaced += 1
            logger.info("Replaced figure image (r_embed=%s, cx=%d, cy=%d, format=%s, url=%s)",
                        r_embed, orig_cx, orig_cy, target_format, download_url[:120])

        except Exception as fig_err:
            logger.error("Figure replacement failed for r_embed=%s: %s", r_embed, fig_err)

    return replaced


def _apply_equation_replacements_to_docx(docx_doc, approved_equations: list) -> int:
    """Replace approved equation updates in a DOCX document.

    Finds original OMML equations and replaces the <m:t> text content
    to reflect the suggested update. Returns count of equations replaced.
    """
    from lxml import etree

    MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    replaced = 0

    for eq_patch in approved_equations:
        original_omml = eq_patch.get("original_omml", "")
        suggested = eq_patch.get("suggested_update", "")
        if not original_omml or not suggested:
            continue

        # Parse the original OMML to get its text content for matching
        try:
            orig_tree = etree.fromstring(original_omml)
            orig_t_elements = orig_tree.findall(f'.//{{{MATH_NS}}}t')
            orig_texts = [t.text or '' for t in orig_t_elements]
            orig_text_joined = ''.join(orig_texts).strip()
        except Exception:
            orig_text_joined = ""

        if not orig_text_joined:
            continue

        # Search all oMath elements in body paragraphs AND table cells
        def _find_omath_in_element(root_el):
            """Find all oMath elements in an XML tree."""
            return root_el.findall(f'.//{{{MATH_NS}}}oMath')

        found = False
        all_omath = _find_omath_in_element(docx_doc.element.body)

        for omath in all_omath:
            # Match by comparing <m:t> text content
            t_elements = omath.findall(f'.//{{{MATH_NS}}}t')
            current_text = ''.join(t.text or '' for t in t_elements).strip()

            if current_text == orig_text_joined:
                # Found the matching equation — apply the update
                # Strategy: Replace all <m:t> text nodes based on the suggested LaTeX

                # For simple text substitutions (e.g., \text{sin} → \sin means
                # replacing "sin" text run), update in-place
                updated_text = suggested.strip()

                # Remove LaTeX command prefixes for plain-text replacement
                # Convert suggested LaTeX back to plain text tokens
                import re as _re_eq_repl
                # Strip common LaTeX wrappers to get the core text changes
                plain_suggested = updated_text
                # Remove \frac{}{}, \sin, \cos etc. — keep the structure but update text
                # Simple approach: replace text content of existing <m:t> elements

                # Strategy 1: If number of <m:t> elements matches, do 1:1 replacement
                # Strategy 2: Combine all <m:t> into one with the suggested text

                # For tracked docs, we add a highlighted annotation after the equation
                # For clean docs, we update the <m:t> elements directly

                # Clear all existing <m:t> and set the first one to suggested text
                # This preserves the OMML structure but updates visible content
                if t_elements:
                    # Set first <m:t> to the full suggested text (cleaned)
                    # Remove LaTeX commands that are OMML structural (not text)
                    display_text = _re_eq_repl.sub(r'\\(sin|cos|tan|log|ln|exp|lim|max|min|det|gcd|partial|leq|geq|neq|approx|equiv|pm|times|cdot|frac|sqrt|sum|int|hat|bar|vec|overline|overbrace|underset|overset|left|right|middle|begin|end|text|mathbb|mathscr|mathfrak)\b', '', updated_text)
                    display_text = _re_eq_repl.sub(r'[{}\\]', '', display_text)
                    display_text = _re_eq_repl.sub(r'\s+', ' ', display_text).strip()

                    # If we can't produce clean display text, use the original with fixes applied
                    if not display_text or len(display_text) < 2:
                        # Fallback: apply specific known fixes to the original text elements
                        _apply_known_fixes_to_omml(omath, eq_patch, MATH_NS)
                    else:
                        # For now, don't replace the structural OMML — instead apply
                        # targeted fixes to the XML tree
                        _apply_known_fixes_to_omml(omath, eq_patch, MATH_NS)

                found = True
                replaced += 1
                logger.info("Replaced equation: '%s' (patch_id=%s)",
                            orig_text_joined[:50], eq_patch.get("patch_id", "?"))
                break

        if not found:
            logger.warning("Equation not found in DOCX for replacement: '%s'", orig_text_joined[:50])

    return replaced


def _apply_known_fixes_to_omml(omath_el, eq_patch: dict, math_ns: str):
    """Apply specific known fixes directly to the OMML XML tree.

    Handles:
    - \\text{sin} → proper function (remove rPr text formatting)
    - \\text{1}, \\text{+} → remove unnecessary text-mode wrapping
    - Missing spaces (\\partialx → \\partial x): add space <m:t> element
    """
    from lxml import etree

    reason = eq_patch.get("reason", "").lower()
    suggested = eq_patch.get("suggested_update", "")
    original = eq_patch.get("original_text", "")

    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # Fix 1: \text{} wrapping math functions — remove the rPr that forces text mode
    # In OMML, \text{sin} appears as <m:r> with <w:rPr><w:rFonts/></w:rPr> and <m:t>sin</m:t>
    # We need to remove the w:rPr that makes it render as text
    if '\\text{' in original or 'text{' in reason:
        for r_elem in omath_el.findall(f'.//{{{math_ns}}}r'):
            t_elem = r_elem.find(f'{{{math_ns}}}t')
            if t_elem is not None and t_elem.text:
                t_text = t_elem.text.strip()
                # Check if this is a function name wrapped in text mode
                func_names = ['sin', 'cos', 'tan', 'log', 'ln', 'exp', 'lim',
                              'max', 'min', 'det', 'gcd', 'sec', 'csc', 'cot',
                              'arcsin', 'arccos', 'arctan', 'sinh', 'cosh', 'tanh']
                if t_text in func_names:
                    # Remove w:rPr (text formatting) — this makes it render as math
                    for rpr in r_elem.findall(f'{{{W_NS}}}rPr'):
                        r_elem.remove(rpr)
                    # Also remove m:rPr > m:sty if it forces text
                    for mrpr in r_elem.findall(f'{{{math_ns}}}rPr'):
                        sty = mrpr.find(f'{{{math_ns}}}sty')
                        if sty is not None:
                            val = sty.get(f'{{{math_ns}}}val', '')
                            if val == 'p':  # 'p' = plain/text style
                                mrpr.remove(sty)

                # Remove \text{} wrapping around numbers and operators
                if t_text in '0123456789+-*/=<>' and len(t_text) == 1:
                    for rpr in r_elem.findall(f'{{{W_NS}}}rPr'):
                        r_elem.remove(rpr)
                    for mrpr in r_elem.findall(f'{{{math_ns}}}rPr'):
                        sty = mrpr.find(f'{{{math_ns}}}sty')
                        if sty is not None:
                            val = sty.get(f'{{{math_ns}}}val', '')
                            if val == 'p':
                                mrpr.remove(sty)

    # Fix 2: Missing spaces (e.g. \partialx → \partial x)
    if 'missing space' in reason or 'invalid command' in reason:
        t_elements = omath_el.findall(f'.//{{{math_ns}}}t')
        for t_elem in t_elements:
            if t_elem.text:
                text = t_elem.text
                import re as _re_space
                # Fix: "∂x" → "∂ x", "≤y" → "≤ y" etc.
                # Look for a symbol character immediately followed by a letter
                fixed = _re_space.sub(r'([\u2202\u2264\u2265\u2260\u2248\u2261])([a-zA-Z])',
                                      r'\1 \2', text)
                if fixed != text:
                    t_elem.text = fixed

    # Fix 3: Greek letter preservation — convert any spelled-out Greek names back to symbols
    # This catches cases where GPT's suggested update accidentally used English names
    _GREEK_NAME_TO_SYMBOL = {
        'Alpha': 'Α', 'Beta': 'Β', 'Gamma': 'Γ', 'Delta': 'Δ', 'Epsilon': 'Ε',
        'Zeta': 'Ζ', 'Eta': 'Η', 'Theta': 'Θ', 'Iota': 'Ι', 'Kappa': 'Κ',
        'Lambda': 'Λ', 'Mu': 'Μ', 'Nu': 'Ν', 'Xi': 'Ξ', 'Pi': 'Π',
        'Rho': 'Ρ', 'Sigma': 'Σ', 'Tau': 'Τ', 'Upsilon': 'Υ', 'Phi': 'Φ',
        'Chi': 'Χ', 'Psi': 'Ψ', 'Omega': 'Ω',
        'alpha': 'α', 'beta': 'β', 'gamma': 'γ', 'delta': 'δ', 'epsilon': 'ε',
        'zeta': 'ζ', 'eta': 'η', 'theta': 'θ', 'iota': 'ι', 'kappa': 'κ',
        'lambda': 'λ', 'mu': 'μ', 'nu': 'ν', 'xi': 'ξ', 'pi': 'π',
        'rho': 'ρ', 'sigma': 'σ', 'tau': 'τ', 'upsilon': 'υ', 'phi': 'φ',
        'chi': 'χ', 'psi': 'ψ', 'omega': 'ω',
    }
    t_elements = omath_el.findall(f'.//{{{math_ns}}}t')
    for t_elem in t_elements:
        if t_elem.text:
            text = t_elem.text
            for name, symbol in _GREEK_NAME_TO_SYMBOL.items():
                # Only replace standalone words (not substrings of other words)
                import re as _re_greek
                text = _re_greek.sub(r'\b' + name + r'\b', symbol, text)
            if text != t_elem.text:
                logger.debug("Greek letter fix in OMML: '%s' → '%s'", t_elem.text, text)
                t_elem.text = text


def _apply_table_updates_to_docx(docx_doc, approved_tables: list, highlight: bool = True) -> int:
    """Apply approved table cell updates to a DOCX document.

    Finds tables by their original DOCX index and updates specific cells.
    When highlight=True, adds light blue highlighting to changed cells for visibility.
    When highlight=False (clean document), no color is applied.
    Handles merged cells by searching for old_value text across all cells in the row.
    Returns count of tables updated.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    LIGHT_BLUE = "B4D8F0"  # Light blue highlight for changed cells

    def _highlight_cell(cell, hex_color=LIGHT_BLUE):
        """Add background shading (light blue) to a table cell."""
        if not highlight:
            return
        tc = cell._element
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.insert(0, tcPr)
        # Remove existing shading if any
        for old_shd in tcPr.findall(qn('w:shd')):
            tcPr.remove(old_shd)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def _set_cell_text_preserve_format(cell, new_val: str):
        """Set cell text while preserving formatting, then optionally highlight."""
        if cell.paragraphs:
            first_para = cell.paragraphs[0]
            # Clear all runs
            for run in first_para.runs:
                run.text = ""
            if first_para.runs:
                first_para.runs[0].text = new_val
            else:
                # No runs — add a new run with the text
                run = first_para.add_run(new_val)

            # Remove extra paragraphs
            for extra_para in list(cell.paragraphs[1:]):
                p_elem = extra_para._element
                p_elem.getparent().remove(p_elem)
        else:
            cell.text = new_val

        # Apply light blue highlight only for tracked documents
        _highlight_cell(cell)

    updated = 0

    for tbl_patch in approved_tables:
        cell_updates = tbl_patch.get("cell_updates", [])
        if not cell_updates:
            continue

        # Use table_idx (original DOCX index) for precise targeting
        t_idx = tbl_patch.get("table_idx")
        if t_idx is None:
            t_idx = tbl_patch.get("table_number", 1) - 1

        if t_idx < 0 or t_idx >= len(docx_doc.tables):
            logger.warning("Table index %d out of range (doc has %d tables)", t_idx, len(docx_doc.tables))
            continue

        tbl = docx_doc.tables[t_idx]
        cells_changed = 0

        for cu in cell_updates:
            r = cu.get("row", -1)
            c = cu.get("col", -1)
            new_val = cu.get("new_value", "")
            old_val = cu.get("old_value", "")

            if r < 0 or r >= len(tbl.rows):
                logger.warning("Row %d out of range for table %d", r, t_idx)
                continue

            row_cells = tbl.rows[r].cells

            if c < 0 or c >= len(row_cells):
                logger.warning("Col %d out of range for table %d row %d", c, t_idx, r)
                continue

            # Strategy 1: Direct cell at [r,c]
            cell = row_cells[c]
            current_text = cell.text.strip()

            # Check if old_value matches at the target cell
            matched_cell = None
            if not old_val:
                # No old_value provided — just update directly
                matched_cell = cell
            elif old_val.strip()[:20] in current_text or current_text[:20] in old_val.strip():
                matched_cell = cell
            else:
                # Strategy 2: Merged cells shift indices — search ALL cells in this row
                for search_c, search_cell in enumerate(row_cells):
                    search_text = search_cell.text.strip()
                    if old_val.strip()[:20] in search_text or search_text[:20] in old_val.strip():
                        matched_cell = search_cell
                        logger.info("Table %d cell [%d,%d]: found old_value in col %d instead of %d (merged cells)",
                                   t_idx, r, c, search_c, c)
                        break

                if matched_cell is None:
                    # Strategy 3: Skip validation entirely — trust GPT's row/col and just update
                    logger.info("Table %d cell [%d,%d]: old_value not found in any col, updating target cell directly",
                               t_idx, r, c)
                    matched_cell = cell

            _set_cell_text_preserve_format(matched_cell, new_val)
            cells_changed += 1
            logger.info("Table %d cell [%d,%d]: '%s' → '%s'",
                        t_idx, r, c, (old_val or current_text)[:30], new_val[:30])

        if cells_changed > 0:
            updated += 1
            logger.info("Updated table %d with %d cell changes", t_idx + 1, cells_changed)

    return updated


# ═══════════════════════════════════════════════════════════════════════════════
# Caption Updates, Renumbering, Page References, Headers/Footers Preservation
# ═══════════════════════════════════════════════════════════════════════════════

_W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_MATH_NS_RENUMBER = 'http://schemas.openxmlformats.org/officeDocument/2006/math'


def _add_figure_source_urls(docx_doc, approved_figures: list) -> int:
    """Add 'Source: <clickable URL>' annotation below the caption of each replaced figure.

    Finds the caption paragraph near the replaced figure and inserts a new
    paragraph right after it with a clickable hyperlink to the source page.
    Uses source_page_url (full webpage) when available, falls back to image URL.
    Returns count of source annotations added.
    """
    import re as _re
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from copy import deepcopy

    fig_caption_pat = _re.compile(
        r'(?:Figure|Fig\.?)\s+\d+[\-\.]\d+',
        _re.IGNORECASE,
    )

    added = 0
    R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'

    for mp in approved_figures:
        sel = mp.get("selected_replacement")
        if not sel:
            continue

        # Only use the source page URL (full website), never the raw image URL
        source_page = sel.get("source_page_url") or ""
        source_name = sel.get("source") or "Web"

        # SAFETY: Reject any URL that is actually a raw image file
        _img_exts = ('.jpg', '.jpeg', '.png', '.gif', '.webp', '.svg', '.bmp', '.tiff', '.ico')
        if source_page:
            from urllib.parse import urlparse as _urlparse_check
            _path_check = _urlparse_check(source_page).path.lower()
            if any(_path_check.endswith(ext) for ext in _img_exts):
                logger.warning("Blocked image URL from appearing as source: %s", source_page[:100])
                source_page = ""

        if not source_page:
            # No valid source page available — skip adding source annotation
            continue
        display_url = source_page
        link_url = source_page

        r_embed = mp.get("r_embed")
        if not r_embed:
            continue

        # Find paragraph containing this figure's drawing
        fig_para_idx = None
        for idx, para in enumerate(docx_doc.paragraphs):
            drawings = para._element.findall(f'.//{{{_W_NS}}}drawing')
            for drawing in drawings:
                blip = drawing.find(f'.//{{{A_NS}}}blip')
                if blip is not None and blip.get(f'{{{R_NS}}}embed') == r_embed:
                    fig_para_idx = idx
                    break
            if fig_para_idx is not None:
                break

        if fig_para_idx is None:
            continue

        # Find caption paragraph near the figure (up to 3 before/after)
        total_paras = len(docx_doc.paragraphs)
        search_range = list(range(max(0, fig_para_idx - 3), min(total_paras, fig_para_idx + 4)))

        caption_para = None
        for check_idx in search_range:
            para = docx_doc.paragraphs[check_idx]
            text = para.text.strip()
            if fig_caption_pat.match(text):
                caption_para = para
                break

        # If no caption found, use the figure paragraph itself
        if caption_para is None:
            caption_para = docx_doc.paragraphs[fig_para_idx]

        # ── Build run formatting (italic, 8pt, gray) ──
        def _make_styled_rPr():
            rPr = OxmlElement('w:rPr')
            i_el = OxmlElement('w:i')
            i_el.set(qn('w:val'), 'true')
            rPr.append(i_el)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '16')  # 8pt
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '16')
            rPr.append(szCs)
            color_el = OxmlElement('w:color')
            color_el.set(qn('w:val'), '666666')
            rPr.append(color_el)
            return rPr

        # ── Create paragraph ──
        new_p = OxmlElement('w:p')

        # Copy paragraph alignment from caption
        caption_pPr = caption_para._element.find(qn('w:pPr'))
        if caption_pPr is not None:
            new_pPr = deepcopy(caption_pPr)
            new_p.append(new_pPr)

        # ── "Source: " label (plain italic text) ──
        label_run = OxmlElement('w:r')
        label_run.append(_make_styled_rPr())
        label_t = OxmlElement('w:t')
        label_t.set(qn('xml:space'), 'preserve')
        label_t.text = f"Source ({source_name}): "
        label_run.append(label_t)
        new_p.append(label_run)

        # ── Clickable hyperlink ──
        # Add external hyperlink relationship directly via the document part
        HYPERLINK_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
        r_id = docx_doc.part.rels._next_rId if hasattr(docx_doc.part.rels, '_next_rId') else f"rId{len(docx_doc.part.rels) + 100}"
        try:
            # python-docx >=0.8.11 approach
            docx_doc.part.relate_to(link_url, HYPERLINK_REL_TYPE, is_external=True)
            # Get the rId that was just created
            for rel_key, rel_val in docx_doc.part.rels.items():
                if rel_val.target_ref == link_url and rel_val.is_external:
                    r_id = rel_key
                    break
        except Exception:
            try:
                # Fallback: add_relationship directly
                docx_doc.part.rels.add_relationship(
                    HYPERLINK_REL_TYPE, link_url, r_id, is_external=True
                )
            except Exception as rel_err:
                logger.warning("Failed to create hyperlink relationship: %s", rel_err)
                # Last resort: just add as plain text
                plain_run = OxmlElement('w:r')
                plain_run.append(_make_styled_rPr())
                plain_t = OxmlElement('w:t')
                plain_t.set(qn('xml:space'), 'preserve')
                plain_t.text = display_url
                plain_run.append(plain_t)
                new_p.append(plain_run)
                caption_para._element.addnext(new_p)
                added += 1
                continue

        # Create w:hyperlink element with the relationship ID
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        # Also set w:history for Word to track it as visited
        hyperlink.set(qn('w:history'), '1')

        # Hyperlink run — blue underlined clickable text
        link_run = OxmlElement('w:r')
        link_rPr = OxmlElement('w:rPr')

        # Hyperlink style (this is what makes Word treat it as clickable)
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')
        link_rPr.append(rStyle)

        # Italic
        i_el = OxmlElement('w:i')
        i_el.set(qn('w:val'), 'true')
        link_rPr.append(i_el)

        # Font size 8pt
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '16')
        link_rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '16')
        link_rPr.append(szCs)

        # Blue color for hyperlink
        color_el = OxmlElement('w:color')
        color_el.set(qn('w:val'), '0563C1')
        link_rPr.append(color_el)

        # Underline (single)
        u_el = OxmlElement('w:u')
        u_el.set(qn('w:val'), 'single')
        link_rPr.append(u_el)

        link_run.append(link_rPr)

        link_t = OxmlElement('w:t')
        link_t.set(qn('xml:space'), 'preserve')
        link_t.text = display_url
        link_run.append(link_t)

        hyperlink.append(link_run)
        new_p.append(hyperlink)

        # Insert the new paragraph right after the caption paragraph
        caption_para._element.addnext(new_p)

        added += 1
        logger.info("Added source URL below caption near para %d: %s (page: %s)",
                    fig_para_idx, source_name, display_url[:100])

    return added


def _update_figure_captions_in_docx(docx_doc, approved_figures: list) -> int:
    """Update figure captions in the DOCX when a figure is replaced.

    Scans paragraphs near the replaced figure for a caption like
    "Figure 1-2. Old caption text" and updates it with the new caption
    from the approved patch's selected_replacement.

    Returns count of captions updated.
    """
    import re as _re

    fig_caption_pat = _re.compile(
        r'((?:Figure|Fig\.?)\s+\d+[\-\.]\d+)([.:]\s*)(.*)',
        _re.IGNORECASE | _re.DOTALL,
    )

    updated = 0
    for mp in approved_figures:
        sel = mp.get("selected_replacement")
        if not sel:
            continue
        new_caption_text = sel.get("caption") or sel.get("title") or ""
        if not new_caption_text:
            continue

        r_embed = mp.get("r_embed")
        if not r_embed:
            continue

        # Find paragraph containing this figure's drawing, then look at nearby paragraphs for caption
        R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'

        fig_para_idx = None
        for idx, para in enumerate(docx_doc.paragraphs):
            drawings = para._element.findall(f'.//{{{_W_NS}}}drawing')
            for drawing in drawings:
                blip = drawing.find(f'.//{{{A_NS}}}blip')
                if blip is not None and blip.get(f'{{{R_NS}}}embed') == r_embed:
                    fig_para_idx = idx
                    break
            if fig_para_idx is not None:
                break

        if fig_para_idx is None:
            continue

        # Search nearby paragraphs (up to 3 before and after) for caption text
        total_paras = len(docx_doc.paragraphs)
        search_range = list(range(max(0, fig_para_idx - 3), min(total_paras, fig_para_idx + 4)))

        for check_idx in search_range:
            para = docx_doc.paragraphs[check_idx]
            text = para.text.strip()
            m = fig_caption_pat.match(text)
            if m:
                # Found caption paragraph — update the description part, keep "Figure X-Y."
                fig_prefix = m.group(1)  # e.g. "Figure 1-2"
                separator = m.group(2)   # e.g. ". " or ": "
                new_full = f"{fig_prefix}{separator}{new_caption_text}"

                # Replace text in runs while preserving formatting
                _replace_paragraph_text_preserve_format(para, new_full)
                updated += 1
                logger.info("Updated figure caption near para %d: '%s' → '%s'",
                            fig_para_idx, text[:50], new_full[:50])
                break

    return updated


def _update_table_captions_in_docx(docx_doc, approved_tables: list) -> int:
    """Update table captions in DOCX when table analysis provides updated captions.

    Returns count of captions updated.
    """
    import re as _re

    tbl_caption_pat = _re.compile(
        r'((?:TABLE|Table|Tbl\.?)\s+\d+[\-\.]\d+)([.:]\s*)(.*)',
        _re.IGNORECASE | _re.DOTALL,
    )

    updated = 0
    for mp in approved_tables:
        new_caption = mp.get("updated_caption", "")
        if not new_caption:
            continue

        table_idx = mp.get("table_idx", -1)
        if table_idx < 0 or table_idx >= len(docx_doc.tables):
            continue

        # Walk body children to find the table element and look at preceding paragraphs
        tbl_element = docx_doc.tables[table_idx]._tbl
        body = docx_doc.element.body
        children = list(body)

        tbl_pos = None
        for i, child in enumerate(children):
            if child is tbl_element:
                tbl_pos = i
                break

        if tbl_pos is None:
            continue

        # Check up to 3 paragraphs before the table for caption
        for offset in range(1, 4):
            check_pos = tbl_pos - offset
            if check_pos < 0:
                break
            child = children[check_pos]
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag != 'p':
                continue

            texts = []
            for t_el in child.iter(f'{{{_W_NS}}}t'):
                if t_el.text:
                    texts.append(t_el.text)
            para_text = ''.join(texts).strip()

            m = tbl_caption_pat.match(para_text)
            if m:
                tbl_prefix = m.group(1)
                separator = m.group(2)
                new_full = f"{tbl_prefix}{separator}{new_caption}"

                # Find the python-docx paragraph for this element
                for para in docx_doc.paragraphs:
                    if para._element is child:
                        _replace_paragraph_text_preserve_format(para, new_full)
                        updated += 1
                        logger.info("Updated table caption: '%s' → '%s'",
                                    para_text[:50], new_full[:50])
                        break
                break

    return updated


def _replace_paragraph_text_preserve_format(para, new_text: str):
    """Replace paragraph text while preserving the formatting of the first run."""
    runs = para.runs
    if not runs:
        # No runs — create one
        para.text = new_text
        return

    # Keep first run's formatting, set its text, clear the rest
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def _update_equation_numbers_in_omml(docx_doc, number_map: dict) -> int:
    """Update equation number labels in OMML elements within the DOCX.

    number_map: {old_number: new_number} e.g. {"6-4": "6-3", "6-5": "6-4"}

    Equation numbers in OMML appear as <m:t> text like "(6-4)" near the end
    of the equation's oMath element (often in a separate oMath or run).
    Also updates standalone text runs that contain equation numbers like "(6-4)".

    Returns count of equation numbers updated.
    """
    import re as _re

    if not number_map:
        return 0

    updated = 0

    # Pattern to find equation numbers like (6-4), (1.2), etc.
    eq_num_pat = _re.compile(r'\((\d+[\-\.]\d+)\)')

    # 1. Update equation numbers in OMML <m:t> elements
    for omath in docx_doc.element.body.iter(f'{{{_MATH_NS_RENUMBER}}}oMath'):
        for mt in omath.iter(f'{{{_MATH_NS_RENUMBER}}}t'):
            if mt.text:
                m = eq_num_pat.search(mt.text)
                if m and m.group(1) in number_map:
                    old_num = m.group(1)
                    new_num = number_map[old_num]
                    mt.text = mt.text.replace(f"({old_num})", f"({new_num})")
                    updated += 1
                    logger.info("OMML equation number: (%s) → (%s)", old_num, new_num)

    # 2. Update equation numbers in regular text runs (some DOCX files have the
    #    equation label as a plain text run next to the oMath element)
    for para in docx_doc.paragraphs:
        for run in para.runs:
            if run.text:
                m = eq_num_pat.search(run.text)
                if m and m.group(1) in number_map:
                    old_num = m.group(1)
                    new_num = number_map[old_num]
                    run.text = run.text.replace(f"({old_num})", f"({new_num})")
                    updated += 1

    return updated


def _renumber_captions_and_references(docx_doc, ref_type: str, number_map: dict) -> int:
    """Renumber figure/table/equation captions and all text references in the DOCX.

    ref_type: "figure", "table", or "equation"
    number_map: {old_number: new_number} e.g. {"1-5": "1-4", "1-6": "1-5"}

    Updates:
    1. Caption paragraphs (e.g., "Figure 1-5." → "Figure 1-4.")
    2. All text references (e.g., "see Figure 1-5" → "see Figure 1-4")
    3. For equations: also updates OMML labels

    Returns total count of replacements made.
    """
    import re as _re

    if not number_map:
        return 0

    updated = 0

    # Build regex patterns for this reference type
    if ref_type == "figure":
        patterns = [
            _re.compile(rf'\b((?:Figure|Fig\.?)\s+){_re.escape(old)}\b', _re.IGNORECASE)
            for old in number_map
        ]
    elif ref_type == "table":
        patterns = [
            _re.compile(rf'\b((?:TABLE|Table|Tbl\.?)\s+){_re.escape(old)}\b', _re.IGNORECASE)
            for old in number_map
        ]
    elif ref_type == "equation":
        patterns = [
            _re.compile(rf'\b((?:Equation|Eq\.?)\s*\(){_re.escape(old)}(\))', _re.IGNORECASE)
            for old in number_map
        ]
    else:
        return 0

    old_numbers = list(number_map.keys())

    for para in docx_doc.paragraphs:
        original_text = para.text
        if not original_text:
            continue

        # Check if any old number exists in this paragraph
        has_match = False
        for old_num in old_numbers:
            if old_num in original_text:
                has_match = True
                break
        if not has_match:
            continue

        # Apply all replacements to each run
        for run in para.runs:
            if not run.text:
                continue
            new_text = run.text
            for i, old_num in enumerate(old_numbers):
                if old_num not in new_text:
                    continue
                new_num = number_map[old_num]
                if ref_type == "equation":
                    new_text = new_text.replace(f"({old_num})", f"({new_num})")
                else:
                    new_text = patterns[i].sub(rf'\g<1>{new_num}', new_text)
            if new_text != run.text:
                updated += 1
                run.text = new_text

    # For equations: also update OMML elements
    if ref_type == "equation":
        updated += _update_equation_numbers_in_omml(docx_doc, number_map)

    if updated:
        logger.info("Renumbered %d %s references: %s", updated, ref_type, number_map)

    return updated


def _build_sequential_number_map(
    existing_numbers: list, removed_numbers: set, chapter: str = None
) -> dict:
    """Build a renumbering map to maintain sequential numbering after removals.

    existing_numbers: sorted list of all current numbers e.g. ["1-1", "1-2", "1-3", "1-5"]
    removed_numbers: set of numbers that were removed e.g. {"1-3"}

    Returns: {old_number: new_number} only for items that need renumbering.
    e.g. {"1-5": "1-4"} (1-3 removed, so 1-5 becomes 1-4)
    """
    import re as _re

    if not removed_numbers:
        return {}

    # Group numbers by chapter prefix
    chapter_groups = {}
    num_pat = _re.compile(r'^(\d+)([\-\.])(\d+)$')

    for num in existing_numbers:
        m = num_pat.match(num)
        if m:
            chap = m.group(1)
            sep = m.group(2)
            seq = int(m.group(3))
            chapter_groups.setdefault((chap, sep), []).append((seq, num))

    number_map = {}
    for (chap, sep), items in chapter_groups.items():
        # Sort by sequence number
        items.sort(key=lambda x: x[0])
        # Filter out removed items and assign new sequential numbers
        remaining = [(seq, num) for seq, num in items if num not in removed_numbers]
        for new_idx, (old_seq, old_num) in enumerate(remaining, start=1):
            new_num = f"{chap}{sep}{new_idx}"
            if new_num != old_num:
                number_map[old_num] = new_num

    return number_map


# ── Page Number Reference Detection & Updates ──

_PAGE_REF_PATTERN = re.compile(
    r'\b(?:'
    r'(?:pages?|pp?\.?)\s+(\d{1,5})\s*[-–—]\s*(\d{1,5})'   # "pages 145-147", "pp. 10-15"
    r'|(?:pages?|pp?\.?)\s+(\d{1,5})'                        # "page 145", "p. 42"
    r'|(?:around|approximately|about|near)\s+page\s+(\d{1,5})'  # "around page 145"
    r')\b',
    re.IGNORECASE,
)


def _detect_page_references(docx_doc) -> list:
    """Detect all hardcoded page number references in the document.

    Returns list of dicts: [{paragraph_idx, text, page_numbers, raw_match}, ...]
    """
    refs = []
    for idx, para in enumerate(docx_doc.paragraphs):
        text = para.text
        if not text:
            continue
        for m in _PAGE_REF_PATTERN.finditer(text):
            pages = []
            if m.group(1) and m.group(2):  # range
                pages = list(range(int(m.group(1)), int(m.group(2)) + 1))
            elif m.group(3):  # single
                pages = [int(m.group(3))]
            elif m.group(4):  # approximate
                pages = [int(m.group(4))]

            refs.append({
                "paragraph_idx": idx,
                "text": text[:80],
                "page_numbers": pages,
                "raw_match": m.group(0),
                "start": m.start(),
                "end": m.end(),
            })
    return refs


def _update_page_references(docx_doc, page_offset_map: dict) -> int:
    """Update hardcoded page number references based on an offset map.

    page_offset_map: {old_page: new_page} e.g. {145: 147, 146: 148}

    Returns count of references updated.
    """
    updated = 0
    for para in docx_doc.paragraphs:
        for run in para.runs:
            if not run.text:
                continue
            new_text = run.text
            changed = False
            for m in _PAGE_REF_PATTERN.finditer(run.text):
                if m.group(1) and m.group(2):
                    old_start, old_end = int(m.group(1)), int(m.group(2))
                    new_start = page_offset_map.get(old_start, old_start)
                    new_end = page_offset_map.get(old_end, old_end)
                    if new_start != old_start or new_end != old_end:
                        old_str = m.group(0)
                        new_str = old_str.replace(str(old_start), str(new_start)).replace(str(old_end), str(new_end))
                        new_text = new_text.replace(old_str, new_str)
                        changed = True
                elif m.group(3):
                    old_page = int(m.group(3))
                    new_page = page_offset_map.get(old_page)
                    if new_page and new_page != old_page:
                        new_text = new_text.replace(m.group(0), m.group(0).replace(str(old_page), str(new_page)))
                        changed = True
                elif m.group(4):
                    old_page = int(m.group(4))
                    new_page = page_offset_map.get(old_page)
                    if new_page and new_page != old_page:
                        new_text = new_text.replace(m.group(0), m.group(0).replace(str(old_page), str(new_page)))
                        changed = True
            if changed:
                run.text = new_text
                updated += 1

    if updated:
        logger.info("Updated %d page number references", updated)
    return updated


# ── Headers/Footers Preservation ──

def _preserve_headers_footers(source_doc, target_doc):
    """Copy headers and footers from source DOCX to target DOCX.

    This ensures headers/footers are preserved during DOCX reconstruction.
    Copies header/footer relationships and XML parts from all sections.
    """
    from copy import deepcopy

    try:
        # Copy headers and footers for each section
        for src_section, tgt_section in zip(source_doc.sections, target_doc.sections):
            # Copy header
            if src_section.header and src_section.header.is_linked_to_previous is False:
                try:
                    # Deep copy the header XML
                    src_header_el = src_section.header._element
                    tgt_header_el = tgt_section.header._element
                    # Clear target header and copy source content
                    for child in list(tgt_header_el):
                        tgt_header_el.remove(child)
                    for child in src_header_el:
                        tgt_header_el.append(deepcopy(child))
                except Exception as hdr_err:
                    logger.debug("Header copy failed for a section: %s", hdr_err)

            # Copy footer
            if src_section.footer and src_section.footer.is_linked_to_previous is False:
                try:
                    src_footer_el = src_section.footer._element
                    tgt_footer_el = tgt_section.footer._element
                    for child in list(tgt_footer_el):
                        tgt_footer_el.remove(child)
                    for child in src_footer_el:
                        tgt_footer_el.append(deepcopy(child))
                except Exception as ftr_err:
                    logger.debug("Footer copy failed for a section: %s", ftr_err)

            # Preserve page layout settings from source
            try:
                src_sectPr = src_section._sectPr
                tgt_sectPr = tgt_section._sectPr
                # Copy page size
                for tag in ['pgSz', 'pgMar', 'cols', 'docGrid']:
                    src_el = src_sectPr.find(f'{{{_W_NS}}}{tag}')
                    tgt_el = tgt_sectPr.find(f'{{{_W_NS}}}{tag}')
                    if src_el is not None:
                        if tgt_el is not None:
                            tgt_sectPr.remove(tgt_el)
                        tgt_sectPr.append(deepcopy(src_el))
            except Exception as layout_err:
                logger.debug("Page layout copy failed: %s", layout_err)

        logger.info("Headers/footers preserved across %d sections", len(source_doc.sections))
    except Exception as e:
        logger.warning("Headers/footers preservation failed (non-fatal): %s", e)


def _build_tracked_changes_docx(original_path: str, approved_patches: list, output_path: str,
                                 outline: list | None = None):
    """Build a visual diff DOCX from the original + approved patches.

    Visual format (no Review tab needed — visible immediately):
    - Regular patches: original text with RED highlight + strikethrough,
      new replacement text with GREEN highlight, both inline
    - Ask AI patches (new content): with YELLOW highlight, matching document formatting
    """
    from docx import Document as DocxDocument
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from copy import deepcopy

    docx_doc = DocxDocument(original_path)
    logger.info("TRACKED: building visual diff from %d patches", len(approved_patches))

    regular_applied = 0
    ai_applied = 0
    media_inserted = 0
    regular_patches = []  # collect for grouped-per-paragraph application

    for patch in approved_patches:
        final_text = patch.get("editor_revision") or patch.get("revised_sentence", "")
        ask_ai_meta = patch.get("ask_ai_meta")
        insert_media_meta = patch.get("insert_media_meta")
        research_insert_meta = patch.get("research_insert_meta")

        # --- Research-insert patches: drop drafted content at chosen location
        # (highlighted as tracked-insert). Content field holds the final edited
        # prose. Placement fields mirror insert-media / ask-ai semantics.
        if research_insert_meta:
            try:
                if _apply_research_insert_patch(
                    docx_doc, research_insert_meta, tracked=True,
                    outline=outline or [],
                ):
                    ai_applied += 1
                    logger.info(
                        "TRACKED: research-insert applied (section=%s, position=%s)",
                        (research_insert_meta.get("section_text") or "")[:40],
                        research_insert_meta.get("position"),
                    )
            except Exception as ri_err:
                logger.error(
                    "TRACKED: research-insert failed (patch_id=%s): %s",
                    patch.get("patch_id"), ri_err,
                )
            continue

        # --- Insert Media patches: insert uploaded figure/table + renumber ---
        if insert_media_meta:
            try:
                if _apply_insert_media_patch(docx_doc, insert_media_meta):
                    media_inserted += 1
                    logger.info(
                        "TRACKED: insert-media applied (%s, section=%s)",
                        insert_media_meta.get("media_type"),
                        (insert_media_meta.get("section_text") or "")[:40],
                    )
            except Exception as im_err:
                logger.error("TRACKED: insert-media failed (patch_id=%s): %s",
                             patch.get("patch_id"), im_err)
            continue

        if not final_text:
            continue

        try:
            if ask_ai_meta:
                section_text = ask_ai_meta.get("section_text", "").strip()
                position = ask_ai_meta.get("position", "after")
                para_text = ask_ai_meta.get("paragraph_text", "").strip()
                para_idx = ask_ai_meta.get("paragraph_index", -1)

                target_element = None
                ref_para = None

                if para_text and para_idx >= 0:
                    found_idx = None
                    if para_idx < len(docx_doc.paragraphs) and para_text[:50] in docx_doc.paragraphs[para_idx].text:
                        found_idx = para_idx
                    else:
                        for si, p in enumerate(docx_doc.paragraphs):
                            if para_text[:50] in p.text:
                                found_idx = si
                                break
                    if found_idx is not None:
                        target_element = docx_doc.paragraphs[found_idx]._element
                        ref_para = docx_doc.paragraphs[found_idx]

                if target_element is None and section_text:
                    heading_idx = _find_heading_in_docx(docx_doc, section_text)
                    if heading_idx is not None:
                        target_element = docx_doc.paragraphs[heading_idx]._element
                        ref_para = _find_reference_body_paragraph(docx_doc, heading_idx)

                if target_element is None:
                    logger.warning("TRACKED: target not found for '%s'", section_text[:50])
                    continue

                lines = [line.strip() for line in final_text.split("\n") if line.strip()]

                if position == "before":
                    for line in reversed(lines):
                        new_p = _make_highlighted_paragraph_xml(ref_para, line, "yellow")
                        target_element.addprevious(new_p)
                else:
                    insert_after = target_element
                    for line in lines:
                        new_p = _make_highlighted_paragraph_xml(ref_para, line, "yellow")
                        insert_after.addnext(new_p)
                        insert_after = new_p

                ai_applied += 1
                logger.info("TRACKED: AI patch applied (yellow highlight) near '%s'",
                            (para_text or section_text)[:40])

            else:
                original_text = patch.get("original_sentence", "")
                if not original_text or not final_text:
                    continue
                regular_patches.append({
                    "original_text": original_text,
                    "final_text": final_text,
                    "is_custom": patch.get("is_custom"),
                    "scope": patch.get("scope"),
                    "patch_id": patch.get("patch_id"),
                })

        except Exception as e:
            logger.error("TRACKED: patch failed (patch_id=%s): %s", patch.get("patch_id"), e)

    # ── Apply regular patches grouped by paragraph (single pass per paragraph) ──
    # This prevents multiple patches on the same paragraph from corrupting each other
    def _all_paragraphs(doc):
        for p in doc.paragraphs:
            yield p
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p

    applied_patch_ids = set()
    for para in _all_paragraphs(docx_doc):
        para_text = para.text or ""
        if not para_text:
            continue

        # Find ALL patches that match this paragraph
        matches = []
        for rp in regular_patches:
            old_t = rp["original_text"]
            replace_all = rp.get("is_custom") and rp.get("scope") == "whole_document"
            pid = rp.get("patch_id", "")

            # Skip non-custom patches already applied to a previous paragraph
            if not replace_all and pid in applied_patch_ids:
                continue

            if old_t in para_text:
                if replace_all:
                    search_start = 0
                    while True:
                        idx = para_text.find(old_t, search_start)
                        if idx == -1:
                            break
                        matches.append((idx, old_t, rp["final_text"]))
                        search_start = idx + len(old_t)
                else:
                    idx = para_text.index(old_t)
                    matches.append((idx, old_t, rp["final_text"]))
                    applied_patch_ids.add(pid)

        if not matches:
            continue

        # Sort by position, remove overlaps
        matches.sort(key=lambda m: m[0])
        filtered = []
        last_end = 0
        for start, old_t, new_t in matches:
            if start >= last_end:
                filtered.append((start, old_t, new_t))
                last_end = start + len(old_t)
        matches = filtered

        # Apply all matches in a single pass
        _apply_multi_visual_diff_to_paragraph(para, matches)
        regular_applied += len(matches)
        for _, old_t, _ in matches:
            logger.info("TRACKED: patch applied (strikethrough/green) for '%s...'", old_t[:40])

    # Log patches that were not found
    for rp in regular_patches:
        pid = rp.get("patch_id", "")
        if pid and pid not in applied_patch_ids and not (rp.get("is_custom") and rp.get("scope") == "whole_document"):
            logger.warning("TRACKED: original text not found: '%s...'", rp["original_text"][:50])

    # --- Append new references to the References/Bibliography section (cyan highlight) ---
    refs_added = _append_references_to_docx(docx_doc, approved_patches, tracked=True, highlight_color="cyan")
    if refs_added:
        logger.info("TRACKED: added %d new reference entries to References section", refs_added)

    docx_doc.save(output_path)
    logger.info(
        "TRACKED: visual diff DOCX saved — %d regular, %d AI, %d insert-media patches applied",
        regular_applied, ai_applied, media_inserted,
    )


def _make_highlighted_paragraph_xml(ref_para, text: str, highlight_color: str):
    """Create a paragraph element with highlighted text, copying formatting from ref_para.
    Uses pure XML — no docx_doc.add_paragraph() to avoid document corruption.
    """
    from copy import deepcopy
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    new_p = OxmlElement("w:p")

    # Copy paragraph properties (alignment, spacing, indentation) from reference
    if ref_para:
        source_pPr = ref_para._element.find(qn("w:pPr"))
        if source_pPr is not None:
            new_p.append(deepcopy(source_pPr))

    # Build run with text
    new_r = OxmlElement("w:r")

    # Build run properties: copy from reference + add highlight
    rPr = OxmlElement("w:rPr")
    if ref_para:
        source_runs = ref_para._element.findall(qn("w:r"))
        if source_runs:
            source_rPr = source_runs[0].find(qn("w:rPr"))
            if source_rPr is not None:
                rPr = deepcopy(source_rPr)

        # If rPr is still empty (run inherits from style), resolve font from
        # the paragraph style so highlighted text matches the document body.
        has_font = rPr.find(qn("w:rFonts")) is not None
        has_size = rPr.find(qn("w:sz")) is not None
        if not has_font or not has_size:
            try:
                style = ref_para.style
                # Walk up style hierarchy to find font info
                font_name = None
                font_size = None
                checked = set()
                while style and style.name not in checked:
                    checked.add(style.name)
                    if style.font:
                        if not font_name and style.font.name:
                            font_name = style.font.name
                        if not font_size and style.font.size:
                            font_size = style.font.size
                    if font_name and font_size:
                        break
                    style = style.base_style

                if font_name and not has_font:
                    rFonts = OxmlElement("w:rFonts")
                    rFonts.set(qn("w:ascii"), font_name)
                    rFonts.set(qn("w:hAnsi"), font_name)
                    rFonts.set(qn("w:cs"), font_name)
                    rPr.insert(0, rFonts)
                    logger.debug("_make_highlighted_paragraph_xml: resolved font '%s' from style", font_name)

                if font_size and not has_size:
                    # python-docx font.size is in EMU; Word XML sz is in half-points
                    half_points = str(int(font_size / 6350))
                    sz = OxmlElement("w:sz")
                    sz.set(qn("w:val"), half_points)
                    rPr.append(sz)
                    szCs = OxmlElement("w:szCs")
                    szCs.set(qn("w:val"), half_points)
                    rPr.append(szCs)
                    logger.debug("_make_highlighted_paragraph_xml: resolved size %s half-pts from style", half_points)
            except Exception as exc:
                logger.debug("_make_highlighted_paragraph_xml: style font resolve failed: %s", exc)

    # Add the highlight color
    _add_highlight_to_rPr(rPr, highlight_color)
    new_r.append(rPr)

    # Add text
    new_t = OxmlElement("w:t")
    new_t.set(qn("xml:space"), "preserve")
    new_t.text = text
    new_r.append(new_t)

    new_p.append(new_r)
    return new_p


def _apply_visual_diff_to_paragraph(para, old_text: str, new_text: str):
    """Apply visual diff: original in red strikethrough+red highlight, new in green highlight.
    All visible directly — no Review tab needed.

    Uses para.text (which captures ALL text including hyperlinks, field codes, etc.)
    rather than just para.runs to correctly locate the old text position.
    Then walks the actual XML child elements (w:r, w:hyperlink, etc.) to split
    runs at the exact boundaries of the matched text.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from copy import deepcopy

    # Use para.text which captures text from ALL child elements, not just runs
    full_text = para.text or ""
    if old_text not in full_text:
        return

    # ── Collect ALL text-bearing elements with their character offsets ──
    def _get_run_text(r_elem):
        """Get concatenated text from all w:t children of a run."""
        parts = []
        for t in r_elem.findall(qn("w:t")):
            parts.append(t.text or "")
        return "".join(parts)

    def _get_elem_rPr(r_elem):
        """Get run properties from a run element."""
        rPr = r_elem.find(qn("w:rPr"))
        return deepcopy(rPr) if rPr is not None else None

    # Build a flat list: each entry is (xml_element, text, char_offset)
    # This handles w:r (runs), w:hyperlink > w:r, etc.
    text_runs = []
    char_offset = 0
    for child in list(para._element):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "r":
            t = _get_run_text(child)
            if t:
                text_runs.append({"elem": child, "text": t, "offset": char_offset, "rPr": _get_elem_rPr(child)})
                char_offset += len(t)
        elif tag == "hyperlink":
            for sub_r in child.findall(qn("w:r")):
                t = _get_run_text(sub_r)
                if t:
                    text_runs.append({"elem": sub_r, "parent": child, "text": t, "offset": char_offset, "rPr": _get_elem_rPr(sub_r)})
                    char_offset += len(t)

    # Verify the reconstructed text matches what we expect
    reconstructed = "".join(tr["text"] for tr in text_runs)

    # If our reconstruction doesn't contain the old_text, try a simpler fallback
    if old_text not in reconstructed:
        # Fallback: use para.text directly
        reconstructed = full_text

    # If still no match in reconstructed, give up
    if old_text not in reconstructed:
        return

    # Recalculate start/end based on reconstructed text
    r_start = reconstructed.index(old_text)
    r_end = r_start + len(old_text)

    # Get default formatting from the first run
    default_rPr = None
    if text_runs:
        default_rPr = text_runs[0].get("rPr")

    def make_run(text, rPr_source=None):
        r = OxmlElement("w:r")
        if rPr_source:
            r.append(deepcopy(rPr_source))
        else:
            r.append(OxmlElement("w:rPr"))
        t_elem = OxmlElement("w:t")
        t_elem.set(qn("xml:space"), "preserve")
        t_elem.text = text
        r.append(t_elem)
        return r

    # ── Remove all text-bearing children from the paragraph ──
    # Keep non-text elements like w:pPr (paragraph properties), w:bookmarkStart, etc.
    for child in list(para._element):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("r", "hyperlink", "smartTag", "fldSimple"):
            para._element.remove(child)

    # ── Rebuild the paragraph with the visual diff ──
    before_text = reconstructed[:r_start]
    after_text = reconstructed[r_end:]

    # 1. Unchanged text before — preserve per-run formatting where possible
    if before_text:
        # Try to reconstruct with original per-run formatting
        pos = 0
        for tr in text_runs:
            tr_start = tr["offset"]
            tr_end = tr_start + len(tr["text"])
            if tr_end <= r_start and tr_start >= pos:
                para._element.append(make_run(tr["text"], tr.get("rPr") or default_rPr))
                pos = tr_end
            elif tr_start < r_start and tr_end > tr_start:
                # Partial overlap — take the part before old_text
                overlap_end = min(tr_end, r_start)
                if overlap_end > tr_start and tr_start >= pos:
                    para._element.append(make_run(tr["text"][:overlap_end - tr_start], tr.get("rPr") or default_rPr))
                    pos = overlap_end
                break
        # If we missed some before text, add it with default formatting
        if pos < r_start:
            para._element.append(make_run(reconstructed[pos:r_start], default_rPr))

    # 2. Original text — strikethrough
    old_run = make_run(old_text, default_rPr)
    old_rPr = old_run.find(qn("w:rPr"))
    if old_rPr is None:
        old_rPr = OxmlElement("w:rPr")
        old_run.insert(0, old_rPr)
    _add_strikethrough_to_rPr(old_rPr)
    para._element.append(old_run)

    # 3. Space separator
    sep_run = make_run(" ", default_rPr)
    para._element.append(sep_run)

    # 4. New text — green highlight
    new_run = make_run(new_text, default_rPr)
    new_rPr = new_run.find(qn("w:rPr"))
    if new_rPr is None:
        new_rPr = OxmlElement("w:rPr")
        new_run.insert(0, new_rPr)
    _add_highlight_to_rPr(new_rPr, "green")
    para._element.append(new_run)

    # 5. Unchanged text after — preserve per-run formatting where possible
    if after_text:
        pos = r_end
        added = False
        for tr in text_runs:
            tr_start = tr["offset"]
            tr_end = tr_start + len(tr["text"])
            if tr_start >= r_end:
                para._element.append(make_run(tr["text"], tr.get("rPr") or default_rPr))
                pos = tr_end
                added = True
            elif tr_end > r_end and tr_start < r_end:
                # Partial overlap — take the part after old_text
                skip = r_end - tr_start
                remaining = tr["text"][skip:]
                if remaining:
                    para._element.append(make_run(remaining, tr.get("rPr") or default_rPr))
                    pos = tr_end
                    added = True
        # If we missed some after text, add with default formatting
        if not added or pos < len(reconstructed):
            leftover = reconstructed[max(pos, r_end):]
            if leftover:
                para._element.append(make_run(leftover, default_rPr))


def _apply_multi_visual_diff_to_paragraph(para, matches):
    """Apply multiple visual diffs to a paragraph in a single pass.

    matches: list of (start_idx, old_text, new_text) sorted by start_idx, non-overlapping.

    Rebuilds the paragraph once: unchanged text keeps original formatting,
    old text gets strikethrough, new text gets green highlight.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from copy import deepcopy

    full_text = para.text or ""
    if not full_text or not matches:
        return

    # ── Collect per-run formatting info ──
    def _get_run_text(r_elem):
        return "".join((t.text or "") for t in r_elem.findall(qn("w:t")))

    def _get_elem_rPr(r_elem):
        rPr = r_elem.find(qn("w:rPr"))
        return deepcopy(rPr) if rPr is not None else None

    # Build flat list of (text, rPr, offset) from all text-bearing children
    text_runs = []
    char_offset = 0
    for child in list(para._element):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "r":
            t = _get_run_text(child)
            if t:
                text_runs.append({"text": t, "offset": char_offset, "rPr": _get_elem_rPr(child)})
                char_offset += len(t)
        elif tag == "hyperlink":
            for sub_r in child.findall(qn("w:r")):
                t = _get_run_text(sub_r)
                if t:
                    text_runs.append({"text": t, "offset": char_offset, "rPr": _get_elem_rPr(sub_r)})
                    char_offset += len(t)

    reconstructed = "".join(tr["text"] for tr in text_runs)

    # If reconstruction doesn't match para.text, fall back to para.text
    if not reconstructed:
        reconstructed = full_text

    # Get default formatting
    default_rPr = text_runs[0]["rPr"] if text_runs else None

    def make_run(text, rPr_source=None):
        r = OxmlElement("w:r")
        if rPr_source:
            r.append(deepcopy(rPr_source))
        else:
            r.append(OxmlElement("w:rPr"))
        t_elem = OxmlElement("w:t")
        t_elem.set(qn("xml:space"), "preserve")
        t_elem.text = text
        r.append(t_elem)
        return r

    def get_rPr_at(pos):
        """Get the run properties for the character at given position."""
        for tr in text_runs:
            if tr["offset"] <= pos < tr["offset"] + len(tr["text"]):
                return tr.get("rPr") or default_rPr
        return default_rPr

    # ── Remove all text-bearing children ──
    for child in list(para._element):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("r", "hyperlink", "smartTag", "fldSimple"):
            para._element.remove(child)

    # ── Rebuild paragraph in a single pass ──
    cursor = 0
    for start, old_text, new_text in matches:
        # Verify match position in reconstructed text
        actual_start = reconstructed.find(old_text, cursor)
        if actual_start == -1:
            continue
        actual_end = actual_start + len(old_text)

        # 1. Unchanged text before this match
        if actual_start > cursor:
            before = reconstructed[cursor:actual_start]
            para._element.append(make_run(before, get_rPr_at(cursor)))

        # 2. Old text — strikethrough only
        old_run = make_run(old_text, get_rPr_at(actual_start))
        old_rPr = old_run.find(qn("w:rPr"))
        if old_rPr is None:
            old_rPr = OxmlElement("w:rPr")
            old_run.insert(0, old_rPr)
        _add_strikethrough_to_rPr(old_rPr)
        para._element.append(old_run)

        # 3. Space separator
        para._element.append(make_run(" ", default_rPr))

        # 4. New text — green highlight
        new_run = make_run(new_text, default_rPr)
        new_rPr = new_run.find(qn("w:rPr"))
        if new_rPr is None:
            new_rPr = OxmlElement("w:rPr")
            new_run.insert(0, new_rPr)
        _add_highlight_to_rPr(new_rPr, "green")
        para._element.append(new_run)

        cursor = actual_end

    # 5. Remaining unchanged text after all matches
    if cursor < len(reconstructed):
        after = reconstructed[cursor:]
        para._element.append(make_run(after, get_rPr_at(cursor)))


# ══════════════════════════════════════════════════════════════════════════════
# STAGE 9: DATED STATEMENT AUDIT
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/{session_id}/run-audit")
async def run_audit(
    session_id: str,
    background_tasks: BackgroundTasks,
    user=Depends(get_current_user_dep),
):
    """Run dated-statement audit on in-scope sections."""
    db = get_database()
    session_repo = SessionRepository(db)
    doc_repo = DocumentRepository(db)

    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    doc = await doc_repo.find_by_id(session["document_id"], analysis_mode=True)
    if not doc:
        raise HTTPException(404, "Document not found")

    background_tasks.add_task(_run_audit_task, session_id, session, doc)

    return {
        "session_id": session_id,
        "status": "auditing",
        "message": "Dated statement audit started. Poll session for updates.",
    }


async def _run_audit_task(session_id: str, session: dict, doc: dict):
    """Background task: find temporal language in document."""
    from openai import AsyncOpenAI
    import json

    db = get_database()
    session_repo = SessionRepository(db)

    try:
        client = AsyncOpenAI(api_key=(await _get_openai_key(db)))
        outline = session.get("outline", [])
        text_content = doc.get("text_content", "")

        # Get in-scope text
        paragraphs = text_content.split("\n")
        in_scope = [o for o in outline if o.get("in_scope", True)]
        section_texts = []
        for i, heading in enumerate(in_scope):
            start = heading.get("paragraph_index", 0)
            end = in_scope[i + 1].get("paragraph_index", len(paragraphs)) if i + 1 < len(in_scope) else len(paragraphs)
            section_texts.append({
                "section": heading.get("text", ""),
                "text": "\n".join(paragraphs[start:end]),
            })

        combined = "\n\n".join(
            f"[Section: {s['section']}]\n{s['text']}" for s in section_texts
        )

        # Get existing patched sentences to exclude
        patches = await session_repo.find_patches(session_id)
        patched_sentences = set(p.get("original_sentence", "") for p in patches)

        prompt = f"""Find every sentence containing temporal or date-sensitive language in this text.
Look for: explicit dates ("2018", "March 2019"), relative temporal words ("currently", "recently", "now", "today"),
future tense ("will be", "is planned", "upcoming"), and planned status phrases ("as of", "planned for", "expected to").

For each found, return JSON with:
- "sentence": the exact sentence
- "section_ref": which section it's in
- "trigger_word": the specific word/phrase that is temporal
- "trigger_type": one of "explicit_date", "relative_temporal", "future_tense", "planned_status"
- "risk": "high", "medium", or "low"

Return ONLY a JSON array. No other text.

Text:
{combined[:15000]}"""

        response = await client.chat.completions.create(
            model=settings.GPT_MODEL,
            messages=[
                {"role": "system", "content": "You are a temporal language auditor. Find all date-sensitive statements that might be outdated. Never spell out Greek letters — preserve original symbols (Δ, α, β, θ, etc.) exactly as written. Return only valid JSON."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.2,
            max_tokens=4000,
        )

        raw = response.choices[0].message.content.strip()
        if raw.startswith("```"):
            raw = raw.split("\n", 1)[1].rsplit("```", 1)[0]
        statements = json.loads(raw)

        # Delete old statements
        await session_repo.delete_dated_statements(session_id)

        dated_items = []
        for s in statements:
            sentence = s.get("sentence", "")
            # Exclude already-patched sentences
            if sentence in patched_sentences:
                continue
            dated_items.append({
                "statement_id": str(uuid.uuid4())[:8],
                "session_id": session_id,
                "sentence": sentence,
                "trigger_word": s.get("trigger_word", ""),
                "trigger_type": s.get("trigger_type", "explicit_date"),
                "section_ref": s.get("section_ref", ""),
                "risk": s.get("risk", "medium"),
                "resolved": False,
                "resolution_note": None,
            })

        await session_repo.create_dated_statements(dated_items)
        # Status stays at EDITS_APPLIED; transitions to AUDIT_COMPLETE
        # only when the editor resolves all dated statements (in resolve_statement endpoint)

        logger.info("Audit complete for session %s: %d dated statements", session_id, len(dated_items))

    except Exception as e:
        logger.error("Audit failed for session %s: %s", session_id, e)
        error_str = str(e)
        if "insufficient_quota" in error_str or ("429" in error_str and "quota" in error_str.lower()):
            error_msg = (
                "OpenAI API quota exceeded. Your API key has no remaining credits. "
                "Please check your billing at platform.openai.com or update your API key in Admin → API Keys."
            )
        elif "invalid_api_key" in error_str or "401" in error_str:
            error_msg = (
                "OpenAI API key is invalid or expired. "
                "Please update your API key in Admin → API Keys."
            )
        else:
            error_msg = error_str
        await session_repo.update_session(session_id, {
            "status": SessionStatus.ERROR.value,
            "error_message": error_msg,
        })


@router.get("/{session_id}/dated-statements")
async def get_dated_statements(session_id: str, user=Depends(get_current_user_dep)):
    """Get all dated statements for review."""
    db = get_database()
    session_repo = SessionRepository(db)
    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    statements = await session_repo.find_dated_statements(session_id)
    return {
        "session_id": session_id,
        "statements": statements,
        "total": len(statements),
        "resolved": sum(1 for s in statements if s.get("resolved")),
        "unresolved": sum(1 for s in statements if not s.get("resolved")),
    }


@router.put("/{session_id}/dated-statements/{statement_id}")
async def resolve_statement(
    session_id: str,
    statement_id: str,
    req: DatedStatementResolveRequest,
    user=Depends(get_current_user_dep),
):
    """Resolve a dated statement."""
    db = get_database()
    session_repo = SessionRepository(db)

    if req.resolution not in ("still_current", "flag_for_patch", "acceptable"):
        raise HTTPException(400, "Invalid resolution. Use 'still_current', 'flag_for_patch', or 'acceptable'.")

    updated = await session_repo.resolve_dated_statement(statement_id, req.resolution)

    # Check if all resolved
    statements = await session_repo.find_dated_statements(session_id)
    all_resolved = all(s.get("resolved") for s in statements)

    if all_resolved and statements:
        await session_repo.update_session(session_id, {
            "status": SessionStatus.AUDIT_COMPLETE.value,
        })

    return {
        "statement_id": statement_id,
        "resolution": req.resolution,
        "updated": updated,
        "all_resolved": all_resolved,
    }


# ══════════════════════════════════════════════════════════════════════════════
# STAGE 10: EXPORT
# ══════════════════════════════════════════════════════════════════════════════

@router.get("/{session_id}/export/tracked-docx")
async def export_tracked_docx(session_id: str, token: Optional[str] = None, request: Request = None):
    """Export tracked-changes DOCX with visual diff (red/green/yellow highlights)."""
    from fastapi.responses import FileResponse

    # Auth: accept either query param token (direct browser download) or Authorization header
    user = None
    if token:
        user = verify_download_token(token)
    else:
        # Try Authorization header
        auth_header = request.headers.get("authorization", "") if request else ""
        if auth_header.startswith("Bearer "):
            jwt_token = auth_header[7:]
            payload = decode_token(jwt_token)
            user = {"email": payload.get("sub"), "role": payload.get("role", "user"), "token": jwt_token}
    if not user:
        raise HTTPException(401, "Authentication required — pass token query param or Authorization header")

    db = get_database()
    session_repo = SessionRepository(db)
    doc_repo = DocumentRepository(db)

    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    # Verify session has reached at least edits_applied stage
    allowed_statuses = (
        SessionStatus.EDITS_APPLIED.value,
        SessionStatus.AUDIT_COMPLETE.value,
        SessionStatus.EXPORTED.value,
    )
    if session.get("status") not in allowed_statuses:
        raise HTTPException(400, "Patches must be applied before exporting. Complete stages 1-8 first.")

    doc = await doc_repo.find_by_id(session["document_id"], lightweight=True)
    if not doc:
        raise HTTPException(404, "Document not found")

    original_path = _resolve_file_path(doc.get("file_path", ""))
    working_path = session.get("working_doc_path", "")

    if not working_path or not os.path.exists(working_path):
        raise HTTPException(
            400,
            "Working document not available. The original file may be on a different server. "
            "Use the Changelog export instead, which contains all patch details.",
        )

    output_dir = os.path.join(settings.OUTPUT_DIR, session_id)
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"tracked_{doc.get('original_filename', 'document.docx')}")

    logger.info("EXPORT tracked-docx: building tracked changes from patches")

    # Build tracked changes DOCX using approved patches
    approved_patches = await session_repo.find_approved_patches(session_id)
    try:
        _build_tracked_changes_docx(original_path, approved_patches, output_path,
                                    outline=session.get("outline", []))
        logger.info("EXPORT: tracked changes DOCX built at %s", output_path)
    except Exception as e:
        logger.error("EXPORT: tracked changes build failed: %s", e)
        # Fallback: copy working doc
        shutil.copy2(working_path, output_path)

    # ── Apply approved media patches (figures, equations, tables) to tracked DOCX ──
    try:
        db = get_database()
        approved_media = await db.media_patches.find({
            "session_id": session_id,
            "status": "approved",
        }).to_list(length=500)

        if approved_media:
            from docx import Document as DocxDocument
            tracked_doc = DocxDocument(output_path)
            media_applied = 0

            # Figures: replace images + update captions + add source URLs
            approved_figures = [m for m in approved_media if m.get("type") == "figure"]
            if approved_figures:
                fig_count = await _apply_figure_replacements_to_docx(tracked_doc, approved_figures)
                media_applied += fig_count
                logger.info("EXPORT tracked-docx: replaced %d figures", fig_count)
                cap_count = _update_figure_captions_in_docx(tracked_doc, approved_figures)
                if cap_count:
                    logger.info("EXPORT tracked-docx: updated %d figure captions", cap_count)
                src_count = _add_figure_source_urls(tracked_doc, approved_figures)
                if src_count:
                    logger.info("EXPORT tracked-docx: added %d figure source URLs", src_count)

            # Equations
            approved_equations = [m for m in approved_media if m.get("type") == "equation"]
            if approved_equations:
                eq_count = _apply_equation_replacements_to_docx(tracked_doc, approved_equations)
                media_applied += eq_count
                logger.info("EXPORT tracked-docx: replaced %d equations", eq_count)

            # Tables: update cells + captions
            approved_tables = [m for m in approved_media if m.get("type") == "table"]
            if approved_tables:
                tbl_count = _apply_table_updates_to_docx(tracked_doc, approved_tables, highlight=True)
                media_applied += tbl_count
                logger.info("EXPORT tracked-docx: updated %d tables (highlighted)", tbl_count)
                tcap_count = _update_table_captions_in_docx(tracked_doc, approved_tables)
                if tcap_count:
                    logger.info("EXPORT tracked-docx: updated %d table captions", tcap_count)

            # Headers/footers: preserve from original
            try:
                original_doc = DocxDocument(original_path)
                _preserve_headers_footers(original_doc, tracked_doc)
            except Exception as hf_err:
                logger.debug("Headers/footers preservation skipped: %s", hf_err)

            if media_applied > 0:
                tracked_doc.save(output_path)
                logger.info("EXPORT tracked-docx: saved with %d total media patches", media_applied)
    except Exception as media_err:
        logger.warning("EXPORT tracked-docx: media patch application failed (non-fatal): %s", media_err)

    output_size = os.path.getsize(output_path) if os.path.exists(output_path) else 0
    logger.info("EXPORT: serving %s (%d bytes)", output_path, output_size)

    await session_repo.update_session(session_id, {
        "status": SessionStatus.EXPORTED.value,
    })

    return FileResponse(
        output_path,
        filename=f"tracked_{doc.get('original_filename', 'document.docx')}",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
            "Pragma": "no-cache",
        },
    )


@router.get("/{session_id}/export/clean-docx")
async def export_clean_docx(session_id: str, token: Optional[str] = None, request: Request = None):
    """Export clean DOCX with all changes accepted."""
    from fastapi.responses import FileResponse

    # Auth: accept either query param token or Authorization header
    user = None
    if token:
        user = verify_download_token(token)
    else:
        auth_header = request.headers.get("authorization", "") if request else ""
        if auth_header.startswith("Bearer "):
            jwt_token = auth_header[7:]
            payload = decode_token(jwt_token)
            user = {"email": payload.get("sub"), "role": payload.get("role", "user"), "token": jwt_token}
    if not user:
        raise HTTPException(401, "Authentication required")

    db = get_database()
    session_repo = SessionRepository(db)
    doc_repo = DocumentRepository(db)

    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    doc = await doc_repo.find_by_id(session["document_id"], lightweight=True)
    working_path = session.get("working_doc_path", "")

    if not working_path or not os.path.exists(working_path):
        raise HTTPException(
            400,
            "Working document not available. The original file may be on a different server. "
            "Use the Changelog export instead, which contains all patch details.",
        )

    logger.info("EXPORT clean-docx: serving working_path=%s (%d bytes)",
                working_path, os.path.getsize(working_path) if os.path.exists(working_path) else 0)

    return FileResponse(
        working_path,
        filename=f"clean_{doc.get('original_filename', 'document.docx')}",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
            "Pragma": "no-cache",
        },
    )


@router.get("/{session_id}/export/changelog")
async def export_changelog(session_id: str, user=Depends(get_current_user_dep)):
    """Export changelog as JSON with all patches, evidence, and decisions."""
    db = get_database()
    session_repo = SessionRepository(db)

    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    patches = await session_repo.find_patches(session_id)
    dated_statements = await session_repo.find_dated_statements(session_id)
    opportunities = await session_repo.find_opportunities(session_id)

    # Include media patches (figures, equations, tables)
    media_patches_cursor = db.media_patches.find(
        {"session_id": session_id},
        {"original_image_b64": 0},  # exclude large base64 blobs from export
    )
    media_patches = await media_patches_cursor.to_list(length=1000)
    for mp in media_patches:
        mp["_id"] = str(mp["_id"])

    return {
        "session_id": session_id,
        "document_id": session.get("document_id"),
        "rules": session.get("rules"),
        "diagnostic": session.get("diagnostic"),
        "total_opportunities": len(opportunities),
        "patches": patches,
        "dated_statements": dated_statements,
        "media_patches": media_patches,
        "exported_at": datetime.utcnow().isoformat(),
    }


# ── Media Analysis Endpoints ─────────────────────────────────────────────────


@router.get("/{session_id}/figure-analysis")
async def figure_analysis(session_id: str, user=Depends(get_current_user_dep)):
    """Analyze all figures/images in the DOCX. Uses GPT-4o Vision to assess
    whether each figure is outdated and searches NASA / Wikimedia for replacements."""
    import httpx
    from uuid import uuid4

    db = get_database()
    session_repo = SessionRepository(db)
    doc_repo = DocumentRepository(db)

    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    # Return cached results if they exist
    existing = await db.media_patches.find(
        {"session_id": session_id, "type": "figure"}
    ).to_list(length=500)
    if existing:
        for e in existing:
            e["_id"] = str(e["_id"])
        return {"session_id": session_id, "figures": existing, "cached": True}

    # ── Read figures from DB (extracted once during POST /process) ──
    doc = await doc_repo.find_with_media(session["document_id"])
    if not doc:
        raise HTTPException(404, "Document not found")

    db_figures = doc.get("figures", [])

    # Filter decorative/small images using stored metadata
    figures = []
    for fig in db_figures:
        size_bytes = fig.get("size_bytes") or 0
        cx = fig.get("cx") or 0
        cy = fig.get("cy") or 0
        w_px = cx / 914400 * 96 if cx else 0
        h_px = cy / 914400 * 96 if cy else 0

        # Skip small images (< 5KB)
        if size_bytes > 0 and size_bytes < 5000:
            logger.debug("Skipping small image fig_%s (%d bytes)", fig.get("figure_id"), size_bytes)
            continue
        # Skip tiny dimensions
        if w_px > 0 and (w_px < 50 or h_px < 50):
            logger.debug("Skipping tiny image fig_%s (%dx%d px)", fig.get("figure_id"), w_px, h_px)
            continue
        # Skip decorative bars (wide + short)
        if w_px > 500 and h_px < 30:
            logger.debug("Skipping decorative bar fig_%s (%dx%d px)", fig.get("figure_id"), w_px, h_px)
            continue

        pos = fig.get("position") or {}
        figures.append({
            "para_idx": pos.get("paragraph"),
            "image_b64": fig.get("image_base64", ""),
            "image_url": fig.get("image_url", ""),
            "caption": fig.get("caption") or "",
            "r_embed": fig.get("r_embed") or "",
            "size_bytes": size_bytes,
            "cx": cx,
            "cy": cy,
        })

    logger.info("Found %d content figures from DB (filtered from %d total)", len(figures), len(db_figures))

    if not figures:
        return {"session_id": session_id, "figures": [], "message": "No figures found in document."}

    # ── Analyze figures in PARALLEL for speed ──────────────────────────────
    import asyncio
    import json as _json
    from openai import AsyncOpenAI

    client = AsyncOpenAI(api_key=(await _get_openai_key(db)))

    # Resolve Tavily API key: check MongoDB first, fall back to .env
    _fig_tavily_key = settings.TAVILY_API_KEY
    _fig_db_key_doc = await db.settings.find_one({"key": "tavily_api_key"})
    if _fig_db_key_doc and _fig_db_key_doc.get("value"):
        _fig_tavily_key = _fig_db_key_doc["value"]

    async def _analyze_single_figure(fig_num: int, fig: dict, http_client: httpx.AsyncClient):
        """Analyze one figure with GPT-4o Vision + search for replacements concurrently."""
        # Use Cloudinary URL directly when available (no base64 needed).
        # Otherwise fall back to compressed base64.
        _cloudinary_url = fig.get("image_url", "")
        thumb_b64 = ""
        if not _cloudinary_url:
            try:
                _raw = base64.b64decode(fig["image_b64"])
                _compressed = ImageService.compress_for_storage(_raw)
                thumb_b64 = base64.b64encode(_compressed).decode("utf-8")
            except Exception:
                thumb_b64 = fig["image_b64"]
        caption = fig.get("caption", "") or ""

        # ── Step 1: GPT-4o Vision analysis ────────────────────────────────
        analysis_text = ""
        search_queries = []
        is_outdated = False
        figure_category = "unknown"
        try:
            caption_context = f'\nThe figure caption in the textbook is: "{caption}"' if caption else ""

            # Build the image_url payload for GPT-4o Vision
            if _cloudinary_url:
                _vision_image_url = {"url": _cloudinary_url}
            else:
                # Detect actual image MIME type from base64 data
                _mime = "image/png"  # default
                try:
                    _header = base64.b64decode(thumb_b64[:32])
                    if _header[:3] == b'\xff\xd8\xff':
                        _mime = "image/jpeg"
                    elif _header[:4] == b'\x89PNG':
                        _mime = "image/png"
                    elif _header[:4] == b'GIF8':
                        _mime = "image/gif"
                    elif _header[:4] == b'RIFF':
                        _mime = "image/webp"
                except Exception:
                    pass
                _vision_image_url = {"url": f"data:{_mime};base64,{thumb_b64}"}

            _vision_messages = [{
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "Analyze this figure from a textbook."
                            f"{caption_context}\n\n"
                            "Step 1 — Classify the figure into one of these categories:\n"
                            "  - \"historical_illustration\": artwork, manuscripts, or illustrations "
                            "depicting historical events/models (e.g., Aristotle's universe, medieval diagrams). "
                            "These are NEVER outdated — they are meant to show how things were understood in the past.\n"
                            "  - \"portrait\": photograph or painting of a specific person (scientist, astronaut, etc.). "
                            "Portraits are NEVER outdated — they depict a specific individual.\n"
                            "  - \"technical_diagram\": engineering diagrams, system architectures, process flows, "
                            "block diagrams. These CAN become outdated if the technology or process has changed.\n"
                            "  - \"data_chart\": graphs, charts, tables with numerical data. "
                            "These CAN become outdated if newer data exists.\n"
                            "  - \"photograph\": real-world photos of hardware, facilities, launches. "
                            "These CAN become outdated if the subject has changed significantly.\n"
                            "  - \"conceptual\": generic conceptual illustrations, decorative images.\n\n"
                            "Step 2 — Determine if the figure is outdated:\n"
                            "  - historical_illustration → ALWAYS set is_outdated = false\n"
                            "  - portrait → ALWAYS set is_outdated = false\n"
                            "  - For other categories, assess whether the content shown is still current.\n\n"
                            "Step 3 — If and ONLY if is_outdated is true, suggest exactly 3 different "
                            "image search queries to find an UPDATED replacement. Each query should:\n"
                            "  - Be 3-7 words, specific to the EXACT content shown\n"
                            "  - Target different search angles (e.g., technical term, industry standard name, textbook diagram name)\n"
                            "  - NEVER use the word 'updated' or 'new' — just describe what the replacement should show\n"
                            "  - Include specific domain terms (e.g., 'SMAD space mission architecture elements' not 'space mission diagram')\n"
                            "  - For data charts: describe the axes/variables (e.g., 'spacecraft design life vs cost tradeoff curve')\n"
                            "  - For technical diagrams: use the standard name of the framework/model shown\n\n"
                            "Respond in JSON with keys: "
                            "\"analysis\" (string), \"figure_category\" (string), "
                            "\"is_outdated\" (bool), \"search_queries\" (list of 3 strings, empty list if not outdated)."
                        ),
                    },
                    {
                        "type": "image_url",
                        "image_url": _vision_image_url,
                    },
                ],
            }]

            # Retry up to 3 times on 429 rate-limit errors
            gpt_resp = None
            for _attempt in range(3):
                try:
                    gpt_resp = await client.chat.completions.create(
                        model="gpt-4o",
                        messages=_vision_messages,
                        temperature=0.0,
                        max_tokens=1000,
                    )
                    break  # success
                except Exception as _rate_err:
                    _err_str = str(_rate_err)
                    if "429" in _err_str or "rate_limit" in _err_str.lower():
                        _wait = 1.0 * (_attempt + 1)  # 1s, 2s, 3s
                        logger.info("Rate limit hit for figure %d, retrying in %.1fs (attempt %d/3)",
                                   fig_num, _wait, _attempt + 1)
                        await asyncio.sleep(_wait)
                    else:
                        raise  # non-rate-limit error, propagate

            if gpt_resp is None:
                raise Exception(f"Rate limit exceeded after 3 retries for figure {fig_num}")

            raw = gpt_resp.choices[0].message.content or ""
            analysis_text = raw
            try:
                cleaned = raw.strip()
                if cleaned.startswith("```"):
                    cleaned = cleaned.split("\n", 1)[1] if "\n" in cleaned else cleaned
                    cleaned = cleaned.rsplit("```", 1)[0]
                parsed = _json.loads(cleaned)
                analysis_text = parsed.get("analysis", raw)
                # Support both old single query and new multi-query format
                search_queries = parsed.get("search_queries", [])
                if not search_queries:
                    sq = parsed.get("search_query", "")
                    search_queries = [sq] if sq else []
                is_outdated = parsed.get("is_outdated", False)
                figure_category = parsed.get("figure_category", "unknown")
            except (_json.JSONDecodeError, Exception):
                search_queries = [caption] if caption else []
                is_outdated = False
        except Exception as gpt_err:
            logger.warning("GPT Vision analysis failed for figure %d: %s", fig_num, gpt_err)
            analysis_text = f"Analysis unavailable: {gpt_err}"
            search_queries = [caption] if caption else []
            is_outdated = False

        logger.info("Figure %d: category=%s, is_outdated=%s, queries=%s",
                     fig_num, figure_category, is_outdated, search_queries)

        # ── Step 2: Search for replacements (only if outdated) ────────────
        all_candidates = []

        if not is_outdated:
            logger.info("Figure %d is not outdated (category=%s) — skipping replacement search",
                        fig_num, figure_category)
            return {
                "patch_id": str(uuid4()),
                "session_id": session_id,
                "type": "figure",
                "status": "not_outdated",
                "figure_number": fig_num,
                "caption": caption,
                "original_image_b64": thumb_b64 if not _cloudinary_url else "",
                "original_image_url": _cloudinary_url,
                "analysis": analysis_text,
                "figure_category": figure_category,
                "is_outdated": False,
                "replacement_candidates": [],
                "selected_replacement": None,
                "r_embed": fig.get("r_embed", ""),
                "para_idx": fig.get("para_idx"),
                "original_cx": fig.get("cx", 0),
                "original_cy": fig.get("cy", 0),
                "created_at": datetime.utcnow().isoformat(),
            }

        logger.info("Figure %d is outdated — searching with %d queries: %s",
                     fig_num, len(search_queries), search_queries)

        # ── Primary: Tavily web image search (best for technical diagrams) ──
        async def _search_tavily(query: str):
            candidates = []
            if not _fig_tavily_key:
                return candidates
            try:
                tavily_resp = await http_client.post(
                    "https://api.tavily.com/search",
                    json={
                        "api_key": _fig_tavily_key,
                        "query": f"{query} diagram",
                        "search_depth": "advanced",
                        "include_images": True,
                        "include_image_descriptions": True,
                        "max_results": 5,
                    },
                    timeout=20,
                )
                tavily_data = tavily_resp.json()
                seen_urls = set()

                # Log Tavily response for debugging
                img_count = len(tavily_data.get("images", []))
                res_count = len(tavily_data.get("results", []))
                logger.info("Tavily response for '%s': %d images, %d results, status=%d",
                           query, img_count, res_count, tavily_resp.status_code)
                if tavily_resp.status_code != 200:
                    logger.warning("Tavily non-200 response: %s", tavily_data)

                # Build domain → page URL map from results for source attribution
                from urllib.parse import urlparse
                result_pages = []  # list of (domain_base, full_page_url, title)
                for res in tavily_data.get("results", []):
                    page_url = res.get("url", "")
                    title = res.get("title", "")
                    if page_url:
                        parsed_page = urlparse(page_url)
                        # Extract base domain (remove 'www.', 'cdn.', 'images.' etc.)
                        domain = parsed_page.netloc.lower()
                        domain_base = '.'.join(domain.replace('www.', '').split('.')[-2:])
                        result_pages.append((domain_base, page_url, title))
                        logger.debug("Tavily result page: domain=%s url=%s", domain_base, page_url[:100])

                if not result_pages:
                    logger.warning("Tavily returned 0 result pages for query '%s' — no source URLs available", query)

                # Helper: check if a URL is a raw image file URL (not a webpage)
                _IMAGE_EXTENSIONS = ('.jpg', '.jpeg', '.png', '.gif', '.webp', '.svg', '.bmp', '.tiff', '.ico')
                def _is_image_url(url: str) -> bool:
                    """Return True if url points to an image file, not a webpage."""
                    if not url:
                        return False
                    # Parse and check path (ignore query params)
                    path = urlparse(url).path.lower()
                    return any(path.endswith(ext) for ext in _IMAGE_EXTENSIONS)

                for i, img_item in enumerate(tavily_data.get("images", [])[:5]):
                    # Tavily images can be strings or dicts with url+description
                    if isinstance(img_item, dict):
                        img_url = img_item.get("url", "")
                    elif isinstance(img_item, str):
                        img_url = img_item
                    else:
                        continue

                    if not img_url or not img_url.startswith("http") or img_url in seen_urls:
                        continue
                    seen_urls.add(img_url)

                    # Match image URL to its source page by domain similarity
                    parsed_img = urlparse(img_url)
                    img_domain = parsed_img.netloc.lower()
                    # Strip ALL common CDN/media subdomains: image3., cdn., i0.wp., etc.
                    import re as _re_domain
                    img_domain_clean = _re_domain.sub(
                        r'^(www\.|cdn\.|images?\d*\.|static\.|media\.|assets\.|i\d+\.wp\.)',
                        '', img_domain
                    )
                    img_domain_base = '.'.join(img_domain_clean.split('.')[-2:])

                    source_page = ""
                    source_title = query
                    for domain_base, page_url, title in result_pages:
                        # Only accept page URLs that are actual webpages, NOT image files
                        if _is_image_url(page_url):
                            continue
                        if img_domain_base == domain_base or img_domain_base in domain_base or domain_base in img_domain_base:
                            source_page = page_url
                            if title:
                                source_title = title
                            break

                    # If no domain match, use the FIRST non-image result page URL
                    if not source_page:
                        for _, pg_url, pg_title in result_pages:
                            if not _is_image_url(pg_url):
                                source_page = pg_url
                                if pg_title:
                                    source_title = pg_title
                                break

                    # FINAL SAFETY CHECK: Never store an image URL as source_page_url
                    if _is_image_url(source_page):
                        logger.warning("Rejected image URL as source_page: %s", source_page[:100])
                        source_page = ""

                    candidates.append({
                        "url": img_url,
                        "title": source_title,
                        "source": "Web Search",
                        "thumbnail_url": img_url,
                        "source_page_url": source_page,
                    })
                    logger.info("Tavily candidate: image=%s → source_page=%s",
                                img_url[:80], source_page[:80] if source_page else "(none)")
            except Exception as e:
                logger.warning("Tavily search failed for figure %d query '%s': %s", fig_num, query, e)
            return candidates

        # ── Secondary: NASA Images (only for space/aerospace content) ───────
        space_keywords = {"space", "mission", "satellite", "orbit", "launch", "spacecraft",
                          "nasa", "esa", "rocket", "constellation", "iss", "lunar", "mars"}
        all_query_words = {w.lower() for q in search_queries for w in q.split()}
        is_space_related = bool(all_query_words & space_keywords)

        async def _search_nasa(query: str):
            candidates = []
            if not is_space_related:
                logger.debug("NASA search skipped (not space-related) for figure %d", fig_num)
                return candidates
            try:
                params = {"q": query, "media_type": "image", "page_size": 5}
                if settings.NASA_API_KEY:
                    params["api_key"] = settings.NASA_API_KEY
                resp = await http_client.get(
                    "https://images-api.nasa.gov/search", params=params, timeout=15
                )
                nasa_data = resp.json()
                items = nasa_data.get("collection", {}).get("items", [])
                logger.info("NASA response for '%s': %d items, status=%d",
                           query, len(items), resp.status_code)
                for item in items[:3]:
                    data = item.get("data", [{}])[0]
                    links = item.get("links", [{}])
                    thumb = links[0].get("href", "") if links else ""
                    if thumb:
                        # Derive full-resolution URL from thumbnail
                        # NASA thumbnails: .../image~thumb.jpg → .../image~orig.jpg
                        full_url = thumb.replace("~thumb", "~orig").replace("~small", "~orig").replace("~medium", "~orig")
                        # Construct NASA source page URL from nasa_id
                        nasa_id = data.get("nasa_id", "")
                        source_page = f"https://images.nasa.gov/details/{nasa_id}" if nasa_id else ""
                        candidates.append({
                            "url": full_url,
                            "title": data.get("title", ""),
                            "source": "NASA",
                            "thumbnail_url": thumb,
                            "source_page_url": source_page,
                        })
            except Exception as e:
                logger.warning("NASA search failed for figure %d: %s", fig_num, e)
            return candidates

        # ── Secondary: Wikimedia Commons ────────────────────────────────────
        async def _search_wikimedia(query: str):
            candidates = []
            wiki_headers = {"User-Agent": "AIBookUpdater/1.0 (educational; contact@example.com)"}
            try:
                wiki_params = {
                    "action": "query", "format": "json",
                    "generator": "search", "gsrsearch": query,
                    "gsrnamespace": "6", "gsrlimit": "5",
                    "prop": "imageinfo", "iiprop": "url|thumburl|extmetadata|descriptionurl", "iiurlwidth": "1200",
                }
                resp = await http_client.get(
                    "https://commons.wikimedia.org/w/api.php",
                    params=wiki_params, headers=wiki_headers, timeout=15,
                )
                pages = resp.json().get("query", {}).get("pages", {})
                for page in pages.values():
                    ii = page.get("imageinfo", [])
                    if ii:
                        info = ii[0]
                        img_url = info.get("thumburl") or info.get("url", "")
                        if img_url:
                            page_title = page.get("title", "")
                            # Construct Wikimedia Commons page URL
                            source_page = info.get("descriptionurl") or ""
                            if not source_page and page_title:
                                from urllib.parse import quote
                                source_page = f"https://commons.wikimedia.org/wiki/{quote(page_title)}"
                            candidates.append({
                                "url": info.get("url", img_url),
                                "title": page_title.replace("File:", ""),
                                "source": "Wikimedia Commons",
                                "thumbnail_url": img_url,
                                "source_page_url": source_page,
                            })
            except Exception as e:
                logger.warning("Wikimedia search failed for figure %d: %s", fig_num, e)
            return candidates

        # ── Run all queries across all sources in parallel ──────────────────
        search_tasks = []
        for q in search_queries:
            search_tasks.append(_search_tavily(q))
            search_tasks.append(_search_nasa(q))
            search_tasks.append(_search_wikimedia(q))

        search_results = await asyncio.gather(*search_tasks, return_exceptions=True)
        for res in search_results:
            if isinstance(res, list):
                all_candidates.extend(res)

        # ── Deduplicate by URL ──────────────────────────────────────────────
        seen_urls = set()
        unique_candidates = []
        for c in all_candidates:
            url = c.get("thumbnail_url") or c.get("url", "")
            if url and url not in seen_urls:
                seen_urls.add(url)
                unique_candidates.append(c)
        all_candidates = unique_candidates

        # ── GPT-4o relevance filter: keep only candidates that match ────────
        if all_candidates and len(all_candidates) > 2:
            try:
                candidate_descriptions = "\n".join(
                    f"{i+1}. title=\"{c.get('title', '')}\" source={c.get('source', '')} url={c.get('url', '')}"
                    for i, c in enumerate(all_candidates[:10])
                )
                filter_resp = await client.chat.completions.create(
                    model="gpt-4o",
                    messages=[{
                        "role": "user",
                        "content": (
                            f"I need a replacement image for this textbook figure:\n"
                            f"  Caption: \"{caption}\"\n"
                            f"  AI Analysis: \"{analysis_text}\"\n"
                            f"  Category: {figure_category}\n\n"
                            f"Here are candidate replacement images found via search:\n"
                            f"{candidate_descriptions}\n\n"
                            f"Return a JSON object with key \"relevant_indices\" — a list of the "
                            f"1-based indices of candidates that are RELEVANT replacements for "
                            f"this specific figure. Only include candidates whose title/source "
                            f"suggests they show the SAME type of content (same topic, same kind "
                            f"of diagram/chart). Exclude anything generic, unrelated, or from a "
                            f"completely different domain.\n"
                            f"Example: {{\"relevant_indices\": [1, 3, 5]}}"
                        ),
                    }],
                    temperature=0.0,
                    max_tokens=200,
                )
                filter_raw = filter_resp.choices[0].message.content or ""
                filter_cleaned = filter_raw.strip()
                if filter_cleaned.startswith("```"):
                    filter_cleaned = filter_cleaned.split("\n", 1)[1] if "\n" in filter_cleaned else filter_cleaned
                    filter_cleaned = filter_cleaned.rsplit("```", 1)[0]
                filter_parsed = _json.loads(filter_cleaned)
                relevant_indices = filter_parsed.get("relevant_indices", [])
                if relevant_indices:
                    filtered = [all_candidates[i - 1] for i in relevant_indices
                                if 1 <= i <= len(all_candidates)]
                    if filtered:
                        logger.info("Figure %d: relevance filter kept %d/%d candidates",
                                    fig_num, len(filtered), len(all_candidates))
                        all_candidates = filtered
            except Exception as filter_err:
                logger.warning("Figure %d: relevance filter failed, using all candidates: %s",
                               fig_num, filter_err)

        logger.info("Figure %d: found %d replacement candidates", fig_num, len(all_candidates))
        replacement_candidates = all_candidates[:6]

        return {
            "patch_id": str(uuid4()),
            "session_id": session_id,
            "type": "figure",
            "status": "pending",
            "figure_number": fig_num,
            "caption": caption,
            "original_image_b64": thumb_b64 if not _cloudinary_url else "",
            "original_image_url": _cloudinary_url,
            "analysis": analysis_text,
            "figure_category": figure_category,
            "is_outdated": True,
            "replacement_candidates": replacement_candidates,
            "selected_replacement": None,
            "r_embed": fig.get("r_embed", ""),
            "para_idx": fig.get("para_idx"),
            "original_cx": fig.get("cx", 0),
            "original_cy": fig.get("cy", 0),
            "created_at": datetime.utcnow().isoformat(),
        }

    # Process all figures in parallel (batches of 3 to avoid rate limits)
    result_patches = []
    BATCH_SIZE = 3
    async with httpx.AsyncClient() as shared_http:
        for batch_start in range(0, len(figures), BATCH_SIZE):
            batch = figures[batch_start:batch_start + BATCH_SIZE]
            tasks = [
                _analyze_single_figure(batch_start + i + 1, fig, shared_http)
                for i, fig in enumerate(batch)
            ]
            batch_results = await asyncio.gather(*tasks, return_exceptions=True)
            for res in batch_results:
                if isinstance(res, Exception):
                    logger.error("Figure analysis failed: %s", res)
                else:
                    result_patches.append(res)

    # Store in MongoDB
    if result_patches:
        await db.media_patches.insert_many(result_patches)

    # Sanitise _id for JSON response
    for rp in result_patches:
        if "_id" in rp:
            rp["_id"] = str(rp["_id"])

    return {"session_id": session_id, "figures": result_patches, "cached": False}


@router.get("/{session_id}/equation-analysis")
async def equation_analysis(session_id: str, user=Depends(get_current_user_dep)):
    """Analyze all OMML equations in the DOCX for outdated notation."""
    from uuid import uuid4

    db = get_database()
    session_repo = SessionRepository(db)
    doc_repo = DocumentRepository(db)

    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    # Return cached results if they exist
    existing = await db.media_patches.find(
        {"session_id": session_id, "type": "equation"}
    ).to_list(length=500)
    if existing:
        for e in existing:
            e["_id"] = str(e["_id"])
        return {"session_id": session_id, "equations": existing, "cached": True}

    # ── Read equations from DB (extracted once during POST /process) ──
    doc = await doc_repo.find_with_media(session["document_id"])
    if not doc:
        raise HTTPException(404, "Document not found")

    db_equations = doc.get("equations", [])

    equations = []
    for eq in db_equations:
        pos = eq.get("position") or {}
        readable = eq.get("latex") or ""
        omml_xml = eq.get("raw_omml") or ""

        equations.append({
            "para_idx": pos.get("paragraph"),
            "omml_xml": omml_xml,
            "readable_text": readable,
        })

    # ── Merge continuation lines, classify, filter & deduplicate ────────
    def _merge_classify_filter(eqs):
        """Full equation pipeline:
        1. Merge continuation lines (=...) into the previous equation
        2. Classify: equation / definition / invalid
        3. Remove junk and pure-text fragments
        4. Deduplicate identical equations
        """
        import re as _mre

        # ── helpers ──────────────────────────────────────────────────
        def _to_plain(latex):
            """Convert LaTeX to plain readable text for analysis.
            Strips ALL LaTeX commands, braces, wrappers → raw characters."""
            s = latex
            # Unwrap \text{}, \mathrm{}, \mathbf{}, \operatorname{}
            for _ in range(3):
                s = _mre.sub(r'\\(?:text|mathrm|mathbf|operatorname)\{([^}]*)\}', r'\1', s)
            # Remove \begin{...}{...} and \end{...}
            s = _mre.sub(r'\\begin\{[^}]+\}(?:\{[^}]*\})?', '', s)
            s = _mre.sub(r'\\end\{[^}]+\}', '', s)
            # Replace known LaTeX commands with readable equivalents
            s = s.replace(r'\sqrt', '√').replace(r'\frac', '/')
            s = s.replace(r'\times', '×').replace(r'\cdot', '·')
            s = s.replace(r'\left', '').replace(r'\right', '')
            s = s.replace(r'\pm', '±').replace(r'\mp', '∓')
            s = s.replace(r'\infty', '∞').replace(r'\partial', '∂')
            s = s.replace(r'\leq', '≤').replace(r'\geq', '≥')
            s = s.replace(r'\neq', '≠').replace(r'\approx', '≈')
            # Strip remaining \command sequences
            s = _mre.sub(r'\\[a-zA-Z]+', '', s)
            # Strip \; \, \: \! and other single-char escapes
            s = _mre.sub(r'\\[^a-zA-Z\s]', '', s)
            # Strip braces and excess whitespace
            s = s.replace('{', '').replace('}', '')
            # Strip LaTeX alignment character & (from \begin{aligned} environments)
            s = s.replace('&', '')
            s = _mre.sub(r'\s+', ' ', s).strip()
            return s

        def _is_continuation(latex):
            """Check if this equation is a continuation line (starts with =).
            Uses plain-text conversion to handle all LaTeX wrappings."""
            plain = _to_plain(latex)
            # After stripping everything, does it start with = ≈ ≡ ~ ?
            if plain and plain[0] in '=≈≡~':
                return True
            # Also check for \approx, \equiv at the raw LaTeX start
            # Strip leading &, whitespace, braces (alignment chars from aligned environments)
            stripped = _mre.sub(r'^[\s{}&]+', '', latex)
            if _mre.match(r'^\\(?:approx|equiv|sim)\b', stripped):
                return True
            return False

        def _is_junk(latex):
            """Filter out non-equation junk. CONSERVATIVE — only remove
            things that are clearly NOT equations. Equations with long
            variable names like 'V_transfer at Earth' must NOT be removed."""
            plain = _to_plain(latex)
            if not plain or len(plain) < 2:
                return True
            # "as above" placeholder
            if _mre.search(r'as\s+above', plain, _mre.IGNORECASE):
                return True
            # Has ANY math operator or structure? Then it's NOT junk.
            has_operator = bool(_mre.search(
                r'[=<>+\-*/^√×·±∓≤≥≠≈≡∞∂∑∫∏()]|\\|[0-9]',
                plain
            ))
            if has_operator:
                return False
            # No operators/numbers at all. Could be a Greek var name or pure text.
            # If it's short (< 40 chars) with no operators, it's junk
            # e.g., "planet", "εwith respect to the Sun"
            if len(plain) < 40:
                return True
            return False

        def _is_constant_definition(latex):
            """Detect parameter/constant definitions like:
            R = 1.496 × 10⁸ km, μMars = 43,050 km³/s²
            Pattern: <variable> = <number> <optional unit> (nothing else)
            """
            plain = _to_plain(latex)
            # Match: variable_name = number [unit]
            # Variable: 1-30 chars (letters, digits, subscripts, Greek)
            # Value: number with optional scientific notation
            # Unit: optional short text (km, km/s, km³/s², rad/s, etc.)
            m = _mre.match(
                r'^[A-Za-zα-ωΑ-Ω∞μεδσρ][\w\s,α-ωΑ-Ω∞μεδσρ]{0,30}'  # variable part
                r'\s*=\s*'                                               # equals
                r'-?[\d,]+(?:\.[\d]+)?'                                  # number
                r'(?:\s*[×x]\s*10[\^]?[\d{}\-]+)?'                     # optional ×10^n
                r'\s*'
                r'(?:[a-zA-Z/°³²\s]{0,15})?'                           # optional short unit
                r'\s*$',
                plain
            )
            return m is not None

        def _normalize_for_dedup(latex):
            """Normalize equation for duplicate detection."""
            s = _to_plain(latex).lower()
            # Remove all whitespace
            s = _mre.sub(r'\s+', '', s)
            # Remove parentheses variations
            s = s.replace('(', '').replace(')', '')
            s = s.replace('[', '').replace(']', '')
            # Remove trailing units for comparison
            s = _mre.sub(r'(?:km|m|s|rad|hrs?|hours?|days?|months?)[/\d³²]*$', '', s)
            return s

        # ══════════════════════════════════════════════════════════════
        # Step 1: Merge continuation lines (=...) into previous equation
        # ══════════════════════════════════════════════════════════════
        merged = []
        for idx, eq in enumerate(eqs):
            latex = eq["readable_text"]
            is_cont = _is_continuation(latex)
            # Log every equation to debug merge issues
            if not is_cont:
                logger.debug("EQ[%d] PARENT: %.60s", idx, _to_plain(latex))
            else:
                logger.debug("EQ[%d] CONT  : %.60s", idx, _to_plain(latex))
            if is_cont and merged:
                prev = merged[-1]
                prev["readable_text"] = prev["readable_text"] + " \\\\ " + latex
                if eq.get("omml_xml"):
                    prev["omml_xml"] = (prev.get("omml_xml") or "") + eq["omml_xml"]
            else:
                merged.append(dict(eq))

        logger.info("Equation merge: %d → %d after merging continuations", len(eqs), len(merged))

        # ══════════════════════════════════════════════════════════════
        # Step 2: Classify & filter junk
        # ══════════════════════════════════════════════════════════════
        classified = []
        for eq in merged:
            latex = eq["readable_text"]

            # Filter junk (text fragments, placeholders)
            if _is_junk(latex):
                logger.info("Equation JUNK removed: %s", _to_plain(latex)[:80])
                continue

            # Classify as constant definition
            if _is_constant_definition(latex):
                eq["eq_classification"] = "definition"
                logger.info("Equation classified as DEFINITION: %s", _to_plain(latex)[:80])
            else:
                eq["eq_classification"] = "equation"

            classified.append(eq)

        logger.info("Equation classify: %d → %d after removing junk", len(merged), len(classified))

        # ══════════════════════════════════════════════════════════════
        # Step 3: Deduplicate
        # ══════════════════════════════════════════════════════════════
        seen = set()
        deduped = []
        for eq in classified:
            norm = _normalize_for_dedup(eq["readable_text"])
            if norm in seen:
                logger.info("Equation DUPLICATE removed: %s", _to_plain(eq["readable_text"])[:80])
                continue
            seen.add(norm)
            deduped.append(eq)

        logger.info("Equation dedup: %d → %d after removing duplicates", len(classified), len(deduped))
        return deduped

    equations = _merge_classify_filter(equations)

    # ── If DB has no valid equations, re-extract from DOCX (picks up table cells) ──
    if not equations:
        logger.info("No valid equations in DB, attempting live re-extraction from DOCX")
        file_path = _resolve_file_path(doc.get("file_path", ""))
        working_path = session.get("working_doc_path", "")
        docx_path = working_path if working_path and os.path.exists(working_path) else file_path

        if docx_path and os.path.exists(docx_path):
            try:
                from app.services.document_service import DOCXParser
                parser = DOCXParser(docx_path)
                fresh_equations = parser._extract_equations()
                logger.info("Live re-extraction found %d equations from %s", len(fresh_equations), docx_path)

                # Update DB with fresh equations for future use
                eq_dicts = []
                for feq in fresh_equations:
                    eq_dict = {
                        "equation_id": feq.equation_id,
                        "latex": feq.latex,
                        "raw_omml": feq.raw_omml,
                        "position": {"page": feq.position.page, "paragraph": feq.position.paragraph},
                        "number": feq.number,
                    }
                    eq_dicts.append(eq_dict)

                    readable = feq.latex or ""

                    equations.append({
                        "para_idx": feq.position.paragraph if feq.position else 0,
                        "omml_xml": feq.raw_omml or "",
                        "readable_text": readable,
                    })

                # Persist to DB so next time we don't re-extract
                if eq_dicts:
                    from bson import ObjectId as _ObjId
                    doc_oid = _ObjId(doc["id"]) if "id" in doc else doc.get("_id")
                    if doc_oid:
                        await doc_repo.collection.update_one(
                            {"_id": doc_oid},
                            {"$set": {"equations": eq_dicts}},
                        )
                        logger.info("Updated DB with %d fresh equations", len(eq_dicts))

            except Exception as extract_err:
                logger.warning("Live equation re-extraction failed: %s", extract_err)

        # Apply merge-classify-filter to re-extracted equations too
        equations = _merge_classify_filter(equations)

    if not equations:
        return {"session_id": session_id, "equations": [], "message": "No equations found in document."}

    # Send equations to GPT for analysis (batch them to reduce API calls)
    from openai import AsyncOpenAI

    client = AsyncOpenAI(api_key=(await _get_openai_key(db)))

    # Build a summary of all equations for GPT
    eq_summary_lines = []
    for i, eq in enumerate(equations, start=1):
        eq_summary_lines.append(f"Equation {i}: {eq['readable_text']}")

    # ── GPT analysis (batch if > 30 equations) ──────────────────────────
    from openai import AsyncOpenAI
    import json as _json

    BATCH_SIZE = 30
    gpt_analyses = {}

    gpt_prompt_template = (
        "You are a mathematics and scientific notation expert reviewing equations "
        "extracted from a textbook document.\n\n"
        "IMPORTANT RULES:\n"
        "- A real equation is ANY mathematical expression: fractions, trig functions, "
        "matrices, integrals, inequalities, angle brackets, floor/ceiling notation, "
        "piecewise functions, etc. Even a simple expression like \\sin(x) or "
        "\\left\\langle x \\right\\rangle IS a real equation. "
        "Matrices (\\begin{pmatrix}...) ARE real equations. "
        "Do NOT mark these as not_equation.\n"
        "- Only mark as not_equation: alphabet listings, Word UI text, "
        "placeholder text like 'Type equation here', or comma-separated "
        "single letters/symbols that are clearly a reference palette.\n\n"
        "For each REAL equation, classify it into one of three categories:\n\n"
        "1. **\"outdated\"** — ONLY for things that are WRONG or DEPRECATED:\n"
        "   - \\text{} wrapping standard math functions: \\text{sin} must be \\sin, "
        "\\text{cos} must be \\cos, \\text{tan} must be \\tan, "
        "\\text{log} must be \\log, etc.\n"
        "   - Missing spaces that create INVALID commands: \\partialx is not "
        "a valid LaTeX command (must be \\partial x), \\leqy is not valid "
        "(must be \\leq y)\n"
        "   - Formulas that are mathematically superseded or incorrect\n\n"
        "   CRITICAL — \\text{} on variable names and labels is NORMAL, NOT an error:\n"
        "   These equations come from Word OMML conversion, which wraps ALL normal "
        "text runs in \\text{}. The following uses of \\text{} are CORRECT and must "
        "NOT be flagged as outdated or errors:\n"
        "   - Variable names: \\text{V}, \\text{R}, \\text{a}, \\text{m}, \\text{TOF}\n"
        "   - Descriptive labels: \\text{Earth}, \\text{mission}, \\text{boost}, "
        "\\text{Sun}, \\text{Moon}, \\text{transfer}, \\text{target}, \\text{park}, "
        "\\text{retro}, \\text{planet}, \\text{SOI}, \\text{lead}, \\text{final}\n"
        "   - Units and text in equations: \\text{ km}, \\text{ and }, \\text{mass}\n"
        "   - Numbers and operators: \\text{1}, \\text{+}, \\text{2}, \\text{=}\n"
        "   - ANY other \\text{} wrapping letters, words, or short labels\n"
        "   Do NOT suggest removing \\text{} from these — it is how OMML equations "
        "are represented in LaTeX and is perfectly valid.\n\n"
        "2. **\"formatting\"** — for OPTIONAL style improvements:\n"
        "   - Style preferences (\\times vs \\cdot — both valid)\n"
        "   - Alternative valid notations (1/2 vs \\frac{1}{2} — both valid)\n"
        "   - Removing optional \\left(\\right) delimiters\n"
        "   - Cosmetic spacing\n\n"
        "3. **\"ok\"** — equation is correct, no issues\n\n"
        "Respond in JSON: a list of objects with keys:\n"
        "- \"equation_number\" (int)\n"
        "- \"not_equation\" (bool) — true ONLY if not a real math expression\n"
        "- \"category\" (str) — \"outdated\", \"formatting\", or \"ok\"\n"
        "- \"suggested_update\" (str) — full corrected LaTeX (empty if ok)\n"
        "- \"reason\" (str) — brief explanation (empty if ok)\n\n"
        "CRITICAL RULES:\n"
        "- \\times is NOT outdated (it is a valid multiplication symbol)\n"
        "- 1/2 inline is NOT outdated (it is valid inline fraction notation)\n"
        "- \\left(\\frac{a}{b}\\right) is NOT outdated (parentheses have meaning)\n"
        "- \\text{sin} should be \\sin, \\text{cos} should be \\cos (math operator rule)\n"
        "- \\text{} on variable names, labels, units, numbers, operators is NORMAL "
        "and NOT an error — these come from Word OMML conversion. "
        "Do NOT flag \\text{V}, \\text{Earth}, \\text{mass}, \\text{1}, \\text{+} etc.\n"
        "- Only flag an equation as outdated if the MATH CONTENT itself is wrong, "
        "deprecated, or uses incorrect formulas — NOT because of \\text{} wrappers\n"
        "- NEVER spell out Greek letters — always use the original symbol: "
        "use Δ not Delta, α not alpha, β not beta, γ not gamma, θ not theta, "
        "λ not lambda, μ not mu, σ not sigma, Σ not Sigma, π not pi, Ω not Omega, "
        "ε not epsilon, φ not phi, ψ not psi, ω not omega, ρ not rho, τ not tau, "
        "η not eta, ξ not xi, ζ not zeta, χ not chi, ν not nu, κ not kappa. "
        "In LaTeX use \\Delta, \\alpha, \\beta etc. — never the English name.\n"
        "- NEVER convert subscripts/superscripts to plain text: "
        "keep v₁ not v1, x² not x2, Δv not Delta v\n\n"
    )

    async def _analyze_batch(batch_lines, batch_offset):
        """Send a batch of equations to GPT and return parsed analyses."""
        batch_text = "\n".join(batch_lines)
        try:
            resp = await client.chat.completions.create(
                model="gpt-4o",
                messages=[{
                    "role": "user",
                    "content": gpt_prompt_template + batch_text,
                }],
                temperature=0.0,
                max_tokens=4000,
            )
            raw = resp.choices[0].message.content or ""
            cleaned = raw.strip()
            if cleaned.startswith("```"):
                cleaned = cleaned.split("\n", 1)[1] if "\n" in cleaned else cleaned
                cleaned = cleaned.rsplit("```", 1)[0]
            parsed_list = _json.loads(cleaned)
            items = parsed_list if isinstance(parsed_list, list) else parsed_list.get("equations", [])
            for item in items:
                num = item.get("equation_number", 0)
                gpt_analyses[num] = item
        except (_json.JSONDecodeError, Exception) as e:
            logger.warning("GPT equation batch analysis failed: %s", e)

    try:
        # Split into batches and run in parallel
        batches = []
        for start in range(0, len(eq_summary_lines), BATCH_SIZE):
            batch = eq_summary_lines[start:start + BATCH_SIZE]
            batches.append((batch, start))

        if len(batches) == 1:
            await _analyze_batch(batches[0][0], batches[0][1])
        else:
            await asyncio.gather(
                *[_analyze_batch(b, off) for b, off in batches],
                return_exceptions=True
            )

        logger.info("GPT equation analysis complete: %d equations analyzed, %d results",
                    len(equations), len(gpt_analyses))
    except Exception as gpt_err:
        logger.warning("GPT equation analysis failed: %s", gpt_err)
        for i in range(1, len(equations) + 1):
            gpt_analyses[i] = {"suggested_update": f"Analysis unavailable: {gpt_err}", "is_outdated": False}

    result_patches = []
    for eq_num, eq in enumerate(equations, start=1):
        analysis_item = gpt_analyses.get(eq_num, {})

        # GPT may flag some items as not_equation — ignore that flag and show all
        # equations to the user (the extraction already filters genuine junk)
        if analysis_item.get("not_equation", False):
            logger.info("Equation %d: GPT said not_equation, keeping anyway: %s",
                        eq_num, eq["readable_text"][:80])
            analysis_item["not_equation"] = False
            if "category" not in analysis_item:
                analysis_item["category"] = "ok"

        category = analysis_item.get("category", "ok")

        # ── Post-GPT safety net: catch \text{} misuse GPT might miss ───
        original_text = eq["readable_text"]
        if category == "ok":
            import re as _re_eq
            # \text{sin}, \text{cos}, \text{tan}, \text{log}, etc.
            text_func_match = _re_eq.search(r'\\text\{(sin|cos|tan|log|ln|exp|lim|max|min|det|gcd)\}', original_text)
            # \text{+}, \text{-}, \text{1}, \text{2}, etc. (operators/numbers in text mode)
            text_num_op_match = _re_eq.search(r'\\text\{[0-9+\-*/=<>]\}', original_text)
            # \partialx, \leqy — command run into variable with no space
            invalid_cmd_match = _re_eq.search(r'\\(partial|leq|geq|neq|approx|equiv)[a-zA-Z]', original_text)

            if text_func_match:
                func = text_func_match.group(1)
                category = "outdated"
                if not analysis_item.get("suggested_update"):
                    analysis_item["suggested_update"] = original_text.replace(
                        f"\\text{{{func}}}", f"\\{func}"
                    )
                analysis_item["reason"] = analysis_item.get("reason") or (
                    f"Use \\{func} instead of \\text{{{func}}} for proper math notation."
                )
            elif text_num_op_match:
                category = "outdated"
                if not analysis_item.get("suggested_update"):
                    fixed = _re_eq.sub(r'\\text\{([0-9+\-*/=<>])\}', r'\1', original_text)
                    analysis_item["suggested_update"] = fixed
                analysis_item["reason"] = analysis_item.get("reason") or (
                    "\\text{} around numbers/operators is unnecessary in math mode."
                )
            elif invalid_cmd_match:
                category = "outdated"
                cmd = invalid_cmd_match.group(1)
                analysis_item["reason"] = analysis_item.get("reason") or (
                    f"Missing space after \\{cmd} makes the command invalid."
                )

        # ── Filter out no-op suggestions (suggested_update == original) ──
        suggested = analysis_item.get("suggested_update", "")
        original_normalized = eq["readable_text"].strip().replace(" ", "")
        suggested_normalized = suggested.strip().replace(" ", "") if suggested else ""
        if category == "formatting" and (
            not suggested
            or suggested_normalized == original_normalized
        ):
            category = "ok"
            logger.info("Equation %d: suggestion identical to original — downgraded to ok", eq_num)

        # is_outdated = True ONLY for mathematically wrong / deprecated notation
        is_outdated = category == "outdated"
        # has_suggestion = True for both outdated AND formatting suggestions
        has_suggestion = category in ("outdated", "formatting")

        patch_doc = {
            "patch_id": str(uuid4()),
            "session_id": session_id,
            "type": "equation",
            "status": "pending",
            "equation_number": eq_num,
            "original_text": eq["readable_text"],
            "original_omml": eq["omml_xml"],
            "suggested_update": analysis_item.get("suggested_update", "") if has_suggestion else "",
            "is_outdated": is_outdated,
            "category": category,
            "eq_classification": eq.get("eq_classification", "equation"),
            "reason": analysis_item.get("reason", "") if has_suggestion else "",
            "para_idx": eq["para_idx"],
            "created_at": datetime.utcnow().isoformat(),
        }
        result_patches.append(patch_doc)

    if result_patches:
        await db.media_patches.insert_many(result_patches)

    for rp in result_patches:
        if "_id" in rp:
            rp["_id"] = str(rp["_id"])

    return {"session_id": session_id, "equations": result_patches, "cached": False}


def _is_reference_table(rows: list) -> bool:
    """Detect tables that are reference/syntax material (symbol lists, shortcuts, etc.)
    These should not be analyzed for outdated data."""
    if not rows:
        return False

    # Check header row for reference-table patterns
    header = [str(c).lower().strip() for c in (rows[0] if rows else [])]
    header_text = ' '.join(header)

    # Symbol/syntax reference tables
    ref_headers = ['symbol', 'type', 'accent', 'grouping', 'brackets', 'code', 'use']
    ref_matches = sum(1 for h in header if any(rh in h for rh in ref_headers))
    if ref_matches >= 2:
        return True

    # Keyboard shortcut tables
    if any(kw in header_text for kw in ['shift+', 'ctrl+', 'arrow key', 'shortcut']):
        return True

    # Greek letter reference tables
    if any(kw in header_text for kw in ['lower case', 'upper case', 'lowercase', 'uppercase']):
        return True

    # Large symbol-only tables (>15 rows where most cells are short LaTeX commands)
    if len(rows) > 15:
        short_cells = 0
        total_cells = 0
        for row in rows:
            for cell in row:
                cell_str = str(cell).strip()
                total_cells += 1
                if len(cell_str) < 20 and ('\\' in cell_str or len(cell_str) <= 3):
                    short_cells += 1
        if total_cells > 0 and short_cells / total_cells > 0.7:
            return True

    return False


def _is_empty_or_trivial_table(rows: list) -> bool:
    """Return True if the table has no meaningful content."""
    if not rows:
        return True

    # Count non-empty cells
    non_empty = 0
    for row in rows:
        for cell in row:
            if str(cell).strip():
                non_empty += 1

    # Fewer than 2 non-empty cells = trivial
    return non_empty < 2


def _format_table_for_gpt(rows: list, max_rows: int = 25) -> str:
    """Format table content as a readable markdown-style table for GPT."""
    if not rows:
        return "(empty table)"

    display_rows = rows[:max_rows]
    # Determine column widths
    num_cols = max(len(row) for row in display_rows) if display_rows else 0
    if num_cols == 0:
        return "(empty table)"

    lines = []
    for r_idx, row in enumerate(display_rows):
        cells = []
        for c_idx in range(num_cols):
            val = str(row[c_idx]).strip() if c_idx < len(row) else ""
            cells.append(val if val else "(empty)")
        lines.append(f"| {' | '.join(cells)} |")
        if r_idx == 0:
            lines.append(f"| {' | '.join(['---'] * num_cols)} |")

    if len(rows) > max_rows:
        lines.append(f"... ({len(rows) - max_rows} more rows)")

    return '\n'.join(lines)


@router.get("/{session_id}/table-analysis")
async def table_analysis(session_id: str, user=Depends(get_current_user_dep)):
    """Analyze all tables in the DOCX for outdated data."""
    from uuid import uuid4

    db = get_database()
    session_repo = SessionRepository(db)
    doc_repo = DocumentRepository(db)

    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    # Return cached results if they exist
    existing = await db.media_patches.find(
        {"session_id": session_id, "type": "table"}
    ).to_list(length=500)
    if existing:
        for e in existing:
            e["_id"] = str(e["_id"])
        return {"session_id": session_id, "tables": existing, "cached": True}

    # ── Read tables from DB ──
    doc = await doc_repo.find_with_media(session["document_id"])
    if not doc:
        raise HTTPException(404, "Document not found")

    doc_title = doc.get("original_filename", "") or doc.get("filename", "")

    # ── ALWAYS re-extract tables from the ORIGINAL DOCX file ──
    # DB cache can be corrupted by previously approved GPT updates baked into
    # the working copy, or stale from old extraction bugs.  Fresh extraction
    # from the untouched original is fast and guarantees clean data.
    file_path = _resolve_file_path(doc.get("file_path", ""))
    db_tables = doc.get("tables", [])

    if file_path and os.path.exists(file_path):
        try:
            from app.services.document_service import DOCXParser
            parser = DOCXParser(file_path)
            fresh_tables = parser._extract_tables()
            logger.info("Table extraction from original DOCX: %d tables from %s",
                        len(fresh_tables), file_path)

            # Convert to dicts and update DB
            tbl_dicts = []
            for ft in fresh_tables:
                tbl_dicts.append({
                    "table_id": ft.table_id,
                    "caption": ft.caption,
                    "content": ft.content,
                    "position": {"page": ft.position.page, "paragraph": ft.position.paragraph} if ft.position else {},
                    "number": ft.number,
                })

            # Persist to DB
            if tbl_dicts:
                from bson import ObjectId as _ObjId
                doc_oid = _ObjId(doc["id"]) if "id" in doc else doc.get("_id")
                if doc_oid:
                    await doc_repo.collection.update_one(
                        {"_id": doc_oid},
                        {"$set": {"tables": tbl_dicts}},
                    )
                    logger.info("Updated DB with %d fresh tables", len(tbl_dicts))

            db_tables = tbl_dicts
        except Exception as extract_err:
            logger.warning("Live table re-extraction failed: %s", extract_err)

    # ── Filter tables: skip empty, trivial, and reference tables ──
    tables_data = []
    skipped_empty = 0
    skipped_ref = 0
    for t_idx, tbl in enumerate(db_tables):
        content = tbl.get("content", [])
        caption = tbl.get("caption") or ""
        doc_number = tbl.get("number") or ""

        if _is_empty_or_trivial_table(content):
            skipped_empty += 1
            logger.info("Table %d skipped (empty/trivial)", t_idx + 1)
            continue

        if _is_reference_table(content):
            skipped_ref += 1
            logger.info("Table %d skipped (reference/syntax table): header=%s",
                        t_idx + 1, [str(c)[:30] for c in content[0]] if content else [])
            continue

        tables_data.append({
            "table_idx": t_idx,
            "content": content,
            "caption": caption,
            "doc_number": doc_number,
            "display_number": doc_number or str(t_idx + 1),
        })

    if skipped_empty or skipped_ref:
        logger.info("Table analysis: %d total, %d skipped (empty), %d skipped (reference), %d to analyze",
                    len(db_tables), skipped_empty, skipped_ref, len(tables_data))

    if not tables_data:
        return {"session_id": session_id, "tables": [], "message": "No data tables found to analyze."}

    # ── GPT analysis ──
    import asyncio
    import json as _json
    from openai import AsyncOpenAI

    client = AsyncOpenAI(api_key=(await _get_openai_key(db)))

    async def _analyze_single_table(t_num: int, tdata: dict):
        content_str = _format_table_for_gpt(tdata["content"])
        caption = tdata.get("caption") or "(no caption)"
        display_num = tdata.get("display_number", str(t_num))

        analysis_text = ""
        cell_updates = []
        is_outdated = False
        try:
            gpt_resp = await client.chat.completions.create(
                model="gpt-4o",
                messages=[{
                    "role": "user",
                    "content": (
                        f"You are analyzing Table {display_num} from a document "
                        f"titled \"{doc_title}\".\n"
                        f"Caption: {caption}\n\n"
                        f"{content_str}\n\n"
                        "Your task:\n"
                        "1. Determine if this table contains FACTUAL DATA that could "
                        "become outdated (statistics, dates, measurements, names, "
                        "company data, mission details, etc.)\n"
                        "2. If the table is a REFERENCE table (syntax examples, formulas, "
                        "math demonstrations, notation guides), it is NOT outdated.\n"
                        "3. If the table contains real data, check if any values are "
                        "outdated based on your knowledge. Suggest SPECIFIC cell updates.\n\n"
                        "CRITICAL RULES:\n"
                        "- Set is_outdated to true ONLY when you have concrete cell_updates "
                        "to suggest. If you have NO updates, is_outdated MUST be false.\n"
                        "- The new_value MUST be the EXACT text to put in the cell — "
                        "a specific name, number, date, measurement, or factual value.\n"
                        "- ABSOLUTE PROHIBITION: NEVER use phrases like:\n"
                        "  * 'Updated X based on Y'\n"
                        "  * 'Current technology standards'\n"
                        "  * 'Updated temperature levels, resolution, and location accuracy "
                        "based on current technology'\n"
                        "  * 'Updated availability percentage and outage duration based on "
                        "current standards'\n"
                        "  * 'Updated mission duration based on current mission planning'\n"
                        "  * Any sentence that DESCRIBES a change rather than BEING the value\n"
                        "- The new_value should look like cell data, NOT a sentence about the data.\n"
                        "- If you cannot provide the EXACT replacement value, skip that cell "
                        "entirely. An empty cell_updates list is perfectly fine.\n"
                        "- Only suggest updates for cells where you KNOW the specific "
                        "correct current value (e.g., a renamed satellite, a corrected date, "
                        "a known successor mission).\n"
                        "- Examples of GOOD new_value: 'Intelsat', '2024', 'James Webb "
                        "Space Telescope', '$2.5B', '150 km resolution', 'International "
                        "Space Station'\n"
                        "- Examples of BAD new_value (NEVER do this): 'Updated based on "
                        "current data', 'Modern equivalent', 'Current technology standards', "
                        "'Updated coverage area based on current satellite capabilities'\n"
                        "- Fix ONLY: misspelled names, renamed organizations, cancelled "
                        "missions replaced by known successors, factually wrong dates.\n"
                        "- Do NOT update cells that contain descriptions, requirements, "
                        "section references, or procedural text — these are NOT data.\n"
                        "- Row 0 is usually the header row — never update headers.\n"
                        "- Never spell out Greek letters — always preserve original symbols "
                        "(Δ, α, β, θ, Σ, π, etc.) and subscripts/superscripts exactly as written.\n\n"
                        "Respond ONLY in JSON (no markdown code blocks):\n"
                        "{\n"
                        "  \"analysis\": \"Brief description of the table and findings\",\n"
                        "  \"is_outdated\": true/false,\n"
                        "  \"cell_updates\": [\n"
                        "    {\"row\": 1, \"col\": 0, \"old_value\": \"old\", "
                        "\"new_value\": \"specific concrete replacement\"}\n"
                        "  ]\n"
                        "}\n"
                    ),
                }],
                temperature=0.0,
                max_tokens=2000,
            )
            raw = gpt_resp.choices[0].message.content or ""

            # ── Robust JSON parsing ──
            cleaned = raw.strip()
            # Strip markdown code fences
            if cleaned.startswith("```"):
                # Remove opening fence (```json or ```)
                first_newline = cleaned.find('\n')
                if first_newline != -1:
                    cleaned = cleaned[first_newline + 1:]
                # Remove closing fence
                if cleaned.rstrip().endswith("```"):
                    cleaned = cleaned.rstrip()[:-3].rstrip()

            try:
                parsed = _json.loads(cleaned)
                analysis_text = parsed.get("analysis", "")
                cell_updates = parsed.get("cell_updates", [])
                is_outdated = parsed.get("is_outdated", False)
            except _json.JSONDecodeError:
                # Try to extract JSON from mixed text
                import re as _re_json
                json_match = _re_json.search(r'\{[\s\S]*\}', cleaned)
                if json_match:
                    try:
                        parsed = _json.loads(json_match.group())
                        analysis_text = parsed.get("analysis", "")
                        cell_updates = parsed.get("cell_updates", [])
                        is_outdated = parsed.get("is_outdated", False)
                    except _json.JSONDecodeError:
                        analysis_text = cleaned
                else:
                    # Last resort: use raw text but strip any JSON artifacts
                    analysis_text = cleaned
                    if analysis_text.startswith('{') or analysis_text.startswith('['):
                        analysis_text = "Analysis could not be parsed. Please re-analyze."

            # ── Post-GPT safety: filter out vague/non-concrete updates ──
            if cell_updates:
                import re as _re_vague

                # Patterns that indicate a vague description, NOT a concrete value
                vague_phrases = _re_vague.compile(
                    r'(?i)'
                    r'(?:updated?\s+.+\s+based\s+on)'   # "Updated X based on Y"
                    r'|(?:based\s+on\s+current)'         # "based on current ..."
                    r'|(?:current\s+(?:technology|standards?|data|capabilities?|infrastructure|planning|mapping))'
                    r'|(?:modern\s+equivalent)'
                    r'|(?:updated?\s+(?:temperature|coverage|data|mission|availability|map|number|grid))'
                    r'|(?:revised\s+(?:to|based|for))'
                )

                filtered = []
                for upd in cell_updates:
                    new_val = str(upd.get("new_value", "")).strip()
                    old_val = str(upd.get("old_value", "")).strip()
                    # Skip if new_value is empty or same as old
                    if not new_val or new_val.lower() == old_val.lower():
                        continue
                    # Skip if new_value matches vague description patterns
                    if vague_phrases.search(new_val):
                        logger.info("Table %s: filtered vague update: %r → %r",
                                    display_num, old_val[:40], new_val[:60])
                        continue
                    filtered.append(upd)
                cell_updates = filtered

            # ── Enforce: no updates = not outdated ──
            if not cell_updates:
                is_outdated = False

        except Exception as gpt_err:
            logger.warning("GPT table analysis failed for table %s: %s", display_num, gpt_err)
            analysis_text = f"Analysis unavailable: {gpt_err}"

        return {
            "patch_id": str(uuid4()),
            "session_id": session_id,
            "type": "table",
            "status": "pending",
            "table_number": t_num,
            "table_idx": tdata.get("table_idx", t_num - 1),  # original DOCX table index
            "original_content": tdata["content"],
            "caption": tdata.get("caption", ""),
            "doc_table_number": tdata.get("display_number", ""),
            "cell_updates": cell_updates,
            "analysis": analysis_text,
            "is_outdated": is_outdated,
            "created_at": datetime.utcnow().isoformat(),
        }

    # Process all tables in parallel (batches of 3 to avoid rate limits)
    result_patches = []
    BATCH_SIZE = 3
    for batch_start in range(0, len(tables_data), BATCH_SIZE):
        batch = tables_data[batch_start:batch_start + BATCH_SIZE]
        tasks = [
            _analyze_single_table(batch_start + i + 1, tdata)
            for i, tdata in enumerate(batch)
        ]
        batch_results = await asyncio.gather(*tasks, return_exceptions=True)
        for res in batch_results:
            if isinstance(res, Exception):
                logger.error("Table analysis failed: %s", res)
            else:
                result_patches.append(res)

    if result_patches:
        await db.media_patches.insert_many(result_patches)

    for rp in result_patches:
        if "_id" in rp:
            rp["_id"] = str(rp["_id"])

    return {"session_id": session_id, "tables": result_patches, "cached": False}


@router.put("/{session_id}/media-patches/{patch_id}")
async def review_media_patch(
    session_id: str,
    patch_id: str,
    request: Request,
    user=Depends(get_current_user_dep),
):
    """Approve or reject a media patch (figure / equation / table)."""
    db = get_database()
    session_repo = SessionRepository(db)

    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    body = await request.json()
    action = body.get("action", "").lower()
    if action not in ("approve", "reject"):
        raise HTTPException(400, "action must be 'approve' or 'reject'")

    existing = await db.media_patches.find_one({
        "patch_id": patch_id,
        "session_id": session_id,
    })
    if not existing:
        raise HTTPException(404, "Media patch not found")

    update_fields: Dict = {"status": "approved" if action == "approve" else "rejected"}

    # For figures, allow selecting a replacement candidate
    selected_replacement = body.get("selected_replacement")
    if selected_replacement and existing.get("type") == "figure":
        update_fields["selected_replacement"] = selected_replacement

    await db.media_patches.update_one(
        {"patch_id": patch_id, "session_id": session_id},
        {"$set": update_fields},
    )

    return {
        "patch_id": patch_id,
        "session_id": session_id,
        "status": update_fields["status"],
        "message": f"Media patch {action}d successfully.",
    }


@router.post("/{session_id}/figure-upload/{patch_id}")
async def upload_figure_replacement(
    session_id: str,
    patch_id: str,
    request: Request,
    user=Depends(get_current_user_dep),
):
    """Upload a user's own image as a figure replacement."""
    db = get_database()
    session_repo = SessionRepository(db)

    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    existing = await db.media_patches.find_one({
        "patch_id": patch_id,
        "session_id": session_id,
        "type": "figure",
    })
    if not existing:
        raise HTTPException(404, "Figure patch not found")

    # Parse multipart form data
    form = await request.form()
    file = form.get("file")
    if file is None:
        raise HTTPException(400, "No file uploaded")

    # Validate file type
    content_type = getattr(file, "content_type", "") or ""
    allowed_types = {"image/png", "image/jpeg", "image/jpg", "image/webp", "image/gif"}
    if content_type and content_type not in allowed_types:
        raise HTTPException(400, f"Invalid file type: {content_type}. Allowed: PNG, JPEG, WebP, GIF")

    # Save file
    upload_dir = os.path.join(settings.UPLOAD_DIR, "figure_uploads", session_id)
    os.makedirs(upload_dir, exist_ok=True)

    filename = getattr(file, "filename", "image.png") or "image.png"
    ext = os.path.splitext(filename)[1] or ".png"
    saved_filename = f"{patch_id}{ext}"
    saved_path = os.path.join(upload_dir, saved_filename)

    contents = await file.read()
    with open(saved_path, "wb") as f:
        f.write(contents)

    # Build the URL for the frontend to display and export to use.
    # Append a cache-busting timestamp so re-uploads of the same patch_id
    # (same filename on disk) force the browser to fetch the new image
    # instead of serving from cache.
    import time as _time
    cache_bust = int(_time.time())
    image_url = f"/uploads/figures/{session_id}/{saved_filename}?t={cache_bust}"

    # Update the media patch with the user-uploaded replacement
    user_replacement = {
        "url": image_url,
        "thumbnail_url": image_url,
        "title": filename,
        "source": "User Upload",
        "is_user_upload": True,
        "local_path": saved_path,
    }

    await db.media_patches.update_one(
        {"patch_id": patch_id, "session_id": session_id},
        {"$set": {
            "selected_replacement": user_replacement,
            "user_uploaded_replacement": user_replacement,
        }},
    )

    return {
        "patch_id": patch_id,
        "session_id": session_id,
        "replacement": user_replacement,
        "message": "Image uploaded successfully",
    }


@router.delete("/{session_id}/media-patches/{media_type}")
async def clear_media_patches(
    session_id: str,
    media_type: str,
    user=Depends(get_current_user_dep),
):
    """Clear cached media patches of a given type (figure/equation/table) to allow re-analysis."""
    if media_type not in ("figure", "equation", "table"):
        raise HTTPException(400, "media_type must be 'figure', 'equation', or 'table'")

    db = get_database()
    session_repo = SessionRepository(db)

    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(403, "Not authorized")

    result = await db.media_patches.delete_many({
        "session_id": session_id,
        "type": media_type,
    })

    return {
        "session_id": session_id,
        "type": media_type,
        "deleted": result.deleted_count,
        "message": f"Cleared {result.deleted_count} {media_type} patches. Ready for re-analysis.",
    }


# ── Spell & Grammar Check ──────────────────────────────────────────────────────

class SpellFixPayload(BaseModel):
    approved_issue_ids: List[str] = []


@router.post("/{session_id}/run-spell-check")
async def run_spell_check(
    session_id: str,
    user=Depends(get_current_user_dep),
):
    """
    Run GPT-based spell and grammar check on the document's text content.
    Returns a list of SpellIssue objects for admin review.
    Preserves technical terms, equations, and proper nouns.
    """
    import json as _json
    from openai import AsyncOpenAI

    db = get_database()
    session_repo = SessionRepository(db)
    doc_repo = DocumentRepository(db)

    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Not authorized")

    doc = await doc_repo.find_with_media(session["document_id"])
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    text_content = doc.get("text_content", "")
    if not text_content:
        return {"session_id": session_id, "issues": [], "total_issues": 0, "status": "complete"}

    # Split into chunks of ~2000 chars to keep GPT calls manageable
    chunk_size = 2000
    paragraphs = [p.strip() for p in text_content.split("\n") if p.strip()]
    chunks = []  # list of (start_para_idx, chunk_text)
    current_chunk = []
    current_len = 0
    start_idx = 0

    for idx, para in enumerate(paragraphs):
        current_chunk.append(para)
        current_len += len(para)
        if current_len >= chunk_size:
            chunks.append((start_idx, "\n".join(current_chunk)))
            current_chunk = []
            current_len = 0
            start_idx = idx + 1

    if current_chunk:
        chunks.append((start_idx, "\n".join(current_chunk)))

    client = AsyncOpenAI(api_key=(await _get_openai_key(db)))
    all_issues = []

    system_prompt = (
        "You are a professional copy editor for technical textbooks. "
        "Check the given text for spelling errors, grammar mistakes, and clarity issues. "
        "STRICT RULES:\n"
        "1. Do NOT flag technical terms, scientific notation, equations, Greek letters, "
        "   unit abbreviations (km/s, AU, etc.), or proper nouns (NASA, Hubble, etc.).\n"
        "2. Only flag genuine errors — do not suggest style rewrites.\n"
        "3. Keep corrections minimal — change only what is clearly wrong.\n"
        "4. For each issue return a JSON object with:\n"
        '   {"original_text": "<exact phrase from text>", '
        '"suggested_text": "<corrected phrase>", '
        '"issue_type": "spelling"|"grammar"|"clarity", '
        '"explanation": "<brief reason>"}\n'
        "5. Return ONLY a JSON array. If no issues found, return []."
    )

    for start_para, chunk_text in chunks:
        try:
            resp = await client.chat.completions.create(
                model=settings.GPT_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": chunk_text},
                ],
                temperature=0.0,
                max_tokens=2000,
            )
            raw = resp.choices[0].message.content.strip()
            if raw.startswith("```"):
                raw = re.sub(r'^```(?:json)?\s*', '', raw)
                raw = re.sub(r'\s*```$', '', raw)
            found = _json.loads(raw)
            for item in found:
                all_issues.append({
                    "issue_id": str(uuid.uuid4()),
                    "paragraph_idx": start_para,
                    "original_text": item.get("original_text", ""),
                    "suggested_text": item.get("suggested_text", ""),
                    "issue_type": item.get("issue_type", "grammar"),
                    "explanation": item.get("explanation", ""),
                    "status": "pending",
                })
        except Exception as e:
            logger.warning("Spell check chunk failed: %s", e)
            continue

    # Resolve document _id for update
    from bson import ObjectId as _ObjId
    doc_oid = _ObjId(doc["id"]) if "id" in doc else doc.get("_id")

    await doc_repo.collection.update_one(
        {"_id": doc_oid},
        {"$set": {
            "spell_issues": all_issues,
            "spell_check_status": "pending",
            "updated_at": datetime.utcnow(),
        }},
    )

    # Also store on session for quick access
    session_oid = _ObjId(session["id"]) if "id" in session else session.get("_id")
    await db.sessions.update_one(
        {"_id": session_oid},
        {"$set": {
            "spell_issues": all_issues,
            "status": "spell_check_pending",
            "updated_at": datetime.utcnow(),
        }},
    )

    return {
        "session_id": session_id,
        "issues": all_issues,
        "total_issues": len(all_issues),
        "status": "complete",
    }


@router.post("/{session_id}/apply-spell-fixes")
async def apply_spell_fixes(
    session_id: str,
    payload: SpellFixPayload,
    user=Depends(get_current_user_dep),
):
    """Apply approved spell/grammar fixes to the document text_content in the DB."""
    from bson import ObjectId as _ObjId

    db = get_database()
    session_repo = SessionRepository(db)
    doc_repo = DocumentRepository(db)

    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Not authorized")

    doc = await doc_repo.find_with_media(session["document_id"])
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Spell issues stored on session (set during run-spell-check)
    spell_issues = session.get("spell_issues", [])
    approved_ids = set(payload.approved_issue_ids)

    text_content = doc.get("text_content", "")
    applied = 0

    for issue in spell_issues:
        if issue["issue_id"] in approved_ids and issue.get("original_text"):
            text_content = text_content.replace(
                issue["original_text"], issue["suggested_text"], 1
            )
            applied += 1

    doc_oid = _ObjId(doc["id"]) if "id" in doc else doc.get("_id")
    await doc_repo.collection.update_one(
        {"_id": doc_oid},
        {"$set": {"text_content": text_content, "updated_at": datetime.utcnow()}},
    )

    session_oid = _ObjId(session["id"]) if "id" in session else session.get("_id")
    await db.sessions.update_one(
        {"_id": session_oid},
        {"$set": {"status": "audit_complete", "updated_at": datetime.utcnow()}},
    )

    return {"applied": applied, "status": "audit_complete"}


@router.post("/{session_id}/skip-spell-check")
async def skip_spell_check(
    session_id: str,
    user=Depends(get_current_user_dep),
):
    """Advance session past spell check without applying any fixes."""
    from bson import ObjectId as _ObjId

    db = get_database()
    session_repo = SessionRepository(db)

    session = await session_repo.find_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    if session.get("user_id") != user["email"] and user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Not authorized")

    session_oid = _ObjId(session["id"]) if "id" in session else session.get("_id")
    await db.sessions.update_one(
        {"_id": session_oid},
        {"$set": {"status": "audit_complete", "updated_at": datetime.utcnow()}},
    )

    return {"status": "audit_complete", "message": "Spell check skipped"}